from flask import Flask, render_template_string, request, send_file, jsonify
import os
import json
import copy
import re
import platform

from werkzeug.utils import secure_filename

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

app = Flask(__name__)
app.config["UPLOAD_FOLDER"] = "uploads"
app.config["OUTPUT_FOLDER"] = "outputs"
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024  # 16 MB

os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)
os.makedirs(app.config["OUTPUT_FOLDER"], exist_ok=True)

# ================== GRAMMAR CHECK DEPENDENCIES ==================
IS_WINDOWS = platform.system() == "Windows"
WORD_AVAILABLE = False
LANGUAGE_TOOL_AVAILABLE = False

if IS_WINDOWS:
    try:
        import win32com.client
        import pythoncom
        WORD_AVAILABLE = True
    except ImportError:
        print("⚠ Install pywin32 for Word COM grammar checking")

try:
    import language_tool_python
    _grammar_tool = language_tool_python.LanguageTool("en-US")
    LANGUAGE_TOOL_AVAILABLE = True
except ImportError:
    print("⚠ Install language-tool-python for grammar checking")

# ================== COLORS ==================
HIGHLIGHT_COLORS = {
    "TITLE": WD_COLOR_INDEX.YELLOW,
    "COVER_PAGE": WD_COLOR_INDEX.YELLOW,
    "TOC_TITLE": WD_COLOR_INDEX.YELLOW,
    "TOC_HEADING_1": WD_COLOR_INDEX.GRAY_25,
    "TOC_HEADING_2": WD_COLOR_INDEX.GRAY_50,
    "TOC_HEADING_3": WD_COLOR_INDEX.DARK_YELLOW,
    "TOC_HEADING_4": WD_COLOR_INDEX.DARK_YELLOW,
    "HEADING_1": WD_COLOR_INDEX.BRIGHT_GREEN,
    "HEADING_2": WD_COLOR_INDEX.TURQUOISE,
    "HEADING_3": WD_COLOR_INDEX.PINK,
    "HEADING_4": WD_COLOR_INDEX.VIOLET,
    "HEADING_5": WD_COLOR_INDEX.TEAL,
    "HEADING_6": WD_COLOR_INDEX.DARK_BLUE,
    "PARAGRAPH": WD_COLOR_INDEX.GRAY_25,
    "LIST_ITEM": WD_COLOR_INDEX.GRAY_25,
    "TABLE": WD_COLOR_INDEX.TURQUOISE,
}

# ================== MAIN UI ==================
MAIN_HTML = r"""
<!DOCTYPE html>
<html>
<head>
    <title>Document Formatter</title>
    <style>
        body {
            font-family: Arial, sans-serif;
            background-color: #f2f4f8;
            margin: 0;
            padding: 0;
        }
        .container {
            width: 80%;
            margin: 40px auto;
            text-align: center;
        }
        h1 { margin-bottom: 30px; }
        .tabs {
            display: flex;
            justify-content: center;
            gap: 20px;
        }
        .tab-button {
            padding: 12px 25px;
            border: none;
            border-radius: 8px;
            background-color: #e0e0e0;
            cursor: pointer;
            font-size: 16px;
            font-weight: bold;
        }
        .tab-button:hover { background-color: #d0d0d0; }
        .active { background-color: #4CAF50; color: white; }
        .tab-content {
            margin-top: 40px;
            padding: 30px;
            border-radius: 10px;
            background: white;
            box-shadow: 0px 4px 10px rgba(0,0,0,0.1);
            text-align: left;
        }

        /* ── Smart Doc Formatter styles (embedded) ── */
        .sdf-box { background: linear-gradient(135deg,#f3f4f6,#e5e7eb); padding: 20px; border-radius: 12px; }
        .sdf-inner { background: white; padding: 30px; border-radius: 15px; max-height: 75vh; overflow-y: auto; }
        .sdf-inner h2 { text-align:center; margin-top:0; }
        .sdf-inner label { font-weight:bold; display:block; margin-top:6px; }
        .sdf-inner select, .sdf-inner input[type=number], .sdf-inner input[type=file] { width:100%; padding:8px; margin-bottom:10px; box-sizing:border-box; }
        .sdf-btn { width:100%; padding:14px; background:#000000; border:none; border-radius:30px; color:white; font-size:16px; cursor:pointer; margin-top:8px; }
        .sdf-btn:disabled { background:#ccc; cursor:not-allowed; }
        .sdf-info-box { background:#f0f0f0; padding:15px; border-radius:8px; margin-bottom:20px; }
        .sdf-element-list { list-style:none; padding-left:0; }
        .sdf-element-list li { padding:5px 0; }
        .sdf-element-list li.toc-item { color:#2563eb; font-weight:500; }
        .sdf-element-list li.list-item { color:#16a34a; font-weight:500; }
        .sdf-loading { text-align:center; padding:20px; color:#c7c6c1; }
        .sdf-hierarchy-info { background:#e8f4f8; padding:10px; border-left:4px solid #c7c6c1; margin:10px 0; font-size:14px; }
        .sdf-toc-section { background:#f0f9ff; padding:15px; border-radius:8px; margin:15px 0; border:2px solid #2563eb; }
        .sdf-toc-section h3 { color:#2563eb; margin-top:0; }
        .sdf-list-section { background:#f0fdf4; padding:15px; border-radius:8px; margin:15px 0; border:2px solid #16a34a; }
        .sdf-list-section h3 { color:#16a34a; margin-top:0; }
        .sdf-preview-container { display:flex; gap:15px; align-items:center; }
        .sdf-preview-box { border:1px solid #d1d5db; border-radius:6px; padding:10px 14px; background:#fafafa; min-width:220px; font-size:14px; }
        .sdf-preview-label { font-size:12px; color:#6b7280; margin-bottom:4px; }
        .sdf-subbox { background:#fff; border:1px solid #e5e7eb; border-radius:8px; padding:10px 12px; margin-top:12px; }
        .sdf-subbox h4 { margin:0 0 8px 0; }
        .sdf-small { color:#6b7280; font-size:12px; }
        .sdf-cover-section { background:#fffbeb; padding:15px; border-radius:8px; margin:15px 0; border:2px solid #f59e0b; }
        .sdf-cover-section h3 { color:#b45309; margin-top:0; }
        .sdf-checkbox-label { font-weight:normal; display:inline-flex; align-items:center; gap:6px; margin:4px 0; cursor:pointer; }
        /* ── Grammar Checker styles ── */
        .gc-btn { padding:8px 18px; border:none; border-radius:20px; font-size:14px; cursor:pointer; margin-right:10px; background:#000; color:white; }
        .gc-btn-outline { background:#e5e7eb; color:#111; }
        .gc-btn-outline:hover { background:#d0d0d0; }
        .gc-accept { background:#16a34a; color:white; }
        .gc-ignore { background:#e5e7eb; color:#111; }
        .gc-section { margin-bottom:25px; padding:15px; background:#f9fafb; border-left:5px solid #dc2626; border-radius:4px; }
        .gc-section p { margin:6px 0; }
        /* ── Alignment & Formatting styles ── */
        .al-btn { padding:12px; width:100%; background:#4f46e5; color:white; border:none; border-radius:20px; margin-top:12px; cursor:pointer; font-size:15px; }
        .al-section-title { margin-top:25px; font-weight:bold; font-size:15px; }
        .al-label { font-weight:bold; display:block; margin-top:15px; color:black; }
        .al-select { width:100%; padding:8px; margin-top:5px; color:black; border:1px solid #ccc; border-radius:6px; }
        .al-dropdown { position:relative; margin-top:10px; }
        .al-dropbtn { width:100%; padding:8px; background:white; border:1px solid #ccc; text-align:left; cursor:pointer; border-radius:6px; color:black; font-weight:500; }
        .al-dropdown-content { display:none; position:absolute; background:white; width:100%; border:1px solid #ccc; padding:10px; z-index:10; border-radius:6px; color:black; }
        .al-dropdown-content label { display:block; margin-bottom:6px; font-weight:normal; color:black; }
        .al-dropdown-content input[type=checkbox] { accent-color:#4f46e5; }
        .al-show { display:block; }
    </style>
    <script>
        function showTab(tabName) {
            document.getElementById("font").style.display = "none";
            document.getElementById("grammar").style.display = "none";
            document.getElementById("formatting").style.display = "none";
            document.getElementById(tabName).style.display = "block";
            var buttons = document.getElementsByClassName("tab-button");
            for (var i = 0; i < buttons.length; i++) buttons[i].classList.remove("active");
            document.getElementById(tabName + "-btn").classList.add("active");
        }
        window.onload = function() { showTab("font"); }
        // ══════════════════════════════════════════
        //  Grammar & Spell Checker logic
        // ══════════════════════════════════════════
        let gcApiData = null;

        async function gcUpload() {
            const file = document.getElementById('gc-file').files[0];
            if (!file) return alert('Select a .docx file');
            const fd = new FormData();
            fd.append('file', file);
            document.getElementById('gc-result').innerHTML = '<b>Analyzing...</b>';
            const res = await fetch('/gc_upload', { method: 'POST', body: fd });
            gcApiData = await res.json();
            if (!gcApiData.success) {
                document.getElementById('gc-result').innerHTML = '<p style="color:red">Error analyzing document.</p>';
                return;
            }
            gcRender('SPELLING');
        }

        function gcShowSpelling() { if (gcApiData) gcRender('SPELLING'); }
        function gcShowGrammar()  { if (gcApiData) gcRender('GRAMMAR'); }

        function gcRender(type) {
            let html = '';
            gcApiData.results.forEach(item => {
                const errs = item.error_details.filter(e => e.type === type);
                if (!errs.length) return;
                html += '<div class="gc-section" id="gc-row-' + item.id + '">' +
                    '<p><b>Wrong sentence:</b><br>' + gcEscape(item.current) + '</p>' +
                    '<p><b>Corrected:</b><br><span style="color:green">' + gcEscape(item.corrected) + '</span></p>';
                if (type === 'SPELLING') {
                    html += '<button class="gc-btn gc-accept" onclick="gcAccept(' + item.id + ')">✔ Accept</button>' +
                            '<button class="gc-btn gc-ignore" onclick="gcIgnore(' + item.id + ')">Ignore</button>';
                }
                html += '</div>';
            });
            document.getElementById('gc-result').innerHTML =
                html || '<p style="color:green"><b>No ' + type.toLowerCase() + ' errors 🎉</b></p>';
        }

        function gcAccept(id) {
            const item = gcApiData.results.find(r => r.id === id);
            item.current = item.corrected;
            document.getElementById('gc-row-' + id).innerHTML =
                '<p style="color:green"><b>✔ Accepted</b></p><p>' + gcEscape(item.current) + '</p>';
        }

        function gcIgnore(id) {
            document.getElementById('gc-row-' + id).remove();
        }

        function gcEscape(text) {
            const d = document.createElement('div');
            d.textContent = text;
            return d.innerHTML;
        }
        // ══════════════════════════════════════════
        //  Formatting & Alignment logic
        // ══════════════════════════════════════════
        let alTypes = [];

        function alToggleDropdown() {
            document.getElementById('al-spacingDropdown').classList.toggle('al-show');
        }

        // Close when clicking outside
        document.addEventListener('click', function(event) {
            if (!event.target.closest('.al-dropdown')) {
                const dd = document.getElementById('al-spacingDropdown');
                if (dd) dd.classList.remove('al-show');
            }
        });

        function alUpdateDropdownText() {
            const checkboxes = document.querySelectorAll('#al-spacingDropdown input[type=checkbox]');
            const selected = [];
            checkboxes.forEach(cb => { if (cb.checked) selected.push(cb.parentElement.innerText.trim()); });
            const btn = document.getElementById('al-spacingButton');
            btn.innerText = selected.length === 0 ? 'Select Spacing Options ▼' : selected.join(', ') + ' ▼';
            btn.style.color = 'black';
            document.getElementById('al-spacingDropdown').classList.remove('al-show');
        }

        async function alAnalyse() {
            const file = document.getElementById('al-file').files[0];
            if (!file) return alert('Select a .docx file');
            const fd = new FormData();
            fd.append('file', file);
            document.getElementById('al-ui').innerHTML = '<b>Analysing...</b>';
            const res = await fetch('/al_analyse', { method: 'POST', body: fd });
            const data = await res.json();
            alTypes = data.types;

            let html = '<h3>Formatting Options</h3>';
            alTypes.forEach(t => {
                html += `<div class="al-section-title">${t.toUpperCase()}</div>
                <label class="al-label">Alignment</label>
                <select class="al-select" id="al_${t}_alignment">
                    <option value="left">Align Left</option>
                    <option value="right">Align Right</option>
                    <option value="justify">Justify</option>
                </select>
                <label class="al-label">Line Spacing</label>
                <select class="al-select" id="al_${t}_spacing">
                    <option value="1">Single</option>
                    <option value="1.5">1.5</option>
                    <option value="2">Double</option>
                    <option value="2.5">2.5</option>
                    <option value="3">Triple</option>
                </select>`;
            });

            if (alTypes.includes('paragraph')) {
                html += `<div class="al-section-title" style="margin-top:20px;">Paragraph Spacing (Word Style)</div>
                <div class="al-dropdown">
                    <button type="button" id="al-spacingButton" onclick="alToggleDropdown()" class="al-dropbtn">
                        Select Spacing Options ▼
                    </button>
                    <div id="al-spacingDropdown" class="al-dropdown-content">
                        <label><input type="checkbox" value="add_before" onchange="alUpdateDropdownText()"> Add Space Before Paragraph</label>
                        <label><input type="checkbox" value="add_after" onchange="alUpdateDropdownText()"> Add Space After Paragraph</label>
                        <label><input type="checkbox" value="remove_before" onchange="alUpdateDropdownText()"> Remove Space Before Paragraph</label>
                        <label><input type="checkbox" value="remove_after" onchange="alUpdateDropdownText()"> Remove Space After Paragraph</label>
                    </div>
                </div>`;
            }

            html += '<button class="al-btn" onclick="alGenerate()">Generate Formatted DOCX</button>';
            document.getElementById('al-ui').innerHTML = html;
        }

        async function alGenerate() {
            const file = document.getElementById('al-file').files[0];
            const cfg = {};
            alTypes.forEach(t => {
                cfg[`${t}_alignment`] = document.getElementById(`al_${t}_alignment`).value;
                cfg[`${t}_spacing`]   = document.getElementById(`al_${t}_spacing`).value;
            });
            if (alTypes.includes('paragraph')) {
                const checkboxes = document.querySelectorAll('#al-spacingDropdown input[type=checkbox]');
                const selected = [];
                checkboxes.forEach(cb => { if (cb.checked) selected.push(cb.value); });
                cfg['paragraph_spacing_options'] = selected;
            }
            const fd = new FormData();
            fd.append('file', file);
            fd.append('config', JSON.stringify(cfg));
            const res = await fetch('/al_format', { method: 'POST', body: fd });
            const blob = await res.blob();
            const a = document.createElement('a');
            a.href = URL.createObjectURL(blob);
            a.download = 'formatted_alignment.docx';
            a.click();
        }

        // ══════════════════════════════════════════
        //  Smart Doc Formatter logic
        // ══════════════════════════════════════════
        let sdfAnalysisData = null;

        async function sdfAnalyzeDoc() {
            let file = document.getElementById("sdf-file").files[0];
            if (!file) { alert("Please upload a DOCX file"); return; }
            document.getElementById("sdf-analyzeBtn").disabled = true;
            document.getElementById("sdf-loadingMsg").style.display = "block";
            let fd = new FormData();
            fd.append("file", file);
            try {
                let res = await fetch("/analyze_structure", { method: "POST", body: fd });
                sdfAnalysisData = await res.json();
                if (sdfAnalysisData.error) throw new Error(sdfAnalysisData.error);
                sdfDisplayAnalysis(sdfAnalysisData);
                sdfBuildFormatControls(sdfAnalysisData.detected_elements);
                document.getElementById("sdf-loadingMsg").style.display = "none";
                document.getElementById("sdf-analysisResults").style.display = "block";
                document.getElementById("sdf-formatOptions").style.display = "block";
            } catch (e) {
                alert("Error analyzing document: " + e);
                document.getElementById("sdf-analyzeBtn").disabled = false;
                document.getElementById("sdf-loadingMsg").style.display = "none";
            }
        }

        function sdfRefreshAllPreviews() {
            if (!sdfAnalysisData) return;
            for (let elem of sdfAnalysisData.detected_elements) sdfUpdatePreview(elem);
        }

        function sdfFormatElementName(type) {
            return type.replace(/_/g, ' ').replace(/\b\w/g, l => l.toUpperCase());
        }

        function sdfDisplayAnalysis(data) {
            let list = document.getElementById("sdf-elementsList");
            list.innerHTML = "";
            let tocElements = {}, listElements = {}, coverElements = {}, regularElements = {};
            for (let [elemType, count] of Object.entries(data.element_counts)) {
                if (elemType.startsWith("TOC_")) tocElements[elemType] = count;
                else if (elemType === "BULLET_ITEM" || elemType === "NUMBERED_ITEM" || elemType === "LIST_ITEM") listElements[elemType] = count;
                else if (elemType === "TITLE" || elemType === "COVER_PAGE") coverElements[elemType] = count;
                else regularElements[elemType] = count;
            }
            if (Object.keys(coverElements).length > 0) {
                let li = document.createElement("li");
                li.style.color = "#b45309"; li.style.fontWeight = "500";
                li.innerHTML = `<strong>Cover / Title Page Detected:</strong>`;
                list.appendChild(li);
                for (let [elemType, count] of Object.entries(coverElements)) {
                    let subLi = document.createElement("li");
                    subLi.style.paddingLeft = "20px"; subLi.style.color = "#b45309";
                    subLi.innerHTML = `→ ${sdfFormatElementName(elemType)}: ${count} found`;
                    list.appendChild(subLi);
                }
            }
            for (let [elemType, count] of Object.entries(regularElements)) {
                let li = document.createElement("li");
                li.innerHTML = `<strong>${sdfFormatElementName(elemType)}:</strong> ${count} found`;
                list.appendChild(li);
            }
            if (Object.keys(listElements).length > 0) {
                let li = document.createElement("li");
                li.className = "list-item";
                li.innerHTML = `<strong>Lists Detected:</strong>`;
                list.appendChild(li);
                for (let [elemType, count] of Object.entries(listElements)) {
                    let subLi = document.createElement("li");
                    subLi.className = "list-item"; subLi.style.paddingLeft = "20px";
                    subLi.innerHTML = `→ ${sdfFormatElementName(elemType)}: ${count} found`;
                    list.appendChild(subLi);
                }
            }
            if (Object.keys(tocElements).length > 0) {
                let li = document.createElement("li");
                li.className = "toc-item";
                li.innerHTML = `<strong>Table of Contents Detected:</strong>`;
                list.appendChild(li);
                for (let [elemType, count] of Object.entries(tocElements)) {
                    let subLi = document.createElement("li");
                    subLi.className = "toc-item"; subLi.style.paddingLeft = "20px";
                    subLi.innerHTML = `→ ${sdfFormatElementName(elemType)}: ${count} found`;
                    list.appendChild(subLi);
                }
            }
            // Bullets by heading
            let box = document.getElementById("sdf-bulletsByHeadingBox");
            let ul = document.getElementById("sdf-bulletsByHeadingList");
            ul.innerHTML = "";
            if (data.list_items_by_heading && Object.keys(data.list_items_by_heading).length > 0) {
                box.style.display = "block";
                const keys = Object.keys(data.list_items_by_heading).sort((a, b) => {
                    if (a === "NO_HEADING") return 1;
                    if (b === "NO_HEADING") return -1;
                    return (parseInt(a.split("_").pop(), 10) || 99) - (parseInt(b.split("_").pop(), 10) || 99);
                });
                for (const k of keys) {
                    const c = data.list_items_by_heading[k];
                    const label = k === "NO_HEADING" ? "No Heading" : sdfFormatElementName(k);
                    const li = document.createElement("li");
                    li.innerHTML = `<strong>${label}:</strong> ${c} list items`;
                    ul.appendChild(li);
                }
            } else {
                box.style.display = "none";
            }
            if (data.has_toc || data.has_indented_hierarchy || data.cover_page_end >= 0) {
                let hierarchyDiv = document.getElementById("sdf-hierarchyInfo");
                let messages = [];
                if (data.cover_page_end >= 0) messages.push(`Cover page detected (${data.cover_page_end + 1} paragraph(s))`);
                if (data.has_toc) messages.push(`Table of Contents with ${data.toc_indent_levels} hierarchy levels`);
                if (data.has_indented_hierarchy) messages.push(`${data.indent_levels} indentation levels in document body`);
                document.getElementById("sdf-hierarchyText").textContent = messages.join(" • ");
                hierarchyDiv.style.display = "block";
            }
        }

        function sdfEscapeHtml(text) {
            return (text || "").replace(/&/g,"&amp;").replace(/</g,"&lt;").replace(/>/g,"&gt;").replace(/\t/g,"&nbsp;&nbsp;&nbsp;&nbsp;");
        }

        function sdfOrderKey(elem) {
            if (elem === "TITLE") return 0;
            if (elem === "COVER_PAGE") return 1;
            if (elem === "TOC_TITLE") return 10;
            if (elem.startsWith("TOC_HEADING_")) return 11 + (parseInt(elem.split("_").pop(), 10) || 99);
            if (elem.startsWith("HEADING_")) return 100 + (parseInt(elem.split("_").pop(), 10) || 99);
            if (elem === "BULLET_ITEM") return 190;
            if (elem === "NUMBERED_ITEM") return 195;
            if (elem === "LIST_ITEM") return 200;
            if (elem === "PARAGRAPH") return 210;
            if (elem === "TABLE") return 220;
            return 999;
        }

        function sdfSortElements(elements) {
            return [...elements].sort((a, b) => {
                const ka = sdfOrderKey(a), kb = sdfOrderKey(b);
                return ka !== kb ? ka - kb : a.localeCompare(b);
            });
        }

        function sdfBuildFormatControls(elements) {
            let container = document.getElementById("sdf-styleControls");
            container.innerHTML = "";
            const sorted = sdfSortElements(elements);
            const defaults = {
                "TITLE":        { font: "Times New Roman", size: 26 },
                "COVER_PAGE":   { font: "Times New Roman", size: 12 },
                "TOC_TITLE":    { font: "Calibri", size: 14 },
                "TOC_HEADING_1":{ font: "Calibri", size: 11 },
                "TOC_HEADING_2":{ font: "Calibri", size: 11 },
                "TOC_HEADING_3":{ font: "Calibri", size: 11 },
                "TOC_HEADING_4":{ font: "Calibri", size: 11 },
                "HEADING_1":    { font: "Calibri", size: 18 },
                "HEADING_2":    { font: "Calibri", size: 16 },
                "HEADING_3":    { font: "Calibri", size: 14 },
                "HEADING_4":    { font: "Calibri", size: 13 },
                "HEADING_5":    { font: "Calibri", size: 12 },
                "HEADING_6":    { font: "Calibri", size: 11 },
                "PARAGRAPH":    { font: "Calibri", size: 12 },
                "BULLET_ITEM":  { font: "Calibri", size: 12 },
                "NUMBERED_ITEM":{ font: "Calibri", size: 12 },
                "LIST_ITEM":    { font: "Calibri", size: 12 },
                "TABLE":        { font: "Calibri", size: 11 }
            };
            const fontOptions = [
                "Times New Roman","Arial","Calibri","Georgia","Verdana","Tahoma",
                "Trebuchet MS","Courier New","Garamond","Palatino Linotype",
                "Book Antiqua","Comic Sans MS","Impact","Lucida Sans Unicode","Century Gothic"
            ];
            const tocElements     = sorted.filter(e => e.startsWith("TOC_"));
            const listElements    = sorted.filter(e => e === "BULLET_ITEM" || e === "NUMBERED_ITEM" || e === "LIST_ITEM");
            const coverElements   = sorted.filter(e => e === "TITLE" || e === "COVER_PAGE");
            const regularElements = sorted.filter(e => !e.startsWith("TOC_") && e !== "BULLET_ITEM" && e !== "NUMBERED_ITEM" && e !== "LIST_ITEM" && e !== "TITLE" && e !== "COVER_PAGE");

            if (coverElements.length > 0) {
                let sec = document.createElement("div");
                sec.className = "sdf-cover-section";
                sec.innerHTML = "<h3>Cover / Title Page</h3>";
                container.appendChild(sec);
                for (let elem of coverElements) sdfCreateControl(elem, defaults, fontOptions, sec);
            }
            for (let elem of regularElements) sdfCreateControl(elem, defaults, fontOptions, container);
            if (listElements.length > 0) {
                let sec = document.createElement("div");
                sec.className = "sdf-list-section";
                sec.innerHTML = "<h3>List Item Formatting</h3>";
                container.appendChild(sec);
                for (let elem of listElements) sdfCreateControl(elem, defaults, fontOptions, sec);
            }
            if (tocElements.length > 0) {
                let sec = document.createElement("div");
                sec.className = "sdf-toc-section";
                sec.innerHTML = "<h3>Table of Contents Formatting</h3>";
                container.appendChild(sec);
                for (let elem of tocElements) sdfCreateControl(elem, defaults, fontOptions, sec);
            }
            for (let elem of sorted) sdfUpdatePreview(elem);
        }

        function sdfCreateControl(elem, defaults, fontOptions, container) {
            let def = defaults[elem] || { font: "Calibri", size: 12 };
            let section = document.createElement("div");
            section.style.marginBottom = "15px";
            let fontOptionsHTML = fontOptions.map(f => `<option ${def.font === f ? "selected" : ""}>${f}</option>`).join('');
            section.innerHTML = `
                <h3>${sdfFormatElementName(elem)} Style</h3>
                <div class="sdf-preview-container">
                  <div style="flex:1">
                    <label>Font</label>
                    <select id="sdf_${elem}_font" onchange="sdfUpdatePreview('${elem}')">${fontOptionsHTML}</select>
                    <label>Size</label>
                    <input type="number" id="sdf_${elem}_size" value="${def.size}" onchange="sdfUpdatePreview('${elem}')">
                  </div>
                  <div class="sdf-preview-box">
                    <div class="sdf-preview-label">Preview</div>
                    <div id="sdf_${elem}_preview">${sdfEscapeHtml((sdfAnalysisData && sdfAnalysisData.sample_texts && sdfAnalysisData.sample_texts[elem]) || "Sample text")}</div>
                  </div>
                </div>
            `;
            container.appendChild(section);
        }

        function sdfUpdatePreview(elem) {
            let preview = document.getElementById("sdf_" + elem + "_preview");
            if (!preview) return;
            let fontEl = document.getElementById("sdf_" + elem + "_font");
            let sizeEl = document.getElementById("sdf_" + elem + "_size");
            if (!fontEl || !sizeEl) return;
            preview.style.fontFamily = fontEl.value;
            preview.style.fontSize = sizeEl.value + "px";
            let boldLists = document.getElementById("sdf-bold_lists") ? document.getElementById("sdf-bold_lists").checked : false;
            if (
                ((elem === "TITLE" || elem === "COVER_PAGE") && document.getElementById("sdf-bold_titles").checked) ||
                (elem.startsWith("HEADING_") && document.getElementById("sdf-bold_headings").checked) ||
                (elem === "TOC_TITLE" && document.getElementById("sdf-bold_toc").checked) ||
                ((elem === "BULLET_ITEM" || elem === "NUMBERED_ITEM" || elem === "LIST_ITEM") && boldLists)
            ) preview.style.fontWeight = "bold";
            else preview.style.fontWeight = "normal";
            if (document.getElementById("sdf-highlight").checked) preview.style.background = "#fff59d";
            else preview.style.background = "transparent";
        }

        async function sdfFormatDoc() {
            let file = document.getElementById("sdf-file").files[0];
            if (!file || !sdfAnalysisData) { alert("Please analyze document first"); return; }
            let config = {
                bold_titles:          document.getElementById("sdf-bold_titles").checked,
                bold_headings:        document.getElementById("sdf-bold_headings").checked,
                bold_toc:             document.getElementById("sdf-bold_toc").checked,
                bold_lists:           document.getElementById("sdf-bold_lists") ? document.getElementById("sdf-bold_lists").checked : false,
                highlight:            document.getElementById("sdf-highlight").checked,
                preserve_indentation: document.getElementById("sdf-preserve_indentation").checked
            };
            for (let elem of sdfAnalysisData.detected_elements) {
                let fontEl = document.getElementById("sdf_" + elem + "_font");
                let sizeEl = document.getElementById("sdf_" + elem + "_size");
                if (fontEl) config[elem.toLowerCase() + "_font"] = fontEl.value;
                if (sizeEl) config[elem.toLowerCase() + "_size"] = sizeEl.value;
            }
            let fd = new FormData();
            fd.append("file", file);
            fd.append("config", JSON.stringify(config));
            let res = await fetch("/format", { method: "POST", body: fd });
            if (!res.ok) { alert("Format failed: " + await res.text()); return; }
            let blob = await res.blob();
            let a = document.createElement("a");
            a.href = URL.createObjectURL(blob);
            a.download = file.name.replace(".docx", "_formatted.docx");
            a.click();
        }
    </script>
</head>
<body>
<div class="container">
    <h1>📄 Document Formatter</h1>
    <div class="tabs">
        <button id="font-btn" class="tab-button" onclick="showTab('font')">Font Style &amp; Size</button>
        <button id="grammar-btn" class="tab-button" onclick="showTab('grammar')">Grammar &amp; Spell Check</button>
        <button id="formatting-btn" class="tab-button" onclick="showTab('formatting')">Formatting &amp; Alignment</button>
    </div>

    <!-- ── TAB 1: Font Style & Size (Smart Doc Formatter) ── -->
    <div id="font" class="tab-content">
        <div class="sdf-box">
          <div class="sdf-inner">
            <h2>Smart Document Formatter</h2>

            <input type="file" id="sdf-file" accept=".docx">
            <br>
            <button class="sdf-btn" type="button" onclick="sdfAnalyzeDoc()" id="sdf-analyzeBtn">Analyze Document</button>
            <div id="sdf-loadingMsg" class="sdf-loading" style="display:none;">Analyzing document structure...</div>

            <div id="sdf-analysisResults" style="display:none;">
              <h2>Document Analysis Complete</h2>
              <div class="sdf-info-box">
                <h3>Detected Elements:</h3>
                <ul id="sdf-elementsList" class="sdf-element-list"></ul>
                <div id="sdf-bulletsByHeadingBox" class="sdf-subbox" style="display:none;">
                  <h4>List items (bullets + numbered) by Heading</h4>
                  <div class="sdf-small">List items are categorized under their nearest preceding heading.</div>
                  <ul id="sdf-bulletsByHeadingList" class="sdf-element-list" style="margin-top:8px;"></ul>
                </div>
              </div>
              <div id="sdf-hierarchyInfo" class="sdf-hierarchy-info" style="display:none;">
                <strong>Document Structure:</strong> <span id="sdf-hierarchyText"></span>
              </div>
            </div>

            <div id="sdf-formatOptions" style="display:none;">
              <h2>Formatting Options</h2>
              <p>Configure styling for the detected elements:</p>
              <div id="sdf-styleControls"></div>

              <label class="sdf-checkbox-label"><input type="checkbox" id="sdf-bold_titles" checked onchange="sdfRefreshAllPreviews()"> Bold Titles</label>
              <label class="sdf-checkbox-label"><input type="checkbox" id="sdf-bold_headings" checked onchange="sdfRefreshAllPreviews()"> Bold Headings</label>
              <label class="sdf-checkbox-label"><input type="checkbox" id="sdf-bold_toc" checked onchange="sdfRefreshAllPreviews()"> Bold TOC Title</label>
              <label class="sdf-checkbox-label"><input type="checkbox" id="sdf-bold_lists" onchange="sdfRefreshAllPreviews()"> Bold List Items</label>
              <label class="sdf-checkbox-label"><input type="checkbox" id="sdf-highlight" checked onchange="sdfRefreshAllPreviews()"> Highlight Sections</label>
              <label class="sdf-checkbox-label"><input type="checkbox" id="sdf-preserve_indentation" checked> Preserve Indentation</label>
              <br>
              <button class="sdf-btn" type="button" onclick="sdfFormatDoc()">Format Document</button>
            </div>
          </div>
        </div>
    </div>

    <!-- ── TAB 2: Grammar & Spell Check ── -->
    <div id="grammar" class="tab-content" style="display:none;">
        <h2>📝 Grammar &amp; Spell Checker</h2>
        <input type="file" id="gc-file" accept=".docx"><br><br>
        <button class="gc-btn" onclick="gcUpload()">Analyze Document</button>
        <div style="margin:20px 0;">
            <button class="gc-btn gc-btn-outline" onclick="gcShowSpelling()">Spelling Errors</button>
            <button class="gc-btn gc-btn-outline" onclick="gcShowGrammar()">Grammar Errors</button>
        </div>
        <div id="gc-result"></div>
    </div>

    <!-- ── TAB 3: Formatting & Alignment ── -->
    <div id="formatting" class="tab-content" style="display:none;">
        <h2>📄 Formatting &amp; Alignment</h2>
        <input type="file" id="al-file"><br><br>
        <button class="al-btn" onclick="alAnalyse()">Analyse Document</button>
        <div id="al-ui"></div>
    </div>
</div>
</body>
</html>
"""

# ================== UTILITIES (from smart_doc_formatter.py) ==================

def get_paragraph_indentation(para) -> float:
    try:
        if para.paragraph_format.left_indent:
            return para.paragraph_format.left_indent.inches
        return 0.0
    except Exception:
        return 0.0


def get_list_level(para) -> int:
    try:
        ilvl_nodes = para._element.xpath(".//w:numPr/w:ilvl")
        if ilvl_nodes:
            return int(ilvl_nodes[0].get(qn("w:val"), 0))
    except Exception:
        pass
    return 0


def is_bullet_list(para) -> bool:
    try:
        style_name = para.style.name if para.style else ""
        if "List Bullet" in style_name:
            return True
        numPr_list = para._element.xpath(".//w:numPr")
        if not numPr_list:
            return False
        numId_nodes = para._element.xpath(".//w:numPr/w:numId")
        if not numId_nodes:
            return False
        doc = para._p.getparent()
        while doc is not None and doc.tag != qn("w:body"):
            doc = doc.getparent()
        if doc is not None:
            try:
                numbering_part = para.part.numbering_part
                if numbering_part is None:
                    return True
                numId_val = numId_nodes[0].get(qn("w:val"), "0")
                ilvl_nodes = para._element.xpath(".//w:numPr/w:ilvl")
                ilvl_val = ilvl_nodes[0].get(qn("w:val"), "0") if ilvl_nodes else "0"
                num_elements = numbering_part._element.xpath(
                    f".//w:num[@w:numId='{numId_val}']",
                    namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                )
                if not num_elements:
                    return False
                abstractNumId_nodes = num_elements[0].xpath(
                    ".//w:abstractNumId",
                    namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                )
                if not abstractNumId_nodes:
                    return False
                abstract_id = abstractNumId_nodes[0].get(qn("w:val"), "0")
                abstract_num = numbering_part._element.xpath(
                    f".//w:abstractNum[@w:abstractNumId='{abstract_id}']",
                    namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                )
                if not abstract_num:
                    return False
                lvl_elements = abstract_num[0].xpath(
                    f".//w:lvl[@w:ilvl='{ilvl_val}']/w:numFmt",
                    namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                )
                if lvl_elements:
                    fmt = lvl_elements[0].get(qn("w:val"), "")
                    if fmt in ("bullet", "none"):
                        return True
                    if fmt in ("decimal","lowerLetter","upperLetter","lowerRoman","upperRoman",
                               "ordinal","cardinalText","decimalZero"):
                        return False
                    return True
            except Exception:
                return False
    except Exception:
        pass
    return False


def is_numbered_list(para) -> bool:
    try:
        style_name = para.style.name if para.style else ""
        if "List Number" in style_name:
            return True
        numPr_list = para._element.xpath(".//w:numPr")
        if not numPr_list:
            return False
        numId_nodes = para._element.xpath(".//w:numPr/w:numId")
        if not numId_nodes:
            return False
        try:
            numbering_part = para.part.numbering_part
            if numbering_part is None:
                return False
            numId_val = numId_nodes[0].get(qn("w:val"), "0")
            ilvl_nodes = para._element.xpath(".//w:numPr/w:ilvl")
            ilvl_val = ilvl_nodes[0].get(qn("w:val"), "0") if ilvl_nodes else "0"
            num_elements = numbering_part._element.xpath(
                f".//w:num[@w:numId='{numId_val}']",
                namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            )
            if not num_elements:
                return False
            abstractNumId_nodes = num_elements[0].xpath(
                ".//w:abstractNumId",
                namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            )
            if not abstractNumId_nodes:
                return False
            abstract_id = abstractNumId_nodes[0].get(qn("w:val"), "0")
            abstract_num = numbering_part._element.xpath(
                f".//w:abstractNum[@w:abstractNumId='{abstract_id}']",
                namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            )
            if not abstract_num:
                return False
            lvl_elements = abstract_num[0].xpath(
                f".//w:lvl[@w:ilvl='{ilvl_val}']/w:numFmt",
                namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            )
            if lvl_elements:
                fmt = lvl_elements[0].get(qn("w:val"), "")
                if fmt in ("decimal","lowerLetter","upperLetter","lowerRoman","upperRoman",
                           "ordinal","cardinalText","decimalZero"):
                    return True
        except Exception:
            pass
    except Exception:
        pass
    return False


def is_any_list_paragraph(para) -> bool:
    try:
        if para._element.xpath(".//w:numPr"):
            return True
    except Exception:
        pass
    try:
        style_name = para.style.name if para.style is not None else ""
        if "List" in style_name:
            return True
    except Exception:
        pass
    return False


def get_list_type(para) -> str:
    if is_bullet_list(para):
        return "BULLET_ITEM"
    if is_numbered_list(para):
        return "NUMBERED_ITEM"
    if is_any_list_paragraph(para):
        return "LIST_ITEM"
    return None


NUMBERED_HEADING_RE = re.compile(r"^\s*(\d+(?:\.\d+)+)\s+(.*\S.*)$")
MANUAL_NUMBERED_LIST_RE = re.compile(r"^\s*(\d+[\.\)]\s+|[a-zA-Z][\.\)]\s+|[ivxIVX]+[\.\)]\s+)(.*\S.*)$")


def detect_numbered_heading(text: str):
    if not text:
        return None, None, None
    m = NUMBERED_HEADING_RE.match(text)
    if not m:
        return None, None, None
    num = m.group(1)
    title = m.group(2).strip()
    if len(num.split(".")) < 2:
        return None, None, None
    level = min(6, len(num.split(".")))
    return level, num, title


def clear_paragraph_runs(para):
    p = para._p
    for r in list(para.runs):
        try:
            p.remove(r._r)
        except Exception:
            pass


def has_page_break_before(para) -> bool:
    try:
        pPr = para._element.find(qn("w:pPr"))
        if pPr is not None:
            pb = pPr.find(qn("w:pageBreakBefore"))
            if pb is not None:
                val = pb.get(qn("w:val"), "true")
                if val.lower() not in ("false", "0", "off"):
                    return True
        for br in para._element.xpath(".//w:br"):
            if br.get(qn("w:type"), "") == "page":
                return True
        if "lastRenderedPageBreak" in para._element.xml:
            return True
    except Exception:
        pass
    return False


def detect_cover_page(doc) -> int:
    paragraphs = doc.paragraphs
    n = len(paragraphs)
    for i, para in enumerate(paragraphs[:60]):
        for run in para.runs:
            for br in run._element.xpath(".//w:br"):
                if br.get(qn("w:type"), "") == "page":
                    return i
        if i > 0 and has_page_break_before(para):
            return i - 1
    for i, para in enumerate(paragraphs[:40]):
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading") and i > 0:
            return i - 1
    COVER_KEYWORDS = re.compile(
        r"\b(prepared\s+(by|for)|author|version|copyright|confidential|"
        r"restricted|january|february|march|april|may|june|july|august|"
        r"september|october|november|december|\d{4})\b",
        re.IGNORECASE,
    )
    non_empty = [(i, p) for i, p in enumerate(paragraphs) if (p.text or "").strip()]
    cover_candidate = -1
    for rank, (i, para) in enumerate(non_empty[:25]):
        text = (para.text or "").strip()
        is_short = len(text) < 200
        is_centered = para.alignment == WD_ALIGN_PARAGRAPH.CENTER
        has_large_font = False
        style_name = para.style.name if para.style else ""
        if para.runs:
            sz = para.runs[0].font.size
            if sz and sz >= Pt(13):
                has_large_font = True
        has_keyword = bool(COVER_KEYWORDS.search(text))
        is_title_style = style_name in ("Title", "Subtitle")
        if is_title_style or is_centered or has_large_font or has_keyword or is_short:
            cover_candidate = i
        else:
            if rank > 2:
                break
    return cover_candidate


def _get_style_base(para) -> str:
    try:
        name = para.style.name if para.style else ""
        return name.split(" Char")[0].strip()
    except Exception:
        return ""


def _is_toc_entry_by_style(para) -> bool:
    return bool(re.match(r"^TOC \d+$", _get_style_base(para)))


def _toc_level_from_style(para) -> int:
    m = re.match(r"^TOC (\d+)$", _get_style_base(para))
    if m:
        return int(m.group(1))
    return 0


_TOC_ENTRY_RE = re.compile(
    r"^\s*(\d+(?:\.\d+)*)?\s*(.+?)[\s\.…\-]{2,}(\d{1,4})\s*$",
    re.UNICODE,
)
_TOC_TAB_RE = re.compile(r"^.+\t\d{1,4}\s*$")


def _is_toc_entry_by_heuristic(para) -> bool:
    text = (para.text or "").strip()
    if not text:
        return False
    if _TOC_ENTRY_RE.match(text):
        return True
    if _TOC_TAB_RE.match(text):
        return True
    xml = para._element.xml
    if "PAGEREF" in xml or ("w:instr" in xml and "TOC" in xml):
        return True
    return False


def _toc_level_from_heuristic(text: str) -> int:
    m = re.match(r"^\s*(\d+(?:\.\d+)*)\s+", text.strip())
    if m:
        return min(len(m.group(1).split(".")), 6)
    return 1


def detect_toc_section(doc):
    paragraphs = doc.paragraphs
    first_toc_style_idx = -1
    last_toc_style_idx = -1
    for i, para in enumerate(paragraphs):
        if _is_toc_entry_by_style(para):
            if first_toc_style_idx == -1:
                first_toc_style_idx = i
            last_toc_style_idx = i
        elif first_toc_style_idx != -1:
            if i - last_toc_style_idx > 5:
                break
    if first_toc_style_idx != -1:
        toc_start = first_toc_style_idx
        for j in range(first_toc_style_idx - 1, max(-1, first_toc_style_idx - 5), -1):
            text = (paragraphs[j].text or "").strip().lower()
            style_b = _get_style_base(paragraphs[j])
            if style_b == "TOC Heading" or "table of contents" in text or text == "contents":
                toc_start = j
                break
            if text:
                break
        return toc_start, last_toc_style_idx, True
    TOC_TITLE_RE = re.compile(r"^\s*(table\s+of\s+contents|contents)\s*$", re.IGNORECASE)
    toc_start = -1
    search_limit = min(len(paragraphs), 100)
    for i, para in enumerate(paragraphs[:search_limit]):
        text = (para.text or "").strip()
        style_b = _get_style_base(para)
        if style_b == "TOC Heading" or TOC_TITLE_RE.match(text):
            toc_start = i
            break
        if "TOC" in para._element.xml and "w:fldChar" in para._element.xml:
            toc_start = i
            break
    if toc_start == -1:
        return -1, -1, False
    toc_end = toc_start
    consecutive_non_toc = 0
    for i in range(toc_start + 1, min(len(paragraphs), toc_start + 300)):
        para = paragraphs[i]
        text = (para.text or "").strip()
        if not text:
            continue
        if _is_toc_entry_by_heuristic(para) or _is_toc_entry_by_style(para):
            toc_end = i
            consecutive_non_toc = 0
        else:
            consecutive_non_toc += 1
            if consecutive_non_toc >= 3:
                break
    return toc_start, toc_end, True


def analyze_toc_hierarchy(doc, toc_start, toc_end):
    if toc_start == -1 or toc_end == -1:
        return {}, 0
    style_levels = set()
    for i in range(toc_start, toc_end + 1):
        if i >= len(doc.paragraphs):
            break
        lvl = _toc_level_from_style(doc.paragraphs[i])
        if lvl:
            style_levels.add(lvl)
    if style_levels:
        return {lvl: lvl for lvl in style_levels}, len(style_levels)
    level_counts = {}
    for i in range(toc_start + 1, toc_end + 1):
        if i >= len(doc.paragraphs):
            break
        para = doc.paragraphs[i]
        text = (para.text or "").strip()
        if not text:
            continue
        if _is_toc_entry_by_heuristic(para):
            lvl = _toc_level_from_heuristic(text)
            level_counts[lvl] = level_counts.get(lvl, 0) + 1
    return {lvl: lvl for lvl in sorted(level_counts)}, len(level_counts)


def detect_hierarchy_by_indentation(doc, exclude_range=None):
    indent_levels = {}
    for i, para in enumerate(doc.paragraphs):
        if exclude_range and exclude_range[0] <= i <= exclude_range[1]:
            continue
        if not (para.text or "").strip():
            continue
        if is_any_list_paragraph(para):
            continue
        indent = get_paragraph_indentation(para)
        if indent > 3:
            continue
        if indent > 0:
            rounded_indent = round(indent * 4) / 4
            indent_levels[rounded_indent] = indent_levels.get(rounded_indent, 0) + 1
    sorted_indents = sorted(indent_levels.keys())
    indent_to_level = {}
    for i, indent in enumerate(sorted_indents[:6]):
        indent_to_level[indent] = i + 1
    return indent_to_level, len(sorted_indents)


def analyze_document_structure(docx_path):
    doc = Document(docx_path)
    elements = []
    element_counts = {}
    detected_types = set()
    sample_texts = {}
    list_items_by_heading = {}

    cover_end = detect_cover_page(doc)
    toc_start, toc_end, has_toc = detect_toc_section(doc)
    toc_indent_to_level = {}
    toc_indent_levels = 0
    if has_toc:
        toc_indent_to_level, toc_indent_levels = analyze_toc_hierarchy(doc, toc_start, toc_end)

    indent_to_level, total_indent_levels = detect_hierarchy_by_indentation(
        doc, exclude_range=(toc_start, toc_end) if has_toc else None
    )
    has_indented_hierarchy = len(indent_to_level) > 0

    heading_stack = {}

    def current_parent_heading():
        if not heading_stack:
            return None
        return heading_stack[max(heading_stack.keys())]

    def set_heading(level: int):
        heading_stack[level] = f"HEADING_{level}"
        for k in list(heading_stack.keys()):
            if k > level:
                del heading_stack[k]

    first_content_idx = next(
        (i for i, p in enumerate(doc.paragraphs) if (p.text or "").strip()), 0
    )

    for idx, para in enumerate(doc.paragraphs):
        if not (para.text or "").strip():
            continue

        style_name = para.style.name if para.style is not None else ""
        style_base = style_name.split(" Char")[0].strip()

        ptype = None
        parent_heading = None
        in_toc_section = has_toc and toc_start <= idx <= toc_end
        in_cover_page = cover_end >= 0 and idx <= cover_end

        if style_base == "Title":
            ptype = "TITLE"
        elif style_base == "TOC Heading":
            ptype = "TOC_TITLE"
        elif style_base in ("Subtitle", "Document Map"):
            ptype = "COVER_PAGE"
        elif style_base in ("Author", "Date", "Company", "Abstract",
                             "Document Label", "Revision", "Version"):
            ptype = "COVER_PAGE"
        elif re.match(r"^TOC \d+$", style_base):
            try:
                level = int(style_base.split()[-1])
            except Exception:
                level = 1
            ptype = f"TOC_HEADING_{level}"

        if ptype is None and in_toc_section:
            text_lower = para.text.strip().lower()
            style_b = _get_style_base(para)
            if style_b == "TOC Heading" or re.match(r"^\s*(table\s+of\s+contents|contents)\s*$", text_lower):
                ptype = "TOC_TITLE"
            else:
                lvl = _toc_level_from_style(para)
                if lvl:
                    ptype = f"TOC_HEADING_{lvl}"
                else:
                    text_raw = (para.text or "").strip()
                    lvl = _toc_level_from_heuristic(text_raw)
                    ptype = f"TOC_HEADING_{lvl}"

        if ptype is None and in_cover_page and not in_toc_section:
            if idx == first_content_idx:
                ptype = "TITLE"
            else:
                ptype = "COVER_PAGE"

        if ptype is None:
            raw_text = (para.text or "").strip()
            if style_base.startswith("Heading"):
                try:
                    level = int(style_base.split()[-1])
                    level = max(1, min(6, level))
                    ptype = f"HEADING_{level}"
                except Exception:
                    ptype = "HEADING_1"
            if not ptype:
                list_type = get_list_type(para)
                if list_type:
                    ptype = list_type
            if not ptype:
                lvl, num, title = detect_numbered_heading(raw_text)
                if lvl:
                    ptype = f"HEADING_{lvl}"
            if not ptype:
                indent = get_paragraph_indentation(para)
                if indent > 0 and indent_to_level:
                    rounded_indent = round(indent * 4) / 4
                    if rounded_indent in indent_to_level:
                        ptype = f"HEADING_{indent_to_level[rounded_indent]}"
            if not ptype:
                ptype = "PARAGRAPH"

        if ptype and ptype.startswith("HEADING_"):
            try:
                lvl = int(ptype.split("_")[-1])
            except Exception:
                lvl = 1
            set_heading(lvl)

        if ptype in ("BULLET_ITEM", "NUMBERED_ITEM", "LIST_ITEM"):
            parent_heading = current_parent_heading()
            key = parent_heading if parent_heading else "NO_HEADING"
            list_items_by_heading[key] = list_items_by_heading.get(key, 0) + 1

        item = {
            "type": ptype,
            "para_idx": idx,
            "indent": get_paragraph_indentation(para),
            "list_level": get_list_level(para) if ptype in ("BULLET_ITEM", "NUMBERED_ITEM", "LIST_ITEM") else 0
        }
        if parent_heading:
            item["parent_heading"] = parent_heading

        if not in_toc_section and not in_cover_page and ptype and ptype.startswith("HEADING_"):
            lvl2, num2, title2 = detect_numbered_heading((para.text or "").strip())
            if num2 and title2:
                item["heading_number"] = num2
                item["heading_title"] = title2

        elements.append(item)
        element_counts[ptype] = element_counts.get(ptype, 0) + 1
        detected_types.add(ptype)

        if ptype not in sample_texts:
            t = (para.text or "").strip()
            if t:
                sample_texts[ptype] = t[:250]

    for i, table in enumerate(doc.tables):
        elements.append({"type": "TABLE", "table_idx": i})
        element_counts["TABLE"] = element_counts.get("TABLE", 0) + 1
        detected_types.add("TABLE")
        if "TABLE" not in sample_texts:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        sample_texts["TABLE"] = cell.text.strip()[:250]
                        break

    def sort_key(x):
        if x == "TITLE": return 0
        if x == "COVER_PAGE": return 1
        if x == "TOC_TITLE": return 2
        if x.startswith("TOC_HEADING_"): return 3
        if x.startswith("HEADING_"): return 4
        if x in ("BULLET_ITEM", "NUMBERED_ITEM", "LIST_ITEM"): return 5
        if x == "PARAGRAPH": return 6
        return 7

    detected_elements = sorted(detected_types, key=sort_key)

    return {
        "elements": elements,
        "detected_elements": detected_elements,
        "element_counts": element_counts,
        "sample_texts": sample_texts,
        "has_toc": has_toc,
        "toc_indent_levels": toc_indent_levels,
        "has_indented_hierarchy": has_indented_hierarchy,
        "indent_levels": total_indent_levels,
        "list_items_by_heading": list_items_by_heading,
        "cover_page_end": cover_end,
    }


# ================== XML HELPERS ==================

def get_or_add_tblPr(ct_tbl):
    tblPr = ct_tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        ct_tbl.insert(0, tblPr)
    return tblPr


def set_run_font_everywhere(run, font_name, font_size_pt=None, bold=None):
    run.font.name = font_name
    if font_size_pt is not None:
        run.font.size = Pt(font_size_pt)
    if bold is not None:
        run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    for attr in (qn("w:asciiTheme"), qn("w:hAnsiTheme"), qn("w:cstheme")):
        if rFonts.get(attr) is not None:
            del rFonts.attrib[attr]
    if font_size_pt is not None:
        half_pts = str(int(font_size_pt * 2))
        for tag_name in ("w:sz", "w:szCs"):
            el = rPr.find(qn(tag_name))
            if el is None:
                el = OxmlElement(tag_name)
                rPr.append(el)
            el.set(qn("w:val"), half_pts)
    if bold is not None:
        for tag_name in ("w:b", "w:bCs"):
            el = rPr.find(qn(tag_name))
            if el is None:
                el = OxmlElement(tag_name)
                rPr.append(el)
            el.set(qn("w:val"), "1" if bold else "0")


def apply_para_direct_format(para, font_name, font_size_pt, bold, highlight_color_name=None):
    pPr = para._element.get_or_add_pPr()
    rPr = pPr.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        pPr.append(rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    for attr in (qn("w:asciiTheme"), qn("w:hAnsiTheme"), qn("w:cstheme")):
        if rFonts.get(attr) is not None:
            del rFonts.attrib[attr]
    half_pts = str(int(font_size_pt * 2))
    for tag_name in ("w:sz", "w:szCs"):
        el = rPr.find(qn(tag_name))
        if el is None:
            el = OxmlElement(tag_name)
            rPr.append(el)
        el.set(qn("w:val"), half_pts)
    for tag_name in ("w:b", "w:bCs"):
        el = rPr.find(qn(tag_name))
        if el is None:
            el = OxmlElement(tag_name)
            rPr.append(el)
        el.set(qn("w:val"), "1" if bold else "0")
    if highlight_color_name:
        hl = rPr.find(qn("w:highlight"))
        if hl is None:
            hl = OxmlElement("w:highlight")
            rPr.append(hl)
        hl.set(qn("w:val"), highlight_color_name)


def _get_all_runs(para):
    all_r = para._element.xpath(".//w:r")
    nested_tbl_runs = set()
    for tbl in para._element.xpath(".//w:tbl"):
        for r in tbl.xpath(".//w:r"):
            nested_tbl_runs.add(id(r))
    return [r for r in all_r if id(r) not in nested_tbl_runs]


def _patch_doc_style(doc, style_name, font_name, font_size_pt, bold):
    try:
        style = doc.styles[style_name]
    except (KeyError, Exception):
        return
    style_el = style.element
    rPr = style_el.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        style_el.append(rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    for attr in (qn("w:asciiTheme"), qn("w:hAnsiTheme"), qn("w:cstheme")):
        if rFonts.get(attr) is not None:
            del rFonts.attrib[attr]
    half_pts = str(int(font_size_pt * 2))
    for tag in ("w:sz", "w:szCs"):
        el = rPr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            rPr.append(el)
        el.set(qn("w:val"), half_pts)
    for tag in ("w:b", "w:bCs"):
        el = rPr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            rPr.append(el)
        el.set(qn("w:val"), "1" if bold else "0")


def apply_numbering_runprops(para, font_name, font_size, bold=False, highlight_val=None):
    if not para._element.xpath(".//w:numPr"):
        return
    pPr = para._element.get_or_add_pPr()
    rPr = pPr.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        pPr.append(rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    half_pts = str(int(font_size) * 2)
    for tag in ("w:sz", "w:szCs"):
        el = rPr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            rPr.append(el)
        el.set(qn("w:val"), half_pts)
    b_elem = rPr.find(qn("w:b"))
    if bold:
        if b_elem is None:
            b_elem = OxmlElement("w:b")
            rPr.append(b_elem)
    else:
        if b_elem is not None:
            rPr.remove(b_elem)
    if highlight_val:
        highlight = rPr.find(qn("w:highlight"))
        if highlight is None:
            highlight = OxmlElement("w:highlight")
            rPr.append(highlight)
        highlight.set(qn("w:val"), highlight_val)


def get_config_for_type(ptype, config):
    font_key = ptype.lower() + "_font"
    size_key = ptype.lower() + "_size"
    font_name = config.get(font_key)
    font_size_raw = config.get(size_key)
    if not font_name:
        if ptype == "COVER_PAGE":
            font_name = config.get("title_font") or config.get("paragraph_font") or "Calibri"
        elif ptype in ("BULLET_ITEM", "NUMBERED_ITEM"):
            font_name = config.get("list_item_font") or config.get("paragraph_font") or "Calibri"
        else:
            font_name = config.get("paragraph_font") or "Calibri"
    if not font_size_raw:
        if ptype == "COVER_PAGE":
            font_size_raw = config.get("title_size") or 12
        elif ptype in ("BULLET_ITEM", "NUMBERED_ITEM"):
            font_size_raw = config.get("list_item_size") or config.get("paragraph_size") or 12
        else:
            font_size_raw = config.get("paragraph_size") or 12
    try:
        font_size = int(font_size_raw)
    except Exception:
        font_size = 12
    return font_name, font_size


def format_docx(input_path, elements, output_path, config):
    doc = Document(input_path)
    para_map = {e["para_idx"]: e for e in elements if "para_idx" in e}
    color_map = {
        WD_COLOR_INDEX.YELLOW: "yellow",
        WD_COLOR_INDEX.BRIGHT_GREEN: "green",
        WD_COLOR_INDEX.TURQUOISE: "cyan",
        WD_COLOR_INDEX.PINK: "magenta",
        WD_COLOR_INDEX.VIOLET: "magenta",
        WD_COLOR_INDEX.TEAL: "cyan",
        WD_COLOR_INDEX.DARK_BLUE: "darkBlue",
        WD_COLOR_INDEX.GRAY_25: "lightGray",
        WD_COLOR_INDEX.GRAY_50: "darkGray",
        WD_COLOR_INDEX.DARK_YELLOW: "darkYellow",
    }
    STYLE_TYPE_MAP = {
        "TITLE":         ["Title"],
        "COVER_PAGE":    ["Subtitle"],
        "TOC_TITLE":     ["TOC Heading"],
        "TOC_HEADING_1": ["TOC 1"],
        "TOC_HEADING_2": ["TOC 2"],
        "TOC_HEADING_3": ["TOC 3"],
        "TOC_HEADING_4": ["TOC 4"],
        "HEADING_1":     ["Heading 1"],
        "HEADING_2":     ["Heading 2"],
        "HEADING_3":     ["Heading 3"],
        "HEADING_4":     ["Heading 4"],
        "HEADING_5":     ["Heading 5"],
        "HEADING_6":     ["Heading 6"],
        "PARAGRAPH":     ["Normal"],
        "BULLET_ITEM":   ["List Bullet", "List Paragraph"],
        "NUMBERED_ITEM": ["List Number", "List Paragraph"],
    }
    patched_styles = set()
    for ptype, style_names in STYLE_TYPE_MAP.items():
        fn, fs = get_config_for_type(ptype, config)
        is_bold = (
            (ptype.startswith("HEADING_") and config.get("bold_headings", True)) or
            (ptype in ("TITLE", "COVER_PAGE") and config.get("bold_titles", True)) or
            (ptype == "TOC_TITLE" and config.get("bold_toc", True)) or
            (ptype in ("BULLET_ITEM", "NUMBERED_ITEM", "LIST_ITEM") and config.get("bold_lists", False))
        )
        for sname in style_names:
            if sname not in patched_styles:
                _patch_doc_style(doc, sname, fn, fs, is_bold)
                patched_styles.add(sname)

    for idx, para in enumerate(doc.paragraphs):
        elem_info = para_map.get(idx)
        if not elem_info:
            continue
        ptype = elem_info["type"]
        original_indent = elem_info.get("indent", 0.0)
        font_name, font_size = get_config_for_type(ptype, config)
        is_heading   = ptype.startswith("HEADING_")
        is_title     = ptype == "TITLE"
        is_toc_title = ptype == "TOC_TITLE"
        is_cover     = ptype == "COVER_PAGE"
        is_list      = ptype in ("BULLET_ITEM", "NUMBERED_ITEM", "LIST_ITEM")
        bold_run = (
            (is_heading   and config.get("bold_headings", True)) or
            (is_title     and config.get("bold_titles", True))   or
            (is_cover     and config.get("bold_titles", True))   or
            (is_toc_title and config.get("bold_toc", True))      or
            (is_list      and config.get("bold_lists", False))
        )
        highlight_color_name = None
        if config.get("highlight", False):
            hc = HIGHLIGHT_COLORS.get(ptype, WD_COLOR_INDEX.GRAY_25)
            highlight_color_name = color_map.get(hc)
        apply_para_direct_format(para, font_name, font_size, bold_run, highlight_color_name)
        if is_heading:
            text = (para.text or "").strip()
            lvl, num, title = detect_numbered_heading(text)
            if num and title:
                clear_paragraph_runs(para)
                para.add_run(f"{num}\t")
                para.add_run(title)
        for r_el in _get_all_runs(para):
            if r_el.xpath(".//w:drawing"):
                continue
            rPr = r_el.find(qn("w:rPr"))
            if rPr is None:
                rPr = OxmlElement("w:rPr")
                r_el.insert(0, rPr)
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.insert(0, rFonts)
            rFonts.set(qn("w:ascii"), font_name)
            rFonts.set(qn("w:hAnsi"), font_name)
            rFonts.set(qn("w:cs"), font_name)
            rFonts.set(qn("w:eastAsia"), font_name)
            for attr in (qn("w:asciiTheme"), qn("w:hAnsiTheme"), qn("w:cstheme")):
                if rFonts.get(attr) is not None:
                    del rFonts.attrib[attr]
            half_pts = str(int(font_size * 2))
            for tag in ("w:sz", "w:szCs"):
                el = rPr.find(qn(tag))
                if el is None:
                    el = OxmlElement(tag)
                    rPr.append(el)
                el.set(qn("w:val"), half_pts)
            for tag in ("w:b", "w:bCs"):
                el = rPr.find(qn(tag))
                if el is None:
                    el = OxmlElement(tag)
                    rPr.append(el)
                el.set(qn("w:val"), "1" if bold_run else "0")
            if highlight_color_name:
                hl = rPr.find(qn("w:highlight"))
                if hl is None:
                    hl = OxmlElement("w:highlight")
                    rPr.append(hl)
                hl.set(qn("w:val"), highlight_color_name)
        if is_list:
            apply_numbering_runprops(para, font_name=font_name, font_size=font_size,
                                     bold=bold_run, highlight_val=highlight_color_name)
        if config.get("preserve_indentation", True) and original_indent and original_indent > 0:
            para.paragraph_format.left_indent = Inches(original_indent)

    for e in elements:
        if e["type"] == "TABLE" and e.get("table_idx", -1) < len(doc.tables):
            table = doc.tables[e["table_idx"]]
            font_name, font_size = get_config_for_type("TABLE", config)
            half_pts = str(int(font_size * 2))
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        apply_para_direct_format(para, font_name, font_size, False,
                                                 "cyan" if config.get("highlight") else None)
                        for r_el in _get_all_runs(para):
                            if r_el.xpath(".//w:drawing"):
                                continue
                            rPr = r_el.find(qn("w:rPr"))
                            if rPr is None:
                                rPr = OxmlElement("w:rPr")
                                r_el.insert(0, rPr)
                            rFonts = rPr.find(qn("w:rFonts"))
                            if rFonts is None:
                                rFonts = OxmlElement("w:rFonts")
                                rPr.insert(0, rFonts)
                            rFonts.set(qn("w:ascii"), font_name)
                            rFonts.set(qn("w:hAnsi"), font_name)
                            rFonts.set(qn("w:cs"), font_name)
                            rFonts.set(qn("w:eastAsia"), font_name)
                            for tag in ("w:sz", "w:szCs"):
                                el = rPr.find(qn(tag))
                                if el is None:
                                    el = OxmlElement(tag)
                                    rPr.append(el)
                                el.set(qn("w:val"), half_pts)
                            if config.get("highlight", False):
                                hl = rPr.find(qn("w:highlight"))
                                if hl is None:
                                    hl = OxmlElement("w:highlight")
                                    rPr.append(hl)
                                hl.set(qn("w:val"), "cyan")
    doc.save(output_path)




# ================== ALIGNMENT & FORMATTING UTILITIES ==================

def al_analyse_docx(path):
    doc = Document(path)
    elements = []
    types_present = set()
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        is_list = bool(p._element.xpath('.//w:numPr'))
        if is_list:
            elements.append({"type": "list", "index": i})
            types_present.add("list")
        else:
            elements.append({"type": "paragraph", "index": i})
            types_present.add("paragraph")
    return elements, sorted(types_present)


def al_format_docx(input_path, output_path, detected, cfg):
    doc = Document(input_path)
    alignment_map = {
        "left":    WD_ALIGN_PARAGRAPH.LEFT,
        "right":   WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    idx_map = {d["index"]: d["type"] for d in detected}
    for i, p in enumerate(doc.paragraphs):
        etype = idx_map.get(i)
        if not etype:
            continue
        pf = p.paragraph_format
        align_value = cfg.get(f"{etype}_alignment", "left")
        p.alignment = alignment_map.get(align_value, WD_ALIGN_PARAGRAPH.LEFT)
        spacing_value = float(cfg.get(f"{etype}_spacing", 1.0))
        pf.line_spacing = spacing_value
        if etype == "paragraph":
            options = cfg.get("paragraph_spacing_options", [])
            pf.space_before = None
            pf.space_after  = None
            if "add_before"    in options: pf.space_before = Pt(12)
            if "add_after"     in options: pf.space_after  = Pt(12)
            if "remove_before" in options: pf.space_before = Pt(0)
            if "remove_after"  in options: pf.space_after  = Pt(0)
            pf.line_spacing_rule = None
    doc.save(output_path)

# ================== GRAMMAR CHECK UTILITIES ==================

def gc_allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() == 'docx'


def gc_split_into_sentences(text):
    return re.split(r'(?<=[.!?])\s+(?=[A-Z<])', text.strip())


def gc_ai_correct_sentence(sentence):
    if not LANGUAGE_TOOL_AVAILABLE:
        return sentence
    matches = _grammar_tool.check(sentence)
    corrected = sentence
    for m in sorted(matches, key=lambda x: x.offset, reverse=True):
        length = getattr(m, 'errorLength', getattr(m, 'error_length', 0))
        replacement = m.replacements[0] if m.replacements else corrected[m.offset:m.offset + length]
        corrected = corrected[:m.offset] + replacement + corrected[m.offset + length:]
    return corrected


def gc_split_identifier(text):
    if not text or ' ' in text:
        return text
    text = re.sub(r'([A-Z]+)([A-Z][a-z])', r' ', text)
    text = re.sub(r'([a-z0-9])([A-Z])', r' ', text)
    return text.strip()


def gc_smart_correct(original, corrected):
    original = original.strip()
    corrected = corrected.strip()
    if original == corrected:
        split_text = gc_split_identifier(original)
        if split_text != original and len(split_text.split()) > 1:
            return split_text
        return None
    return corrected


def gc_check_with_word_com(file_path):
    pythoncom.CoInitialize()
    word = doc = None
    try:
        word = win32com.client.Dispatch('Word.Application')
        word.Visible = False
        word.DisplayAlerts = 0          # wdAlertsNone — suppress all dialogs
        word.Options.CheckSpellingAsYouType = False
        word.Options.CheckGrammarAsYouType = False
        # ConfirmConversions=False, ReadOnly=True, AddToRecentFiles=False,
        # Revert=False suppresses all Word dialogs that would hang the server
        doc = word.Documents.Open(
            os.path.abspath(file_path),
            False,   # ConfirmConversions
            True,    # ReadOnly
            False,   # AddToRecentFiles
            '',      # PasswordDocument
            '',      # PasswordTemplate
            False,   # Revert
        )

        sentences = {}
        spelling = grammar = 0

        for para in doc.Paragraphs:
            text = para.Range.Text.strip()
            if not text:
                continue
            for s in gc_split_into_sentences(text):
                sentences[s] = {'original': s, 'corrected': s, 'errors': []}

        for err in doc.SpellingErrors:
            spelling += 1
            for s in sentences:
                if err.Text in s:
                    corrected_sentence = gc_ai_correct_sentence(sentences[s]['corrected'])
                    final_corrected = gc_smart_correct(sentences[s]['original'], corrected_sentence)
                    if final_corrected is None:
                        continue
                    sentences[s]['corrected'] = final_corrected
                    sentences[s]['errors'].append({'type': 'SPELLING', 'text': err.Text})
                    break

        for err in doc.GrammaticalErrors:
            grammar += 1
            for s in sentences:
                if err.Text in s:
                    corrected_sentence = gc_ai_correct_sentence(sentences[s]['corrected'])
                    final_corrected = gc_smart_correct(sentences[s]['original'], corrected_sentence)
                    if final_corrected is None:
                        continue
                    sentences[s]['corrected'] = final_corrected
                    sentences[s]['errors'].append({'type': 'GRAMMAR', 'text': err.Text})
                    break

        results = []
        idx = 0
        for v in sentences.values():
            if not v['errors'] or v['original'] == v['corrected']:
                continue
            results.append({
                'id': idx,
                'original': v['original'],
                'corrected': v['corrected'],
                'current': v['original'],
                'error_details': v['errors']
            })
            idx += 1

        return results, spelling, grammar
    finally:
        if doc:
            doc.Close(False)
        if word:
            word.Quit()
        pythoncom.CoUninitialize()

# ================== ROUTES ==================

@app.route("/")
def home():
    return render_template_string(MAIN_HTML)


@app.route("/analyze_structure", methods=["POST"])
def analyze_structure():
    file = request.files.get("file")
    if not file or not file.filename.lower().endswith(".docx"):
        return jsonify({"error": "Invalid file"}), 400
    original_name = secure_filename(file.filename)
    input_path = os.path.join(app.config["UPLOAD_FOLDER"], original_name)
    file.save(input_path)
    analysis = analyze_document_structure(input_path)
    return jsonify(analysis)


@app.route("/format", methods=["POST"])
def format_document():
    file = request.files.get("file")
    config_raw = request.form.get("config", "{}")
    try:
        config = json.loads(config_raw)
    except Exception:
        return jsonify({"error": "Invalid config JSON"}), 400
    if not file or not file.filename.lower().endswith(".docx"):
        return jsonify({"error": "Invalid file"}), 400
    original_name = secure_filename(file.filename)
    base, ext = os.path.splitext(original_name)
    input_path = os.path.join(app.config["UPLOAD_FOLDER"], original_name)
    output_filename = f"{base}_formatted{ext}"
    output_path = os.path.join(app.config["OUTPUT_FOLDER"], output_filename)
    file.save(input_path)
    analysis = analyze_document_structure(input_path)
    format_docx(input_path, analysis["elements"], output_path, config)
    return send_file(output_path, as_attachment=True, download_name=output_filename)


@app.route("/al_analyse", methods=["POST"])
def al_analyse():
    f = request.files.get("file")
    if not f:
        return jsonify({"error": "No file"}), 400
    path = os.path.join(app.config["UPLOAD_FOLDER"], secure_filename(f.filename))
    f.save(path)
    detected, types = al_analyse_docx(path)
    return jsonify({"types": types})


@app.route("/al_format", methods=["POST"])
def al_format():
    f = request.files.get("file")
    cfg = json.loads(request.form.get("config", "{}"))
    if not f:
        return jsonify({"error": "No file"}), 400
    in_path  = os.path.join(app.config["UPLOAD_FOLDER"], secure_filename(f.filename))
    out_path = os.path.join(app.config["OUTPUT_FOLDER"], "formatted_alignment.docx")
    f.save(in_path)
    detected, _ = al_analyse_docx(in_path)
    al_format_docx(in_path, out_path, detected, cfg)
    return send_file(out_path, as_attachment=True, download_name="formatted_alignment.docx")


@app.route("/gc_upload", methods=["POST"])
def gc_upload():
    import time
    file = request.files.get("file")
    if not file or not gc_allowed_file(file.filename):
        return jsonify({"success": False, "error": "Invalid file"})

    # Seek to start in case stream was partially read
    file.stream.seek(0)

    # Use a timestamped name to avoid stale/locked file collisions
    base_name = secure_filename(file.filename)
    unique_name = f"gc_{int(time.time())}_{base_name}"
    path = os.path.abspath(os.path.join(app.config["UPLOAD_FOLDER"], unique_name))
    file.save(path)

    # Verify the saved file is a valid non-empty docx before handing to Word
    if not os.path.exists(path) or os.path.getsize(path) == 0:
        return jsonify({"success": False, "error": "File was empty after upload. Please try again."})

    try:
        if WORD_AVAILABLE:
            results, s, g = gc_check_with_word_com(path)
        else:
            return jsonify({"success": False, "error": "Word COM not available on this platform. Install pywin32 on Windows."})
        return jsonify({"success": True, "results": results, "spelling_errors": s, "grammar_errors": g})
    except Exception as e:
        return jsonify({"success": False, "error": f"Grammar check failed: {str(e)}"})
    finally:
        try:
            if os.path.exists(path):
                os.remove(path)
        except Exception:
            pass


if __name__ == "__main__":
    app.run(debug=True, port=5000)
