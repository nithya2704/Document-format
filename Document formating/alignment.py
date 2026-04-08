import os
import json
import uuid
import platform
import html as _html

from flask import (Blueprint, request, render_template,
                   jsonify, send_file, session, current_app)
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn
from werkzeug.utils import secure_filename

IS_WINDOWS = platform.system() == "Windows"
WORD_AVAILABLE = False
if IS_WINDOWS:
    try:
        import pythoncom
        from docx2pdf import convert
        WORD_AVAILABLE = True
    except ImportError:
        pass

alignment_bp = Blueprint(
    "alignment",
    __name__,
    template_folder="templates",
    url_prefix="/alignment"
)

OUTPUT = "outputs/alignment"
os.makedirs(OUTPUT, exist_ok=True)


# ── Session store helpers ─────────────────────────────────────────────────

def _get_path():
    sid = session.get("sid")
    if not sid:
        return None
    return current_app.config["GET_WORKING_PATH"](sid)


def _set_path(path: str):
    sid = session.get("sid")
    if sid:
        current_app.config["SET_WORKING_PATH"](sid, path)


# ── Cover page detection ──────────────────────────────────────────────────

_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
_NS = {'w': _W}


def _tag(name):
    """Return Clark-notation tag for a w: element."""
    return f'{{{_W}}}{name}'


def _find_cover_page_end(doc):
    """
    Return (first_para_index_after_cover, first_table_index_after_cover).

    Strategy — walk body children in document order, counting paragraphs
    and tables.  The cover page ends at whichever comes first:

    1. A paragraph containing an explicit page-break run  <w:br w:type="page"/>
    2. A paragraph whose pPr contains <w:pageBreakBefore/>
    3. A paragraph whose pPr contains a <w:sectPr> (section break)
    4. A paragraph containing a <w:lastRenderedPageBreak/> rendering hint
    5. The body-level <w:sectPr> that ends the first section — Word stores
       this as the last child of <w:body> but we detect it by checking
       whether the very first <w:sectPr> child of <w:body> (not nested in
       a paragraph) appears before we've seen more than ~40 paragraphs.
       In practice cover pages end at the first sectPr-carrying paragraph.

    Returns (0, 0) if no cover page boundary is detected.
    """
    para_idx = 0
    tbl_idx = 0

    for child in doc.element.body:
        local = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if local == 'p':
            try:
                has_page_br = bool(child.xpath(
                    './/w:br[@w:type="page"]', namespaces=_NS
                ))
                has_pbefore = bool(child.xpath(
                    'w:pPr/w:pageBreakBefore', namespaces=_NS
                ))
                has_sect = bool(child.xpath(
                    'w:pPr/w:sectPr', namespaces=_NS
                ))
                has_last_rendered = bool(child.xpath(
                    './/w:lastRenderedPageBreak', namespaces=_NS
                ))
            except Exception:
                has_page_br = has_pbefore = has_sect = has_last_rendered = False

            para_idx += 1

            if has_page_br or has_pbefore or has_sect or has_last_rendered:
                return para_idx, tbl_idx

        elif local == 'tbl':
            tbl_idx += 1

        elif local == 'sectPr':
            # A body-level sectPr that is NOT inside a paragraph means the
            # first section (cover page) has ended.  Only treat it as a cover
            # page boundary if we are still early in the document (para_idx
            # is small), to avoid false-positives on documents that have
            # section breaks later in the body.
            if para_idx <= 60:
                return para_idx, tbl_idx

    return 0, 0   # no cover page detected


# ── Analyse ───────────────────────────────────────────────────────────────

def analyse_docx(path):
    doc = Document(path)
    elements = []
    types_present = set()

    cover_para_end, cover_tbl_end = _find_cover_page_end(doc)

    # Detect paragraphs and lists (skip cover page)
    for i, p in enumerate(doc.paragraphs):
        if i < cover_para_end:
            continue
        text = p.text.strip()
        if not text:
            continue
        is_list = bool(p._element.xpath('.//w:numPr'))
        if is_list:
            elements.append({"type": "list",      "index": i})
            types_present.add("list")
        else:
            elements.append({"type": "paragraph", "index": i})
            types_present.add("paragraph")

    # Detect tables (skip cover page)
    if doc.tables:
        for i, tbl in enumerate(doc.tables):
            if i < cover_tbl_end:
                continue
            elements.append({"type": "table", "index": i})
        types_present.add("table")

    return elements, sorted(types_present)


# ── Format ────────────────────────────────────────────────────────────────

def format_docx(input_path, output_path, detected, cfg):
    doc = Document(input_path)

    alignment_map = {
        "left":    WD_ALIGN_PARAGRAPH.LEFT,
        "right":   WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }

    # Monospace font families that indicate code/preformatted content.
    _CODE_FONTS = {
        "courier", "courier new", "consolas", "lucida console",
        "monaco", "menlo", "dejavu sans mono", "inconsolata",
        "source code pro", "roboto mono", "cascadia code",
    }
    _CODE_STYLE_KEYWORDS = {"code", "preformat",
                            "verbatim", "monospace", "html"}

    def _is_code_paragraph(p) -> bool:
        """Return True if this paragraph looks like code / preformatted text."""
        # 1. Style name contains a code-related keyword
        sn = (p.style.name or "").lower() if p.style else ""
        if any(kw in sn for kw in _CODE_STYLE_KEYWORDS):
            return True
        # 2. Any run uses a monospace font
        for run in p.runs:
            rn = run.font.name or ""
            if rn.lower() in _CODE_FONTS:
                return True
            # Check rPr/rFonts in XML directly
            try:
                rpr = run._element.find(qn("w:rPr"))
                if rpr is not None:
                    rfonts = rpr.find(qn("w:rFonts"))
                    if rfonts is not None:
                        for attr in rfonts.attrib.values():
                            if attr.lower() in _CODE_FONTS:
                                return True
            except Exception:
                pass
        # 3. Text starts with < or looks like XML/code
        text = p.text.lstrip()
        if text.startswith("<") or text.startswith("//") or text.startswith("/*"):
            return True
        return False

    cover_para_end, cover_tbl_end = _find_cover_page_end(doc)

    # Build a set of paragraph element ids that belong to the cover page,
    # so we can skip them even when encountered inside table cells.
    cover_para_ids = {id(p._element) for p in doc.paragraphs[:cover_para_end]}

    # Build index maps for paragraphs/lists vs tables separately
    para_idx_map = {d["index"]: d["type"]
                    for d in detected if d["type"] in ("paragraph", "list")}
    table_indices = [d["index"] for d in detected if d["type"] == "table"]

    # ── Paragraphs & lists (skip cover page) ─────────────────────────────
    for i, p in enumerate(doc.paragraphs):
        if i < cover_para_end:
            continue
        etype = para_idx_map.get(i)
        if not etype:
            continue

        pf = p.paragraph_format

        align_value = cfg.get(f"{etype}_alignment", "")
        if align_value:
            # Never justify code / preformatted paragraphs
            if align_value == "justify" and _is_code_paragraph(p):
                pass
            else:
                p.alignment = alignment_map.get(align_value)

        spacing_value = cfg.get(f"{etype}_spacing", "")
        if spacing_value:
            pf.line_spacing = float(spacing_value)

        if etype == "paragraph":
            options = cfg.get("paragraph_spacing_options", [])
            if "add_before" in options:
                pf.space_before = Pt(12)
            if "add_after" in options:
                pf.space_after = Pt(12)
            if "remove_before" in options:
                pf.space_before = Pt(0)
            if "remove_after" in options:
                pf.space_after = Pt(0)

        if etype == "list":
            options = cfg.get("list_spacing_options", [])
            if "add_before" in options:
                pf.space_before = Pt(12)
            if "add_after" in options:
                pf.space_after = Pt(12)
            if "remove_before" in options:
                pf.space_before = Pt(0)
            if "remove_after" in options:
                pf.space_after = Pt(0)

    # ── Tables (skip cover page) ──────────────────────────────────────────
    table_align_value = cfg.get("table_alignment", "")
    table_spacing_value = cfg.get("table_spacing", "")
    table_options = cfg.get("table_spacing_options", [])

    if table_indices and doc.tables:
        for tbl_i, tbl in enumerate(doc.tables):
            if tbl_i < cover_tbl_end:
                continue
            if tbl_i not in table_indices:
                continue
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if id(p._element) in cover_para_ids:
                            continue
                        pf = p.paragraph_format
                        if table_align_value:
                            # Never justify code / preformatted paragraphs
                            if table_align_value == "justify" and _is_code_paragraph(p):
                                pass
                            else:
                                p.alignment = alignment_map.get(
                                    table_align_value)
                        if table_spacing_value:
                            pf.line_spacing = float(table_spacing_value)
                        if "add_before" in table_options:
                            pf.space_before = Pt(12)
                        if "add_after" in table_options:
                            pf.space_after = Pt(12)
                        if "remove_before" in table_options:
                            pf.space_before = Pt(0)
                        if "remove_after" in table_options:
                            pf.space_after = Pt(0)

    doc.save(output_path)


# ── HTML preview fallback ─────────────────────────────────────────────────

def _docx_to_html(docx_path: str) -> str:
    doc = Document(docx_path)
    page_css = (
        '<!DOCTYPE html><html><head><meta charset="utf-8"><style>'
        'body{font-family:Arial,sans-serif;margin:0;padding:0;background:#e8e8e8;}'
        '.page{background:white;width:794px;min-height:1123px;margin:30px auto;'
        'padding:72px 80px;box-shadow:0 2px 12px rgba(0,0,0,0.18);box-sizing:border-box;}'
        'h1,h2,h3,h4,h5,h6{margin:0.4em 0 0.2em;}'
        'p{margin:0.3em 0;line-height:1.5;}'
        'ul,ol{margin:0.5em 0 0.5em 1.8em;padding:0;}'
        'li{margin:0.2em 0;line-height:1.5;}'
        'table{border-collapse:collapse;width:100%;margin:0.8em 0;}'
        'td,th{border:1px solid #c8c8c8;padding:6px 10px;vertical-align:top;font-size:0.95em;}'
        'tr:nth-child(even) td{background:#f9f9f9;}'
        '</style></head><body><div class="page">'
    )
    parts = [page_css]

    # Build a set of paragraph elements that belong to tables so we can skip
    # them when iterating doc.paragraphs (they appear in both places via lxml).
    table_para_ids = set()
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    table_para_ids.add(id(p._element))

    # We iterate the body's direct children to preserve document order
    # (paragraphs and tables interleaved).
    from docx.oxml.ns import qn as _qn
    body = doc.element.body

    in_list = False
    para_index = 0  # track which doc.paragraphs entry we're on
    tbl_index = 0  # track which doc.tables entry we're on

    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if tag == 'p':
            # Grab the matching paragraph object
            if para_index >= len(doc.paragraphs):
                para_index += 1
                continue
            para = doc.paragraphs[para_index]
            para_index += 1

            text = (para.text or "").strip()
            if not text:
                if in_list:
                    parts.append("</ul>")
                    in_list = False
                parts.append("<p>&nbsp;</p>")
                continue

            is_list_item = bool(
                para._element.find(_qn("w:pPr")) is not None
                and para._element.find(_qn("w:pPr")).find(_qn("w:numPr")) is not None
            )
            sn = para.style.name.lower() if para.style else ""
            if is_list_item:
                if not in_list:
                    parts.append("<ul>")
                    in_list = True
                parts.append(f"<li>{_html.escape(text)}</li>")
            else:
                if in_list:
                    parts.append("</ul>")
                    in_list = False
                tag_name = "p"
                for i in range(1, 7):
                    if f"heading {i}" in sn:
                        tag_name = f"h{i}"
                        break
                align = ""
                if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    align = ' style="text-align:center"'
                elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                    align = ' style="text-align:right"'
                elif para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                    align = ' style="text-align:justify"'
                parts.append(
                    f"<{tag_name}{align}>{_html.escape(text)}</{tag_name}>")

        elif tag == 'tbl':
            if in_list:
                parts.append("</ul>")
                in_list = False
            if tbl_index >= len(doc.tables):
                tbl_index += 1
                continue
            tbl = doc.tables[tbl_index]
            tbl_index += 1

            parts.append('<table>')
            for r_i, row in enumerate(tbl.rows):
                parts.append('<tr>')
                for cell in row.cells:
                    cell_text = " ".join(
                        p.text for p in cell.paragraphs).strip()
                    # Detect alignment from first paragraph in cell
                    cell_align = ""
                    if cell.paragraphs:
                        cp = cell.paragraphs[0]
                        if cp.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                            cell_align = ' style="text-align:center"'
                        elif cp.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                            cell_align = ' style="text-align:right"'
                        elif cp.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                            cell_align = ' style="text-align:justify"'
                    cell_tag = 'th' if r_i == 0 else 'td'
                    parts.append(
                        f'<{cell_tag}{cell_align}>{_html.escape(cell_text)}</{cell_tag}>')
                parts.append('</tr>')
            parts.append('</table>')

    if in_list:
        parts.append("</ul>")
    parts.append("</div></body></html>")
    return "".join(parts)


def _session_out_dir():
    sid = session.get("sid", "default")
    d = os.path.join(OUTPUT, sid)
    os.makedirs(d, exist_ok=True)
    return d


# ── Routes ────────────────────────────────────────────────────────────────

@alignment_bp.route("/")
def home():
    return render_template("alignmentUI.html")


@alignment_bp.route("/analyse", methods=["POST"])
def analyse():
    path = _get_path()
    if not path or not os.path.exists(path):
        return jsonify({"error": "No file uploaded yet. Please upload from the home page."}), 400
    try:
        _, types = analyse_docx(path)
        return jsonify({"types": types})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@alignment_bp.route("/original", methods=["POST"])
def original():
    path = _get_path()
    if not path or not os.path.exists(path):
        return jsonify({"error": "No file uploaded yet. Please upload from the home page."}), 400

    try:
        if IS_WINDOWS and WORD_AVAILABLE:
            try:
                pythoncom.CoInitialize()
                pdf_path = os.path.join(
                    _session_out_dir(), f"orig_{uuid.uuid4().hex}.pdf")
                convert(path, pdf_path)
                pythoncom.CoUninitialize()
                return send_file(pdf_path, mimetype="application/pdf")
            except Exception:
                try:
                    pythoncom.CoUninitialize()
                except Exception:
                    pass
        html_content = _docx_to_html(path)
        return html_content, 200, {"Content-Type": "text/html; charset=utf-8"}
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@alignment_bp.route("/preview", methods=["POST"])
def preview():
    path = _get_path()
    if not path or not os.path.exists(path):
        return jsonify({"error": "No file uploaded yet. Please upload from the home page."}), 400

    try:
        cfg = json.loads(request.form.get("config", "{}"))
        out_docx = os.path.join(
            _session_out_dir(), f"preview_{uuid.uuid4().hex}.docx")
        detected, _ = analyse_docx(path)
        format_docx(path, out_docx, detected, cfg)

        if IS_WINDOWS and WORD_AVAILABLE:
            pdf_path = os.path.join(
                _session_out_dir(), f"preview_{uuid.uuid4().hex}.pdf")
            pythoncom.CoInitialize()
            try:
                convert(out_docx, pdf_path)
            finally:
                pythoncom.CoUninitialize()
            return send_file(pdf_path, mimetype="application/pdf")

        # Non-Windows or Word not installed — fall back to HTML renderer
        html_content = _docx_to_html(out_docx)
        return html_content, 200, {"Content-Type": "text/html; charset=utf-8"}
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@alignment_bp.route("/format", methods=["POST"])
def format_doc():
    path = _get_path()
    if not path or not os.path.exists(path):
        return jsonify({"error": "No file uploaded yet. Please upload from the home page."}), 400
    try:
        cfg = json.loads(request.form.get("config", "{}"))
        out_path = os.path.join(
            _session_out_dir(), f"aligned_{uuid.uuid4().hex}.docx")
        detected, _ = analyse_docx(path)
        format_docx(path, out_path, detected, cfg)

        # Update working file so the next tool in the pipeline uses this output.
        _set_path(out_path)

        return send_file(
            out_path,
            as_attachment=True,
            download_name="formatted.docx",
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500


if __name__ == "__main__":
    from flask import Flask
    _app = Flask(__name__, template_folder="templates")
    _app.secret_key = "dev"
    _app.register_blueprint(alignment_bp)
    _app.run(debug=True, threaded=False)
