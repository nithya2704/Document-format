from flask import Flask, request, send_file, jsonify, render_template_string
from flask_cors import CORS
from docx import Document
from docx.shared import Pt
import os
from werkzeug.utils import secure_filename
import tempfile
import traceback

app = Flask(__name__)
CORS(app)

UPLOAD_FOLDER = tempfile.gettempdir()
ALLOWED_EXTENSIONS = {'docx'}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def format_docx(input_path, output_path, font_name, font_size):
    """Format the Word document with specified font and size"""
    try:
        doc = Document(input_path)

        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = font_name
                run.font.size = Pt(int(font_size))

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = font_name
                            run.font.size = Pt(int(font_size))

        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                for run in paragraph.runs:
                    run.font.name = font_name
                    run.font.size = Pt(int(font_size))

            for paragraph in section.footer.paragraphs:
                for run in paragraph.runs:
                    run.font.name = font_name
                    run.font.size = Pt(int(font_size))

        doc.save(output_path)
        return True
    except Exception as e:
        print(f"Error formatting document: {str(e)}")
        traceback.print_exc()
        return False


# HTML Template embedded in Python
HTML_TEMPLATE = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Word Document Formatter</title>
    <style>
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }

        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            min-height: 100vh;
            display: flex;
            justify-content: center;
            align-items: center;
            padding: 20px;
        }

        .container {
            background: white;
            border-radius: 20px;
            box-shadow: 0 20px 60px rgba(0, 0, 0, 0.3);
            max-width: 600px;
            width: 100%;
            padding: 40px;
        }

        h1 {
            color: #333;
            margin-bottom: 10px;
            font-size: 28px;
            text-align: center;
        }

        .subtitle {
            color: #666;
            text-align: center;
            margin-bottom: 30px;
            font-size: 14px;
        }

        .status-badge {
            background: #d4edda;
            color: #155724;
            padding: 8px 15px;
            border-radius: 20px;
            font-size: 12px;
            display: inline-block;
            margin-bottom: 20px;
        }

        .upload-area {
            border: 3px dashed #667eea;
            border-radius: 15px;
            padding: 40px;
            text-align: center;
            transition: all 0.3s ease;
            cursor: pointer;
            margin-bottom: 30px;
            background: #f8f9ff;
        }

        .upload-area:hover {
            border-color: #764ba2;
            background: #f0f1ff;
        }

        .upload-area.drag-over {
            border-color: #764ba2;
            background: #e8e9ff;
            transform: scale(1.02);
        }

        .upload-icon {
            font-size: 48px;
            margin-bottom: 10px;
        }

        .upload-text {
            color: #667eea;
            font-weight: 600;
            margin-bottom: 5px;
        }

        .upload-subtext {
            color: #999;
            font-size: 13px;
        }

        input[type="file"] {
            display: none;
        }

        .form-group {
            margin-bottom: 25px;
        }

        label {
            display: block;
            margin-bottom: 8px;
            color: #333;
            font-weight: 600;
            font-size: 14px;
        }

        select, input[type="number"] {
            width: 100%;
            padding: 12px 15px;
            border: 2px solid #e0e0e0;
            border-radius: 10px;
            font-size: 15px;
            transition: all 0.3s ease;
            background: white;
        }

        select:focus, input[type="number"]:focus {
            outline: none;
            border-color: #667eea;
            box-shadow: 0 0 0 3px rgba(102, 126, 234, 0.1);
        }

        .file-info {
            background: #f0f1ff;
            padding: 15px;
            border-radius: 10px;
            margin-bottom: 20px;
            display: none;
        }

        .file-info.show {
            display: block;
        }

        .file-name {
            color: #333;
            font-weight: 600;
            margin-bottom: 5px;
        }

        .file-size {
            color: #666;
            font-size: 13px;
        }

        .btn {
            width: 100%;
            padding: 15px;
            border: none;
            border-radius: 10px;
            font-size: 16px;
            font-weight: 600;
            cursor: pointer;
            transition: all 0.3s ease;
            margin-top: 10px;
        }

        .btn-primary {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
        }

        .btn-primary:hover:not(:disabled) {
            transform: translateY(-2px);
            box-shadow: 0 10px 25px rgba(102, 126, 234, 0.4);
        }

        .btn-primary:disabled {
            opacity: 0.5;
            cursor: not-allowed;
        }

        .preview-area {
            margin-top: 30px;
            padding: 20px;
            background: #f8f9ff;
            border-radius: 10px;
            display: none;
        }

        .preview-area.show {
            display: block;
        }

        .preview-title {
            color: #333;
            font-weight: 600;
            margin-bottom: 15px;
            font-size: 16px;
        }

        .preview-text {
            padding: 15px;
            background: white;
            border-radius: 8px;
            border: 1px solid #e0e0e0;
            min-height: 100px;
        }

        .alert {
            padding: 15px;
            border-radius: 10px;
            margin-bottom: 20px;
            display: none;
        }

        .alert.show {
            display: block;
        }

        .alert-success {
            background: #d4edda;
            color: #155724;
            border: 1px solid #c3e6cb;
        }

        .alert-error {
            background: #f8d7da;
            color: #721c24;
            border: 1px solid #f5c6cb;
        }

        .loading {
            text-align: center;
            padding: 20px;
            display: none;
        }

        .loading.show {
            display: block;
        }

        .spinner {
            border: 3px solid #f3f3f3;
            border-top: 3px solid #667eea;
            border-radius: 50%;
            width: 40px;
            height: 40px;
            animation: spin 1s linear infinite;
            margin: 0 auto 10px;
        }

        @keyframes spin {
            0% { transform: rotate(0deg); }
            100% { transform: rotate(360deg); }
        }
    </style>
</head>
<body>
    <div class="container">
        <h1>📄 Word Document Formatter</h1>
        <p class="subtitle">Upload a .docx file and customize its font styling</p>
        
        <div style="text-align: center;">
            <span class="status-badge">✅ Backend Connected</span>
        </div>

        <div id="alert" class="alert"></div>

        <div class="upload-area" id="uploadArea">
            <div class="upload-icon">📤</div>
            <div class="upload-text">Click to upload or drag and drop</div>
            <div class="upload-subtext">Word Document (.docx) only</div>
            <input type="file" id="fileInput" accept=".docx">
        </div>

        <div id="fileInfo" class="file-info">
            <div class="file-name" id="fileName"></div>
            <div class="file-size" id="fileSize"></div>
        </div>

        <div class="form-group">
            <label for="fontFamily">Font Family</label>
            <select id="fontFamily">
                <option value="Arial">Arial</option>
                <option value="Times New Roman">Times New Roman</option>
                <option value="Calibri" selected>Calibri</option>
                <option value="Georgia">Georgia</option>
                <option value="Verdana">Verdana</option>
                <option value="Courier New">Courier New</option>
                <option value="Comic Sans MS">Comic Sans MS</option>
                <option value="Tahoma">Tahoma</option>
                <option value="Trebuchet MS">Trebuchet MS</option>
            </select>
        </div>

        <div class="form-group">
            <label for="fontSize">Font Size (pt)</label>
            <input type="number" id="fontSize" value="12" min="8" max="72">
        </div>

        <button class="btn btn-primary" id="formatBtn" disabled>
            Format Document
        </button>

        <div id="loading" class="loading">
            <div class="spinner"></div>
            <p>Processing your document...</p>
        </div>

        <div id="previewArea" class="preview-area">
            <div class="preview-title">Preview (Sample text with selected formatting):</div>
            <div class="preview-text" id="previewText">
                This is how your text will look after formatting. The quick brown fox jumps over the lazy dog. Lorem ipsum dolor sit amet, consectetur adipiscing elit.
            </div>
        </div>
    </div>

    <script>
        const API_URL = '/api';
        let uploadedFile = null;

        const uploadArea = document.getElementById('uploadArea');
        const fileInput = document.getElementById('fileInput');
        const fileInfo = document.getElementById('fileInfo');
        const fileName = document.getElementById('fileName');
        const fileSize = document.getElementById('fileSize');
        const formatBtn = document.getElementById('formatBtn');
        const fontFamily = document.getElementById('fontFamily');
        const fontSize = document.getElementById('fontSize');
        const previewArea = document.getElementById('previewArea');
        const previewText = document.getElementById('previewText');
        const alert = document.getElementById('alert');
        const loading = document.getElementById('loading');

        uploadArea.addEventListener('click', () => fileInput.click());

        uploadArea.addEventListener('dragover', (e) => {
            e.preventDefault();
            uploadArea.classList.add('drag-over');
        });

        uploadArea.addEventListener('dragleave', () => {
            uploadArea.classList.remove('drag-over');
        });

        uploadArea.addEventListener('drop', (e) => {
            e.preventDefault();
            uploadArea.classList.remove('drag-over');
            const files = e.dataTransfer.files;
            if (files.length > 0) {
                handleFile(files[0]);
            }
        });

        fileInput.addEventListener('change', (e) => {
            if (e.target.files.length > 0) {
                handleFile(e.target.files[0]);
            }
        });

        function handleFile(file) {
            if (!file.name.endsWith('.docx')) {
                showAlert('Please upload a .docx file only', 'error');
                return;
            }

            uploadedFile = file;
            fileName.textContent = file.name;
            fileSize.textContent = `Size: ${(file.size / 1024).toFixed(2)} KB`;
            fileInfo.classList.add('show');
            formatBtn.disabled = false;
            showAlert('File uploaded successfully!', 'success');
            updatePreview();
        }

        function updatePreview() {
            previewText.style.fontFamily = fontFamily.value;
            previewText.style.fontSize = fontSize.value + 'pt';
            previewArea.classList.add('show');
        }

        fontFamily.addEventListener('change', updatePreview);
        fontSize.addEventListener('input', updatePreview);

        formatBtn.addEventListener('click', formatDocument);

        async function formatDocument() {
            if (!uploadedFile) return;

            loading.classList.add('show');
            formatBtn.disabled = true;

            try {
                const formData = new FormData();
                formData.append('file', uploadedFile);
                formData.append('fontFamily', fontFamily.value);
                formData.append('fontSize', fontSize.value);

                const response = await fetch(`${API_URL}/format`, {
                    method: 'POST',
                    body: formData
                });

                if (!response.ok) {
                    const error = await response.json();
                    throw new Error(error.error || 'Failed to format document');
                }

                const blob = await response.blob();
                const url = window.URL.createObjectURL(blob);
                const a = document.createElement('a');
                a.href = url;
                a.download = `formatted_${uploadedFile.name}`;
                document.body.appendChild(a);
                a.click();
                window.URL.revokeObjectURL(url);
                document.body.removeChild(a);

                showAlert('✅ Document formatted and downloaded successfully!', 'success');
                
            } catch (error) {
                showAlert('❌ Error: ' + error.message, 'error');
            } finally {
                loading.classList.remove('show');
                formatBtn.disabled = false;
            }
        }

        function showAlert(message, type) {
            alert.textContent = message;
            alert.className = 'alert show alert-' + type;
            setTimeout(() => {
                alert.classList.remove('show');
            }, 5000);
        }
    </script>
</body>
</html>
"""


@app.route('/')
def home():
    """Serve the HTML interface"""
    return render_template_string(HTML_TEMPLATE)


@app.route('/api/health', methods=['GET'])
def health_check():
    return jsonify({
        'status': 'healthy',
        'message': 'Word formatter API is running',
        'version': '1.0.0'
    })


@app.route('/api/format', methods=['POST'])
def format_document():
    try:
        print("=" * 50)
        print("Received format request")

        if 'file' not in request.files:
            print("Error: No file in request")
            return jsonify({'error': 'No file provided'}), 400

        file = request.files['file']
        font_name = request.form.get('fontFamily', 'Calibri')
        font_size = request.form.get('fontSize', '12')

        print(f"File: {file.filename}")
        print(f"Font: {font_name}, Size: {font_size}")

        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400

        if not allowed_file(file.filename):
            return jsonify({'error': 'Only .docx files are allowed'}), 400

        filename = secure_filename(file.filename)
        input_path = os.path.join(
            app.config['UPLOAD_FOLDER'], f"input_{filename}")
        file.save(input_path)

        output_filename = f"formatted_{filename}"
        output_path = os.path.join(
            app.config['UPLOAD_FOLDER'], output_filename)

        print("Formatting document...")
        success = format_docx(input_path, output_path, font_name, font_size)

        if not success:
            return jsonify({'error': 'Failed to format document'}), 500

        print("Sending formatted file...")
        response = send_file(
            output_path,
            as_attachment=True,
            download_name=output_filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

        @response.call_on_close
        def cleanup():
            try:
                if os.path.exists(input_path):
                    os.remove(input_path)
                if os.path.exists(output_path):
                    os.remove(output_path)
            except:
                pass

        print("✅ Success!")
        print("=" * 50)
        return response

    except Exception as e:
        print(f"❌ Error: {str(e)}")
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


if __name__ == '__main__':
    print("\n" + "=" * 60)
    print("🚀 Word Document Formatter Starting...")
    print("=" * 60)
    print("📍 Open your browser and go to:")
    print("   👉 http://localhost:5000")
    print("=" * 60)
    print("✅ Server is ready!")
    print("=" * 60 + "\n")

    app.run(debug=True, host='0.0.0.0', port=5000)
