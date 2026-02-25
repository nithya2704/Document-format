# smart_doc_formatter.py
from flask import Flask, render_template_string, request, send_file, jsonify
import os
import json
import copy
import re

from werkzeug.utils import secure_filename

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ================== FLASK SETUP ==================
app = Flask(__name__)
app.config["UPLOAD_FOLDER"] = "uploads"
app.config["OUTPUT_FOLDER"] = "outputs"
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024  # 16 MB

os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)
os.makedirs(app.config["OUTPUT_FOLDER"], exist_ok=True)


# ================== COLORS ==================
HIGHLIGHT_COLORS = {
    "TITLE": WD_COLOR_INDEX.YELLOW,
    "COVER_PAGE": WD_COLOR_INDEX.YELLOW,
    "TOC_TITLE": WD_COLOR_INDEX.YELLOW,
    "TOC_HEADING_1": WD_COLOR_INDEX.GRAY_25,
    "TOC_HEADING_2": WD_COLOR_INDEX.GRAY_50,
    "TOC_HEADING_3": WD_COLOR_INDEX.DARK_YELLOW,
    "TOC_HEADING_4": WD_COLOR_INDEX.DARK_YELLOW,
    "HEADING_1": WD_COLOR_INDEX.BRIGHT_GREEN,
    "HEADING_2": WD_COLOR_INDEX.TURQUOISE,
    "HEADING_3": WD_COLOR_INDEX.PINK,
    "HEADING_4": WD_COLOR_INDEX.VIOLET,
    "HEADING_5": WD_COLOR_INDEX.TEAL,
    "HEADING_6": WD_COLOR_INDEX.DARK_BLUE,
    "PARAGRAPH": WD_COLOR_INDEX.GRAY_25,
    "LIST_ITEM": WD_COLOR_INDEX.GRAY_25,
    "TABLE": WD_COLOR_INDEX.TURQUOISE,
}


# ================== UI ==================
HTML_TEMPLATE = r"""
<!DOCTYPE html>
<html>
<head>
<title>Smart Document Formatter</title>
<style>
body { background: linear-gradient(135deg,#f3f4f6,#e5e7eb); font-family: Arial; padding: 40px; }
.box { background: white; width: 750px; padding: 30px; border-radius: 15px; margin: auto; max-height: 90vh; overflow-y: auto; }
h1 { text-align:center; }
label { font-weight:bold; }
select,input { width:100%; padding:8px; margin-bottom:10px; }
button { width:100%; padding:14px; background:#000000; border:none; border-radius:30px; color:white; font-size:16px; cursor: pointer; }
button:disabled { background:#ccc; cursor: not-allowed; }
.info-box { background: #f0f0f0; padding: 15px; border-radius: 8px; margin-bottom: 20px; }
.element-list { list-style: none; padding-left: 0; }
.element-list li { padding: 5px 0; }
.element-list li.toc-item { color: #2563eb; font-weight: 500; }
.element-list li.list-item { color: #16a34a; font-weight: 500; }
#formatOptions { display: none; }
.loading { text-align: center; padding: 20px; color: #c7c6c1; }
.hierarchy-info { background: #e8f4f8; padding: 10px; border-left: 4px solid #c7c6c1; margin: 10px 0; font-size: 14px; }
.toc-section { background: #f0f9ff; padding: 15px; border-radius: 8px; margin: 15px 0; border: 2px solid #2563eb; }
.toc-section h3 { color: #2563eb; margin-top: 0; }
.list-section { background: #f0fdf4; padding: 15px; border-radius: 8px; margin: 15px 0; border: 2px solid #16a34a; }
.list-section h3 { color: #16a34a; margin-top: 0; }
.preview-container { display: flex; gap: 15px; align-items: center; }
.preview-box { border: 1px solid #d1d5db; border-radius: 6px; padding: 10px 14px; background: #fafafa; min-width: 220px; font-size: 14px; }
.preview-label { font-size: 12px; color: #6b7280; margin-bottom: 4px; }
.subbox { background: #fff; border: 1px solid #e5e7eb; border-radius: 8px; padding: 10px 12px; margin-top: 12px; }
.subbox h4 { margin: 0 0 8px 0; }
.small { color: #6b7280; font-size: 12px; }
</style>
</head>
<body>
<div class="box">
<h1>Smart Document Formatter</h1>
<input type="file" id="file" accept=".docx">
<br><br>
<button type="button" onclick="analyzeDoc()" id="analyzeBtn">Analyze Document</button>
<div id="loadingMsg" class="loading" style="display:none;">Analyzing document structure...</div>
<div id="analysisResults" style="display:none;">
  <h2>Document Analysis Complete</h2>
  <div class="info-box">
    <h3>Detected Elements:</h3>
    <ul id="elementsList" class="element-list"></ul>
    <div id="bulletsByHeadingBox" class="subbox" style="display:none;">
      <h4>List items (bullets + numbered) by Heading</h4>
      <div class="small">List items are categorized under their nearest preceding heading.</div>
      <ul id="bulletsByHeadingList" class="element-list" style="margin-top:8px;"></ul>
    </div>
  </div>
  <div id="hierarchyInfo" class="hierarchy-info" style="display:none;">
    <strong>Document Structure:</strong> <span id="hierarchyText"></span>
  </div>
</div>
<div id="formatOptions">
  <h2>Formatting Options</h2>
  <p>Configure styling for the detected elements:</p>
  <div id="styleControls"></div>
  <label><input type="checkbox" id="bold_titles" checked onchange="refreshAllPreviews()"> Bold Titles</label>
  <label><input type="checkbox" id="bold_headings" checked onchange="refreshAllPreviews()"> Bold Headings</label>
  <label><input type="checkbox" id="bold_toc" checked onchange="refreshAllPreviews()"> Bold TOC Title</label>
  <label><input type="checkbox" id="bold_lists" onchange="refreshAllPreviews()"> Bold List Items</label>
  <label><input type="checkbox" id="highlight" checked onchange="refreshAllPreviews()"> Highlight Sections</label>
  <label><input type="checkbox" id="preserve_indentation" checked>Preserve Indentation</label>
  <br><br>
  <button type="button" onclick="formatDoc()">Format Document</button>
</div>
</div>
<script>
let analysisData = null;
async function analyzeDoc(){
  let file = document.getElementById("file").files[0];
  if(!file){ alert("Please upload a DOCX file"); return; }
  document.getElementById("analyzeBtn").disabled = true;
  document.getElementById("loadingMsg").style.display = "block";
  let fd = new FormData();
  fd.append("file", file);
  try {
    let res = await fetch("/analyze_structure", {method:"POST", body:fd});
    analysisData = await res.json();
    if (analysisData.error) throw new Error(analysisData.error);
    displayAnalysis(analysisData);
    buildFormatControls(analysisData.detected_elements);
    document.getElementById("loadingMsg").style.display = "none";
    document.getElementById("analysisResults").style.display = "block";
    document.getElementById("formatOptions").style.display = "block";
  } catch(e) {
    alert("Error analyzing document: " + e);
    document.getElementById("analyzeBtn").disabled = false;
    document.getElementById("loadingMsg").style.display = "none";
  }
}
function refreshAllPreviews() {
  if (!analysisData) return;
  for (let elem of analysisData.detected_elements) updatePreview(elem);
}
function formatElementName(type) {
  return type.replace(/_/g, ' ').replace(/\b\w/g, l => l.toUpperCase());
}
function displayAnalysis(data) {
  let list = document.getElementById("elementsList");
  list.innerHTML = "";
  let tocElements = {};
  let listElements = {};
  let coverElements = {};
  let regularElements = {};
  for(let [elemType, count] of Object.entries(data.element_counts)) {
    if(elemType.startsWith("TOC_")) tocElements[elemType] = count;
    else if(elemType === "BULLET_ITEM" || elemType === "NUMBERED_ITEM" || elemType === "LIST_ITEM") listElements[elemType] = count;
    else if(elemType === "TITLE" || elemType === "COVER_PAGE") coverElements[elemType] = count;
    else regularElements[elemType] = count;
  }
  // Cover page items
  if(Object.keys(coverElements).length > 0) {
    let li = document.createElement("li");
    li.style.color = "#b45309"; li.style.fontWeight = "500";
    li.innerHTML = `<strong>Cover / Title Page Detected:</strong>`;
    list.appendChild(li);
    for(let [elemType, count] of Object.entries(coverElements)) {
      let subLi = document.createElement("li");
      subLi.style.paddingLeft = "20px"; subLi.style.color = "#b45309";
      subLi.innerHTML = `→ ${formatElementName(elemType)}: ${count} found`;
      list.appendChild(subLi);
    }
  }
  for(let [elemType, count] of Object.entries(regularElements)) {
    let li = document.createElement("li");
    li.innerHTML = `<strong>${formatElementName(elemType)}:</strong> ${count} found`;
    list.appendChild(li);
  }
  if(Object.keys(listElements).length > 0) {
    let li = document.createElement("li");
    li.className = "list-item";
    li.innerHTML = `<strong>Lists Detected:</strong>`;
    list.appendChild(li);
    for(let [elemType, count] of Object.entries(listElements)) {
      let subLi = document.createElement("li");
      subLi.className = "list-item";
      subLi.style.paddingLeft = "20px";
      subLi.innerHTML = `→ ${formatElementName(elemType)}: ${count} found`;
      list.appendChild(subLi);
    }
  }
  if(Object.keys(tocElements).length > 0) {
    let li = document.createElement("li");
    li.className = "toc-item";
    li.innerHTML = `<strong>Table of Contents Detected:</strong>`;
    list.appendChild(li);
    for(let [elemType, count] of Object.entries(tocElements)) {
      let subLi = document.createElement("li");
      subLi.className = "toc-item";
      subLi.style.paddingLeft = "20px";
      subLi.innerHTML = `→ ${formatElementName(elemType)}: ${count} found`;
      list.appendChild(subLi);
    }
  }
  // Lists by heading (bullets + numbered)
  let box = document.getElementById("bulletsByHeadingBox");
  let ul = document.getElementById("bulletsByHeadingList");
  ul.innerHTML = "";
  if (data.list_items_by_heading && Object.keys(data.list_items_by_heading).length > 0) {
    box.style.display = "block";
    const keys = Object.keys(data.list_items_by_heading).sort((a,b) => {
      if (a === "NO_HEADING") return 1;
      if (b === "NO_HEADING") return -1;
      const na = parseInt(a.split("_").pop(), 10);
      const nb = parseInt(b.split("_").pop(), 10);
      return (na||99) - (nb||99);
    });
    for (const k of keys) {
      const c = data.list_items_by_heading[k];
      const label = (k === "NO_HEADING") ? "No Heading" : formatElementName(k);
      const li = document.createElement("li");
      li.innerHTML = `<strong>${label}:</strong> ${c} list items`;
      ul.appendChild(li);
    }
  } else {
    box.style.display = "none";
  }
  if(data.has_toc || data.has_indented_hierarchy || data.cover_page_end >= 0) {
    let hierarchyDiv = document.getElementById("hierarchyInfo");
    let hierarchyText = document.getElementById("hierarchyText");
    let messages = [];
    if(data.cover_page_end >= 0) messages.push(`Cover page detected (${data.cover_page_end + 1} paragraph(s))`);
    if(data.has_toc) messages.push(`Table of Contents with ${data.toc_indent_levels} hierarchy levels`);
    if(data.has_indented_hierarchy) messages.push(`${data.indent_levels} indentation levels in document body`);
    hierarchyText.textContent = messages.join(" • ");
    hierarchyDiv.style.display = "block";
  }
}
function escapeHtml(text) {
  return (text || "")
    .replace(/&/g, "&amp;")
    .replace(/</g, "&lt;")
    .replace(/>/g, "&gt;")
    .replace(/\t/g, "&nbsp;&nbsp;&nbsp;&nbsp;");
}
function orderKey(elem) {
  if (elem === "TITLE") return 0;
  if (elem === "COVER_PAGE") return 1;
  if (elem === "TOC_TITLE") return 10;
  if (elem.startsWith("TOC_HEADING_")) {
    const n = parseInt(elem.split("_").pop(), 10);
    return 11 + (isNaN(n) ? 99 : n);
  }
  if (elem.startsWith("HEADING_")) {
    const n = parseInt(elem.split("_").pop(), 10);
    return 100 + (isNaN(n) ? 99 : n);
  }
  if (elem === "BULLET_ITEM") return 190;
  if (elem === "NUMBERED_ITEM") return 195;
  if (elem === "LIST_ITEM") return 200;
  if (elem === "PARAGRAPH") return 210;
  if (elem === "TABLE") return 220;
  return 999;
}
function sortElementsForUI(elements) {
  return [...elements].sort((a,b) => {
    const ka = orderKey(a), kb = orderKey(b);
    if (ka !== kb) return ka - kb;
    return a.localeCompare(b);
  });
}
function buildFormatControls(elements) {
  let container = document.getElementById("styleControls");
  container.innerHTML = "";
  const sorted = sortElementsForUI(elements);
  const defaults = {
    "TITLE": {font: "Times New Roman", size: 26},
    "COVER_PAGE": {font: "Times New Roman", size: 12},
    "TOC_TITLE": {font: "Calibri", size: 14},
    "TOC_HEADING_1": {font: "Calibri", size: 11},
    "TOC_HEADING_2": {font: "Calibri", size: 11},
    "TOC_HEADING_3": {font: "Calibri", size: 11},
    "TOC_HEADING_4": {font: "Calibri", size: 11},
    "HEADING_1": {font: "Calibri", size: 18},
    "HEADING_2": {font: "Calibri", size: 16},
    "HEADING_3": {font: "Calibri", size: 14},
    "HEADING_4": {font: "Calibri", size: 13},
    "HEADING_5": {font: "Calibri", size: 12},
    "HEADING_6": {font: "Calibri", size: 11},
    "PARAGRAPH": {font: "Calibri", size: 12},
    "BULLET_ITEM": {font: "Calibri", size: 12},
    "NUMBERED_ITEM": {font: "Calibri", size: 12},
    "LIST_ITEM": {font: "Calibri", size: 12},
    "TABLE": {font: "Calibri", size: 11}
  };
  const fontOptions = [
    "Times New Roman","Arial","Calibri","Georgia","Verdana","Tahoma",
    "Trebuchet MS","Courier New","Garamond","Palatino Linotype",
    "Book Antiqua","Comic Sans MS","Impact","Lucida Sans Unicode",
    "Century Gothic"
  ];
  const tocElements = sorted.filter(e => e.startsWith("TOC_"));
  const listElements = sorted.filter(e => e === "BULLET_ITEM" || e === "NUMBERED_ITEM" || e === "LIST_ITEM");
  const coverElements = sorted.filter(e => e === "TITLE" || e === "COVER_PAGE");
  const regularElements = sorted.filter(e =>
    !e.startsWith("TOC_") &&
    e !== "BULLET_ITEM" && e !== "NUMBERED_ITEM" && e !== "LIST_ITEM" &&
    e !== "TITLE" && e !== "COVER_PAGE"
  );

  // Cover/title page section
  if (coverElements.length > 0) {
    let coverSection = document.createElement("div");
    coverSection.style.cssText = "background:#fffbeb;padding:15px;border-radius:8px;margin:15px 0;border:2px solid #f59e0b;";
    coverSection.innerHTML = "<h3 style='color:#b45309;margin-top:0;'>Cover / Title Page</h3>";
    container.appendChild(coverSection);
    for (let elem of coverElements) createFormatControl(elem, defaults, fontOptions, coverSection);
  }
  for (let elem of regularElements) createFormatControl(elem, defaults, fontOptions, container);
  if (listElements.length > 0) {
    let listSection = document.createElement("div");
    listSection.className = "list-section";
    listSection.innerHTML = "<h3>List Item Formatting</h3>";
    container.appendChild(listSection);
    for (let elem of listElements) createFormatControl(elem, defaults, fontOptions, listSection);
  }
  if (tocElements.length > 0) {
    let tocSection = document.createElement("div");
    tocSection.className = "toc-section";
    tocSection.innerHTML = "<h3>Table of Contents Formatting</h3>";
    container.appendChild(tocSection);
    for (let elem of tocElements) createFormatControl(elem, defaults, fontOptions, tocSection);
  }
  for (let elem of sorted) updatePreview(elem);
}
function createFormatControl(elem, defaults, fontOptions, container) {
  let def = defaults[elem] || {font: "Calibri", size: 12};
  let section = document.createElement("div");
  section.style.marginBottom = "15px";
  let fontOptionsHTML = fontOptions.map(font =>
    `<option ${def.font === font ? "selected" : ""}>${font}</option>`
  ).join('');
  section.innerHTML = `
    <h3>${formatElementName(elem)} Style</h3>
    <div class="preview-container">
      <div style="flex:1">
        <label>Font</label>
        <select id="${elem}_font" onchange="updatePreview('${elem}')">${fontOptionsHTML}</select>
        <label>Size</label>
        <input type="number" id="${elem}_size" value="${def.size}" onchange="updatePreview('${elem}')">
      </div>
      <div class="preview-box">
        <div class="preview-label">Preview</div>
        <div id="${elem}_preview">${escapeHtml((analysisData && analysisData.sample_texts && analysisData.sample_texts[elem]) || "Sample text")}</div>
      </div>
    </div>
  `;
  container.appendChild(section);
}
function updatePreview(elem) {
  let preview = document.getElementById(elem + "_preview");
  if (!preview) return;
  let fontEl = document.getElementById(elem + "_font");
  let sizeEl = document.getElementById(elem + "_size");
  if (!fontEl || !sizeEl) return;
  let font = fontEl.value;
  let size = sizeEl.value;
  preview.style.fontFamily = font;
  preview.style.fontSize = size + "px";
  let boldLists = document.getElementById("bold_lists") ? document.getElementById("bold_lists").checked : false;
  if (
    ((elem === "TITLE" || elem === "COVER_PAGE") && document.getElementById("bold_titles").checked) ||
    (elem.startsWith("HEADING_") && document.getElementById("bold_headings").checked) ||
    (elem === "TOC_TITLE" && document.getElementById("bold_toc").checked) ||
    ((elem === "BULLET_ITEM" || elem === "NUMBERED_ITEM" || elem === "LIST_ITEM") && boldLists)
  ) preview.style.fontWeight = "bold";
  else preview.style.fontWeight = "normal";
  if (document.getElementById("highlight").checked) preview.style.background = "#fff59d";
  else preview.style.background = "transparent";
}
async function formatDoc(){
  let file = document.getElementById("file").files[0];
  if(!file || !analysisData){ alert("Please analyze document first"); return; }
  let config = {
    bold_titles: document.getElementById("bold_titles").checked,
    bold_headings: document.getElementById("bold_headings").checked,
    bold_toc: document.getElementById("bold_toc").checked,
    bold_lists: document.getElementById("bold_lists") ? document.getElementById("bold_lists").checked : false,
    highlight: document.getElementById("highlight").checked,
    preserve_indentation: document.getElementById("preserve_indentation").checked
  };
  for(let elem of analysisData.detected_elements) {
    let fontEl = document.getElementById(elem + "_font");
    let sizeEl = document.getElementById(elem + "_size");
    if (fontEl) config[elem.toLowerCase() + "_font"] = fontEl.value;
    if (sizeEl) config[elem.toLowerCase() + "_size"] = sizeEl.value;
  }
  let fd = new FormData();
  fd.append("file", file);
  fd.append("config", JSON.stringify(config));
  let res = await fetch("/format", {method:"POST", body:fd});
  if (!res.ok) {
    let t = await res.text();
    alert("Format failed: " + t);
    return;
  }
  let blob = await res.blob();
  let a = document.createElement("a");
  a.href = URL.createObjectURL(blob);
  a.download = file.name.replace(".docx","_formatted.docx");
  a.click();
}
</script>
</body>
</html>
"""


# ================== UTILITIES ==================
def get_paragraph_indentation(para) -> float:
    try:
        if para.paragraph_format.left_indent:
            return para.paragraph_format.left_indent.inches
        return 0.0
    except Exception:
        return 0.0


def get_list_level(para) -> int:
    """Return 0-based list nesting level from w:numPr/w:ilvl, or 0 if not found."""
    try:
        ilvl_nodes = para._element.xpath(".//w:numPr/w:ilvl")
        if ilvl_nodes:
            return int(ilvl_nodes[0].get(qn("w:val"), 0))
    except Exception:
        pass
    return 0


def is_bullet_list(para) -> bool:
    """
    Returns True if paragraph is a bullet (unordered) list item.
    Checks numFmt for bullet/wingdings style OR style name containing 'List Bullet'.
    """
    try:
        # Check style name first
        style_name = para.style.name if para.style else ""
        if "List Bullet" in style_name:
            return True

        # Check numPr exists
        numPr_list = para._element.xpath(".//w:numPr")
        if not numPr_list:
            return False

        # Try to find numId
        numId_nodes = para._element.xpath(".//w:numPr/w:numId")
        if not numId_nodes:
            return False

        # If we can access the document's numbering part, check numFmt
        doc = para._p.getparent()
        while doc is not None and doc.tag != qn("w:body"):
            doc = doc.getparent()

        if doc is not None:
            root = doc.getparent()
            # Access numbering part via the document part
            try:
                numbering_part = para.part.numbering_part
                if numbering_part is None:
                    # Has numPr but no numbering part = treat as bullet
                    return True

                numId_val = numId_nodes[0].get(qn("w:val"), "0")
                ilvl_nodes = para._element.xpath(".//w:numPr/w:ilvl")
                ilvl_val = ilvl_nodes[0].get(qn("w:val"), "0") if ilvl_nodes else "0"

                # Look up the abstract num
                num_elements = numbering_part._element.xpath(
                    f".//w:num[@w:numId='{numId_val}']",
                    namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                )
                if not num_elements:
                    return False

                abstractNumId_nodes = num_elements[0].xpath(
                    ".//w:abstractNumId",
                    namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                )
                if not abstractNumId_nodes:
                    return False

                abstract_id = abstractNumId_nodes[0].get(qn("w:val"), "0")
                abstract_num = numbering_part._element.xpath(
                    f".//w:abstractNum[@w:abstractNumId='{abstract_id}']",
                    namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                )
                if not abstract_num:
                    return False

                lvl_elements = abstract_num[0].xpath(
                    f".//w:lvl[@w:ilvl='{ilvl_val}']/w:numFmt",
                    namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                )
                if lvl_elements:
                    fmt = lvl_elements[0].get(qn("w:val"), "")
                    # bullet formats
                    if fmt in ("bullet", "none"):
                        return True
                    # numbered formats
                    if fmt in ("decimal", "lowerLetter", "upperLetter", "lowerRoman", "upperRoman",
                               "ordinal", "cardinalText", "decimalZero"):
                        return False
                    # unknown format with numPr = treat as bullet
                    return True
            except Exception:
                # Has numPr but couldn't determine type = treat generically
                return False

    except Exception:
        pass
    return False


def is_numbered_list(para) -> bool:
    """
    Returns True if paragraph is a numbered (ordered) list item.
    """
    try:
        style_name = para.style.name if para.style else ""
        if "List Number" in style_name:
            return True

        numPr_list = para._element.xpath(".//w:numPr")
        if not numPr_list:
            return False

        numId_nodes = para._element.xpath(".//w:numPr/w:numId")
        if not numId_nodes:
            return False

        try:
            numbering_part = para.part.numbering_part
            if numbering_part is None:
                return False

            numId_val = numId_nodes[0].get(qn("w:val"), "0")
            ilvl_nodes = para._element.xpath(".//w:numPr/w:ilvl")
            ilvl_val = ilvl_nodes[0].get(qn("w:val"), "0") if ilvl_nodes else "0"

            num_elements = numbering_part._element.xpath(
                f".//w:num[@w:numId='{numId_val}']",
                namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            )
            if not num_elements:
                return False

            abstractNumId_nodes = num_elements[0].xpath(
                ".//w:abstractNumId",
                namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            )
            if not abstractNumId_nodes:
                return False

            abstract_id = abstractNumId_nodes[0].get(qn("w:val"), "0")
            abstract_num = numbering_part._element.xpath(
                f".//w:abstractNum[@w:abstractNumId='{abstract_id}']",
                namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            )
            if not abstract_num:
                return False

            lvl_elements = abstract_num[0].xpath(
                f".//w:lvl[@w:ilvl='{ilvl_val}']/w:numFmt",
                namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            )
            if lvl_elements:
                fmt = lvl_elements[0].get(qn("w:val"), "")
                if fmt in ("decimal", "lowerLetter", "upperLetter", "lowerRoman", "upperRoman",
                           "ordinal", "cardinalText", "decimalZero"):
                    return True
        except Exception:
            pass

    except Exception:
        pass
    return False


def is_any_list_paragraph(para) -> bool:
    """
    True for ANY list paragraph (bullet OR numbered).
    Also catches 'List Paragraph' style and w:numPr presence.
    """
    try:
        if para._element.xpath(".//w:numPr"):
            return True
    except Exception:
        pass

    try:
        style_name = para.style.name if para.style is not None else ""
        if "List" in style_name:
            return True
    except Exception:
        pass

    return False


def get_list_type(para) -> str:
    """
    Returns 'BULLET_ITEM', 'NUMBERED_ITEM', or 'LIST_ITEM' (unknown).
    """
    if is_bullet_list(para):
        return "BULLET_ITEM"
    if is_numbered_list(para):
        return "NUMBERED_ITEM"
    # Has numPr but couldn't determine type
    if is_any_list_paragraph(para):
        return "LIST_ITEM"
    return None


# --- detect manual numbered headings like "1.1.1 Data Mapping" ---
NUMBERED_HEADING_RE = re.compile(r"^\s*(\d+(?:\.\d+)+)\s+(.*\S.*)$")

# --- detect manual numbered list items like "1. Item" or "a) Item" ---
MANUAL_NUMBERED_LIST_RE = re.compile(r"^\s*(\d+[\.\)]\s+|[a-zA-Z][\.\)]\s+|[ivxIVX]+[\.\)]\s+)(.*\S.*)$")


def detect_numbered_heading(text: str):
    """
    Detect headings like "1.1.1 Data Mapping".
    Returns (level, number, title) or (None, None, None).
    """
    if not text:
        return None, None, None
    m = NUMBERED_HEADING_RE.match(text)
    if not m:
        return None, None, None
    num = m.group(1)
    title = m.group(2).strip()
    # Must have at least 2 levels (e.g., "1.1") to be a heading
    if len(num.split(".")) < 2:
        return None, None, None
    level = min(6, len(num.split(".")))
    return level, num, title


def clear_paragraph_runs(para):
    """Remove all runs from a paragraph so we can rewrite it cleanly."""
    p = para._p
    for r in list(para.runs):
        try:
            p.remove(r._r)
        except Exception:
            pass


def has_page_break_before(para) -> bool:
    """Check if paragraph has a page break before it (either in pPr or as a run element)."""
    try:
        # Check w:pageBreakBefore in paragraph properties
        pPr = para._element.find(qn("w:pPr"))
        if pPr is not None:
            pb = pPr.find(qn("w:pageBreakBefore"))
            if pb is not None:
                val = pb.get(qn("w:val"), "true")
                if val.lower() not in ("false", "0", "off"):
                    return True

        # Check for explicit page break run (w:br w:type="page")
        for br in para._element.xpath(".//w:br"):
            if br.get(qn("w:type"), "") == "page":
                return True

        # Check if the paragraph's XML contains lastRenderedPageBreak
        if "lastRenderedPageBreak" in para._element.xml:
            return True
    except Exception:
        pass
    return False


def detect_cover_page(doc) -> int:
    """
    Detect the last paragraph index that belongs to the cover/title page.

    Strategy (in order of confidence):
    1. If there is an explicit page break run (w:br type=page) in a paragraph,
       everything UP TO AND INCLUDING that paragraph is the cover page.
    2. If a paragraph has w:pageBreakBefore in its pPr, everything before it
       is the cover page.
    3. If we find the first 'Heading 1' style paragraph (real content starts),
       everything before it is cover page — but only if it falls within the
       first 30 paragraphs (to avoid false positives in heading-heavy docs).
    4. If none of the above fire, fall back to checking the first page via
       font size / style heuristics across the first 25 non-empty paragraphs.

    Returns the index of the LAST cover-page paragraph, or -1 if no cover page detected.
    """
    paragraphs = doc.paragraphs
    n = len(paragraphs)

    # --- Pass 1: look for an explicit page break (most reliable) ---
    for i, para in enumerate(paragraphs[:60]):
        # A page break *inside* a paragraph means everything up to this para
        # (inclusive) is still the cover page.
        for run in para.runs:
            for br in run._element.xpath(".//w:br"):
                if br.get(qn("w:type"), "") == "page":
                    return i  # cover page ends at this paragraph

        # A page break *before* the NEXT paragraph means this para is the last
        # cover page para — but only check paragraphs after index 0.
        if i > 0 and has_page_break_before(para):
            return i - 1

    # --- Pass 2: first Heading 1 marks the start of body content ---
    for i, para in enumerate(paragraphs[:40]):
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading") and i > 0:
            return i - 1

    # --- Pass 3: heuristic — scan up to 25 non-empty paragraphs.
    # A cover page typically has: short lines, possibly centered, possibly
    # large font, and keywords like "prepared by/for", "version", "author",
    # "copyright", "confidential", dates.
    COVER_KEYWORDS = re.compile(
        r"\b(prepared\s+(by|for)|author|version|copyright|confidential|"
        r"restricted|january|february|march|april|may|june|july|august|"
        r"september|october|november|december|\d{4})\b",
        re.IGNORECASE,
    )
    non_empty = [(i, p) for i, p in enumerate(paragraphs) if (p.text or "").strip()]
    cover_candidate = -1
    for rank, (i, para) in enumerate(non_empty[:25]):
        text = (para.text or "").strip()
        is_short = len(text) < 200
        is_centered = para.alignment == WD_ALIGN_PARAGRAPH.CENTER
        has_large_font = False
        style_name = para.style.name if para.style else ""
        if para.runs:
            sz = para.runs[0].font.size
            if sz and sz >= Pt(13):
                has_large_font = True
        has_keyword = bool(COVER_KEYWORDS.search(text))
        is_title_style = style_name in ("Title", "Subtitle")
        # Looks like cover content if any indicator fires
        if is_title_style or is_centered or has_large_font or has_keyword or is_short:
            cover_candidate = i
        else:
            # First paragraph that looks like a body paragraph ends the cover
            if rank > 2:  # give at least 3 paras before giving up
                break

    return cover_candidate


def _get_style_base(para) -> str:
    """Return normalised style name without ' Char' suffix."""
    try:
        name = para.style.name if para.style else ""
        return name.split(" Char")[0].strip()
    except Exception:
        return ""


def _is_toc_entry_by_style(para) -> bool:
    """True if the paragraph has a Word built-in TOC style (TOC 1 … TOC 9)."""
    return bool(re.match(r"^TOC \d+$", _get_style_base(para)))


def _toc_level_from_style(para) -> int:
    """Return 1-based level from 'TOC N' style, or 0 if not a TOC style."""
    m = re.match(r"^TOC (\d+)$", _get_style_base(para))
    if m:
        return int(m.group(1))
    return 0


# Matches TOC entry patterns like:
#   "1      Introduction. 5"
#   "3.1.1         Properties of the Data Flow.. 7"
#   "Introduction\t3"
#   "Introduction .................. 3"
#   "Introduction...3"
_TOC_ENTRY_RE = re.compile(
    r"^\s*"
    r"(\d+(?:\.\d+)*)?"          # optional section number (1, 1.1, 3.1.2, etc.)
    r"\s*"
    r"(.+?)"                      # entry title (non-greedy)
    r"[\s\.…\-]{2,}"             # separator: 2+ spaces, dots, ellipsis, dashes
    r"(\d{1,4})"                  # page number (1-4 digits)
    r"\s*$",
    re.UNICODE,
)

# Also catch tab-separated: "Title\t12"
_TOC_TAB_RE = re.compile(r"^.+\t\d{1,4}\s*$")


def _is_toc_entry_by_heuristic(para) -> bool:
    """
    Detect TOC entries in documents using Normal style.
    Handles:
      - "1      Introduction. 5"          (spaces + dot + space + number)
      - "3.1.1  Data Mapping. 9"          (numbered heading style)
      - "Introduction .............. 3"   (dot leader)
      - "Introduction\t3"                 (tab separated)
      - PAGEREF field codes in XML
    """
    text = (para.text or "").strip()
    if not text:
        return False

    if _TOC_ENTRY_RE.match(text):
        return True

    if _TOC_TAB_RE.match(text):
        return True

    # Word field codes for TOC / PAGEREF
    xml = para._element.xml
    if "PAGEREF" in xml or ("w:instr" in xml and "TOC" in xml):
        return True

    return False


def _toc_level_from_heuristic(text: str) -> int:
    """
    Infer TOC nesting level from the section number prefix.
      "1 Introduction..."     → level 1
      "1.1 Purposes..."       → level 2
      "3.1.1 Properties..."   → level 3
    Falls back to 1 if no section number detected.
    """
    m = re.match(r"^\s*(\d+(?:\.\d+)*)\s+", text.strip())
    if m:
        return min(len(m.group(1).split(".")), 6)
    return 1


def detect_toc_section(doc):
    """
    Detect the start and end paragraph indices of the Table of Contents.

    Priority:
    1. Word built-in TOC styles (TOC 1…TOC 9) — most reliable.
    2. Heuristic: find "Table of Contents" heading then scan forward for
       entry-pattern lines (handles Normal-style TOCs).
    """
    paragraphs = doc.paragraphs

    # ── Strategy 1: style-based ───────────────────────────────────────────────
    first_toc_style_idx = -1
    last_toc_style_idx = -1
    for i, para in enumerate(paragraphs):
        if _is_toc_entry_by_style(para):
            if first_toc_style_idx == -1:
                first_toc_style_idx = i
            last_toc_style_idx = i
        elif first_toc_style_idx != -1:
            gap = i - last_toc_style_idx
            if gap > 5:
                break

    if first_toc_style_idx != -1:
        toc_start = first_toc_style_idx
        for j in range(first_toc_style_idx - 1, max(-1, first_toc_style_idx - 5), -1):
            text = (paragraphs[j].text or "").strip().lower()
            style_b = _get_style_base(paragraphs[j])
            if style_b == "TOC Heading" or "table of contents" in text or text == "contents":
                toc_start = j
                break
            if text:
                break
        return toc_start, last_toc_style_idx, True

    # ── Strategy 2: heuristic (Normal-style TOC) ──────────────────────────────
    TOC_TITLE_RE = re.compile(r"^\s*(table\s+of\s+contents|contents)\s*$", re.IGNORECASE)
    toc_start = -1

    # Search up to paragraph 100 — cover page can be long
    search_limit = min(len(paragraphs), 100)
    for i, para in enumerate(paragraphs[:search_limit]):
        text = (para.text or "").strip()
        style_b = _get_style_base(para)
        if style_b == "TOC Heading" or TOC_TITLE_RE.match(text):
            toc_start = i
            break
        if "TOC" in para._element.xml and "w:fldChar" in para._element.xml:
            toc_start = i
            break

    if toc_start == -1:
        return -1, -1, False

    # Walk forward — accept any line matching TOC entry pattern
    toc_end = toc_start
    consecutive_non_toc = 0
    for i in range(toc_start + 1, min(len(paragraphs), toc_start + 300)):
        para = paragraphs[i]
        text = (para.text or "").strip()
        if not text:
            continue
        if _is_toc_entry_by_heuristic(para) or _is_toc_entry_by_style(para):
            toc_end = i
            consecutive_non_toc = 0
        else:
            consecutive_non_toc += 1
            if consecutive_non_toc >= 3:
                break

    return toc_start, toc_end, True


def analyze_toc_hierarchy(doc, toc_start, toc_end):
    """
    Return (level_map, num_levels).
    For style-based TOCs: maps style level → level number.
    For heuristic TOCs: maps section-number depth to level.
    """
    if toc_start == -1 or toc_end == -1:
        return {}, 0

    # Style-based (most reliable)
    style_levels = set()
    for i in range(toc_start, toc_end + 1):
        if i >= len(doc.paragraphs):
            break
        lvl = _toc_level_from_style(doc.paragraphs[i])
        if lvl:
            style_levels.add(lvl)

    if style_levels:
        return {lvl: lvl for lvl in style_levels}, len(style_levels)

    # Heuristic: infer levels from section number depth
    level_counts = {}
    for i in range(toc_start + 1, toc_end + 1):
        if i >= len(doc.paragraphs):
            break
        para = doc.paragraphs[i]
        text = (para.text or "").strip()
        if not text:
            continue
        if _is_toc_entry_by_heuristic(para):
            lvl = _toc_level_from_heuristic(text)
            level_counts[lvl] = level_counts.get(lvl, 0) + 1

    return {lvl: lvl for lvl in sorted(level_counts)}, len(level_counts)


def detect_hierarchy_by_indentation(doc, exclude_range=None):
    indent_levels = {}
    for i, para in enumerate(doc.paragraphs):
        if exclude_range and exclude_range[0] <= i <= exclude_range[1]:
            continue
        if not (para.text or "").strip():
            continue
        # Skip list paragraphs for indentation-based heading inference
        if is_any_list_paragraph(para):
            continue

        indent = get_paragraph_indentation(para)
        if indent > 3:
            continue

        if indent > 0:
            rounded_indent = round(indent * 4) / 4
            indent_levels[rounded_indent] = indent_levels.get(rounded_indent, 0) + 1

    sorted_indents = sorted(indent_levels.keys())
    indent_to_level = {}
    for i, indent in enumerate(sorted_indents[:6]):
        indent_to_level[indent] = i + 1

    return indent_to_level, len(sorted_indents)


def analyze_document_structure(docx_path):
    doc = Document(docx_path)

    elements = []
    element_counts = {}
    detected_types = set()
    sample_texts = {}

    list_items_by_heading = {}

    # --- Cover page detection (must happen before TOC detection) ---
    cover_end = detect_cover_page(doc)  # last para idx of cover page, or -1

    toc_start, toc_end, has_toc = detect_toc_section(doc)
    toc_indent_to_level = {}
    toc_indent_levels = 0
    if has_toc:
        toc_indent_to_level, toc_indent_levels = analyze_toc_hierarchy(doc, toc_start, toc_end)

    indent_to_level, total_indent_levels = detect_hierarchy_by_indentation(
        doc, exclude_range=(toc_start, toc_end) if has_toc else None
    )
    has_indented_hierarchy = len(indent_to_level) > 0

    # Track current heading hierarchy stack
    heading_stack = {}

    def current_parent_heading():
        if not heading_stack:
            return None
        return heading_stack[max(heading_stack.keys())]

    def set_heading(level: int):
        heading_stack[level] = f"HEADING_{level}"
        for k in list(heading_stack.keys()):
            if k > level:
                del heading_stack[k]

    # Pre-compute the first non-empty paragraph index (for cover page title detection)
    first_content_idx = next(
        (i for i, p in enumerate(doc.paragraphs) if (p.text or "").strip()), 0
    )

    for idx, para in enumerate(doc.paragraphs):
        if not (para.text or "").strip():
            continue

        style_name = para.style.name if para.style is not None else ""
        # Normalise: sometimes styles have suffixes like "Title Char"
        style_base = style_name.split(" Char")[0].strip()

        ptype = None
        parent_heading = None
        in_toc_section = has_toc and toc_start <= idx <= toc_end
        in_cover_page = cover_end >= 0 and idx <= cover_end

        # ── STEP 0: Word built-in styles are authoritative — check these FIRST
        #    regardless of position or section range.
        if style_base == "Title":
            ptype = "TITLE"
        elif style_base == "TOC Heading":
            ptype = "TOC_TITLE"
        elif style_base in ("Subtitle", "Document Map"):
            ptype = "COVER_PAGE"
        # Cover-metadata styles that Oracle/corporate templates commonly use
        elif style_base in ("Author", "Date", "Company", "Abstract",
                             "Document Label", "Revision", "Version"):
            ptype = "COVER_PAGE"
        # Word built-in TOC entry styles — most reliable TOC signal
        elif re.match(r"^TOC \d+$", style_base):
            try:
                level = int(style_base.split()[-1])
            except Exception:
                level = 1
            ptype = f"TOC_HEADING_{level}"

        # ── STEP 1: TOC section overrides remaining unclassified paragraphs
        if ptype is None and in_toc_section:
            text_lower = para.text.strip().lower()
            style_b = _get_style_base(para)

            if style_b == "TOC Heading" or re.match(r"^\s*(table\s+of\s+contents|contents)\s*$", text_lower):
                ptype = "TOC_TITLE"
            else:
                # 1st choice: Word built-in TOC style
                lvl = _toc_level_from_style(para)
                if lvl:
                    ptype = f"TOC_HEADING_{lvl}"
                else:
                    # 2nd choice: infer level from section-number depth  e.g. "3.1.1" → 3
                    text_raw = (para.text or "").strip()
                    lvl = _toc_level_from_heuristic(text_raw)
                    ptype = f"TOC_HEADING_{lvl}"

        # ── STEP 2: Cover page range — Normal-style paragraphs on the cover page
        #    (Title/Subtitle already caught above; this catches "Prepared For:" etc.)
        if ptype is None and in_cover_page and not in_toc_section:
            if idx == first_content_idx:
                ptype = "TITLE"   # very first line of document → title
            else:
                ptype = "COVER_PAGE"

        # ── STEP 3: Body content classification
        if ptype is None:
            raw_text = (para.text or "").strip()

            # Heading styles
            if style_base.startswith("Heading"):
                try:
                    level = int(style_base.split()[-1])
                    level = max(1, min(6, level))
                    ptype = f"HEADING_{level}"
                except Exception:
                    ptype = "HEADING_1"

            # List detection (must come before numbered-heading inference)
            if not ptype:
                list_type = get_list_type(para)
                if list_type:
                    ptype = list_type

            # Manual numbered headings like "1.1.1 Data Mapping"
            if not ptype:
                lvl, num, title = detect_numbered_heading(raw_text)
                if lvl:
                    ptype = f"HEADING_{lvl}"

            # Indentation-based hierarchy inference
            if not ptype:
                indent = get_paragraph_indentation(para)
                if indent > 0 and indent_to_level:
                    rounded_indent = round(indent * 4) / 4
                    if rounded_indent in indent_to_level:
                        ptype = f"HEADING_{indent_to_level[rounded_indent]}"

            # Final fallback
            if not ptype:
                ptype = "PARAGRAPH"

        # ── Update heading stack and list tracking
        if ptype and ptype.startswith("HEADING_"):
            try:
                lvl = int(ptype.split("_")[-1])
            except Exception:
                lvl = 1
            set_heading(lvl)

        if ptype in ("BULLET_ITEM", "NUMBERED_ITEM", "LIST_ITEM"):
            parent_heading = current_parent_heading()
            key = parent_heading if parent_heading else "NO_HEADING"
            list_items_by_heading[key] = list_items_by_heading.get(key, 0) + 1

        # Store element info
        item = {
            "type": ptype,
            "para_idx": idx,
            "indent": get_paragraph_indentation(para),
            "list_level": get_list_level(para) if ptype in ("BULLET_ITEM", "NUMBERED_ITEM", "LIST_ITEM") else 0
        }
        if parent_heading:
            item["parent_heading"] = parent_heading

        # Store manual heading pieces if applicable
        if not in_toc_section and not in_cover_page and ptype and ptype.startswith("HEADING_"):
            lvl2, num2, title2 = detect_numbered_heading((para.text or "").strip())
            if num2 and title2:
                item["heading_number"] = num2
                item["heading_title"] = title2

        elements.append(item)

        element_counts[ptype] = element_counts.get(ptype, 0) + 1
        detected_types.add(ptype)

        if ptype not in sample_texts:
            t = (para.text or "").strip()
            if t:
                sample_texts[ptype] = t[:250]

    for i, table in enumerate(doc.tables):
        elements.append({"type": "TABLE", "table_idx": i})
        element_counts["TABLE"] = element_counts.get("TABLE", 0) + 1
        detected_types.add("TABLE")
        if "TABLE" not in sample_texts:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        sample_texts["TABLE"] = cell.text.strip()[:250]
                        break

    def sort_key(x):
        if x == "TITLE": return 0
        if x == "COVER_PAGE": return 1
        if x == "TOC_TITLE": return 2
        if x.startswith("TOC_HEADING_"): return 3
        if x.startswith("HEADING_"): return 4
        if x in ("BULLET_ITEM", "NUMBERED_ITEM", "LIST_ITEM"): return 5
        if x == "PARAGRAPH": return 6
        return 7

    detected_elements = sorted(detected_types, key=sort_key)

    return {
        "elements": elements,
        "detected_elements": detected_elements,
        "element_counts": element_counts,
        "sample_texts": sample_texts,
        "has_toc": has_toc,
        "toc_indent_levels": toc_indent_levels,
        "has_indented_hierarchy": has_indented_hierarchy,
        "indent_levels": total_indent_levels,
        "list_items_by_heading": list_items_by_heading,
        "cover_page_end": cover_end,
    }


# ---------- XML helpers ----------
def get_or_add_tblPr(ct_tbl):
    tblPr = ct_tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        ct_tbl.insert(0, tblPr)
    return tblPr


def add_border_to_run_images(run, border_pt=0.25):
    drawings = run._element.xpath(".//w:drawing")
    if not drawings:
        return
    for drawing in drawings:
        containers = drawing.xpath(".//wp:inline | .//wp:anchor")
        for container in containers:
            spPr_list = container.xpath(".//pic:spPr")
            if not spPr_list:
                continue
            spPr = spPr_list[0]
            for ln in spPr.xpath(".//a:ln"):
                spPr.remove(ln)

            ln = OxmlElement("a:ln")
            ln.set("w", str(int(border_pt * 12700)))

            solidFill = OxmlElement("a:solidFill")
            srgbClr = OxmlElement("a:srgbClr")
            srgbClr.set("val", "000000")
            solidFill.append(srgbClr)
            ln.append(solidFill)

            prstDash = OxmlElement("a:prstDash")
            prstDash.set("val", "solid")
            ln.append(prstDash)

            spPr.append(ln)


def wrap_run_image_in_table(doc, para, run, border_pt=0.25):
    drawings = run._element.xpath(".//w:drawing")
    if not drawings:
        return

    drawing = drawings[0]
    table = doc.add_table(rows=1, cols=1)
    table.allow_autofit = True
    para._p.addnext(table._tbl)

    cell = table.cell(0, 0)
    cell_para = cell.paragraphs[0]
    new_run = cell_para.add_run()
    new_run._element.append(copy.deepcopy(drawing))

    try:
        para._p.remove(run._element)
    except Exception:
        pass

    tbl_pr = get_or_add_tblPr(table._tbl)
    tbl_borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(int(border_pt * 8)))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        tbl_borders.append(border)

    for old in tbl_pr.xpath("./w:tblBorders"):
        tbl_pr.remove(old)
    tbl_pr.append(tbl_borders)


def set_run_font_everywhere(run, font_name, font_size_pt=None, bold=None):
    """
    Apply font name (and optionally size + bold) at run-level rPr.
    Explicitly overrides theme fonts and clears conflicting attributes.
    """
    run.font.name = font_name
    if font_size_pt is not None:
        run.font.size = Pt(font_size_pt)
    if bold is not None:
        run.font.bold = bold

    rPr = run._element.get_or_add_rPr()

    # Font family — all four slots + remove theme overrides
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    for attr in (qn("w:asciiTheme"), qn("w:hAnsiTheme"), qn("w:cstheme")):
        if rFonts.get(attr) is not None:
            del rFonts.attrib[attr]

    # Size
    if font_size_pt is not None:
        half_pts = str(int(font_size_pt * 2))
        for tag_name in ("w:sz", "w:szCs"):
            el = rPr.find(qn(tag_name))
            if el is None:
                el = OxmlElement(tag_name)
                rPr.append(el)
            el.set(qn("w:val"), half_pts)

    # Bold — use explicit val="0" to override style-level bold
    if bold is not None:
        for tag_name in ("w:b", "w:bCs"):
            el = rPr.find(qn(tag_name))
            if el is None:
                el = OxmlElement(tag_name)
                rPr.append(el)
            el.set(qn("w:val"), "1" if bold else "0")


def apply_para_direct_format(para, font_name, font_size_pt, bold, highlight_color_name=None):
    """
    Write formatting into the paragraph's pPr/rPr block.
    This sits between the named paragraph style and the individual runs in Word's
    style cascade, ensuring our values override the style definition.
    """
    pPr = para._element.get_or_add_pPr()
    rPr = pPr.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        pPr.append(rPr)

    # Font
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    for attr in (qn("w:asciiTheme"), qn("w:hAnsiTheme"), qn("w:cstheme")):
        if rFonts.get(attr) is not None:
            del rFonts.attrib[attr]

    # Size
    half_pts = str(int(font_size_pt * 2))
    for tag_name in ("w:sz", "w:szCs"):
        el = rPr.find(qn(tag_name))
        if el is None:
            el = OxmlElement(tag_name)
            rPr.append(el)
        el.set(qn("w:val"), half_pts)

    # Bold
    for tag_name in ("w:b", "w:bCs"):
        el = rPr.find(qn(tag_name))
        if el is None:
            el = OxmlElement(tag_name)
            rPr.append(el)
        el.set(qn("w:val"), "1" if bold else "0")

    # Highlight
    if highlight_color_name:
        hl = rPr.find(qn("w:highlight"))
        if hl is None:
            hl = OxmlElement("w:highlight")
            rPr.append(hl)
        hl.set(qn("w:val"), highlight_color_name)



def _get_all_runs(para):
    """
    Return ALL w:r elements inside a paragraph, including those nested inside
    w:hyperlink, w:ins, w:del, w:sdtContent etc. that para.runs misses.
    Deliberately excludes runs inside nested w:tbl elements (tables within cells).
    python-docx pre-registers all namespaces so bare 'w:tag' xpath works directly.
    """
    all_r = para._element.xpath(".//w:r")
    nested_tbl_runs = set()
    for tbl in para._element.xpath(".//w:tbl"):
        for r in tbl.xpath(".//w:r"):
            nested_tbl_runs.add(id(r))
    return [r for r in all_r if id(r) not in nested_tbl_runs]


def _patch_doc_style(doc, style_name, font_name, font_size_pt, bold):
    """
    Patch the named paragraph style's rPr inside styles.xml so that even
    paragraphs with ZERO runs pick up the right formatting, and so the style
    definition doesn't fight our run-level overrides.
    """
    try:
        style = doc.styles[style_name]
    except (KeyError, Exception):
        return

    # Paragraph style element rPr (inside w:style > w:rPr, NOT w:pPr > w:rPr)
    style_el = style.element
    rPr = style_el.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        style_el.append(rPr)

    # Font
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    for attr in (qn("w:asciiTheme"), qn("w:hAnsiTheme"), qn("w:cstheme")):
        if rFonts.get(attr) is not None:
            del rFonts.attrib[attr]

    # Size
    half_pts = str(int(font_size_pt * 2))
    for tag in ("w:sz", "w:szCs"):
        el = rPr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            rPr.append(el)
        el.set(qn("w:val"), half_pts)

    # Bold
    for tag in ("w:b", "w:bCs"):
        el = rPr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            rPr.append(el)
        el.set(qn("w:val"), "1" if bold else "0")


def apply_numbering_runprops(para, font_name, font_size, bold=False, highlight_val=None):
    """
    Apply formatting to the bullet/number symbol itself (stored in w:pPr/w:rPr).
    This ensures the list marker matches the run formatting.
    """
    if not para._element.xpath(".//w:numPr"):
        return

    pPr = para._element.get_or_add_pPr()
    rPr = pPr.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        pPr.append(rPr)

    # Font
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)

    # Size
    half_pts = str(int(font_size) * 2)
    for tag in ("w:sz", "w:szCs"):
        el = rPr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            rPr.append(el)
        el.set(qn("w:val"), half_pts)

    # Bold
    b_elem = rPr.find(qn("w:b"))
    if bold:
        if b_elem is None:
            b_elem = OxmlElement("w:b")
            rPr.append(b_elem)
    else:
        if b_elem is not None:
            rPr.remove(b_elem)

    # Highlight
    if highlight_val:
        highlight = rPr.find(qn("w:highlight"))
        if highlight is None:
            highlight = OxmlElement("w:highlight")
            rPr.append(highlight)
        highlight.set(qn("w:val"), highlight_val)


def get_config_for_type(ptype, config):
    """
    Get font and size from config for a given element type.
    Falls back gracefully for list types and cover page.
    """
    font_key = ptype.lower() + "_font"
    size_key = ptype.lower() + "_size"

    font_name = config.get(font_key)
    font_size_raw = config.get(size_key)

    # Fallback chains
    if not font_name:
        if ptype == "COVER_PAGE":
            font_name = config.get("title_font") or config.get("paragraph_font") or "Calibri"
        elif ptype in ("BULLET_ITEM", "NUMBERED_ITEM"):
            font_name = config.get("list_item_font") or config.get("paragraph_font") or "Calibri"
        else:
            font_name = config.get("paragraph_font") or "Calibri"

    if not font_size_raw:
        if ptype == "COVER_PAGE":
            font_size_raw = config.get("title_size") or 12
        elif ptype in ("BULLET_ITEM", "NUMBERED_ITEM"):
            font_size_raw = config.get("list_item_size") or config.get("paragraph_size") or 12
        else:
            font_size_raw = config.get("paragraph_size") or 12

    try:
        font_size = int(font_size_raw)
    except Exception:
        font_size = 12

    return font_name, font_size


def format_docx(input_path, elements, output_path, config):
    doc = Document(input_path)

    para_map = {
        e["para_idx"]: e
        for e in elements
        if "para_idx" in e
    }

    color_map = {
        WD_COLOR_INDEX.YELLOW: "yellow",
        WD_COLOR_INDEX.BRIGHT_GREEN: "green",
        WD_COLOR_INDEX.TURQUOISE: "cyan",
        WD_COLOR_INDEX.PINK: "magenta",
        WD_COLOR_INDEX.VIOLET: "magenta",
        WD_COLOR_INDEX.TEAL: "cyan",
        WD_COLOR_INDEX.DARK_BLUE: "darkBlue",
        WD_COLOR_INDEX.GRAY_25: "lightGray",
        WD_COLOR_INDEX.GRAY_50: "darkGray",
        WD_COLOR_INDEX.DARK_YELLOW: "darkYellow",
    }

    # ── Pre-pass: patch the document's named styles so the style definition
    #    itself matches the user's chosen font/size. This ensures paragraphs
    #    with NO runs (or runs with no explicit formatting) look correct too.
    STYLE_TYPE_MAP = {
        "TITLE":        ["Title"],
        "COVER_PAGE":   ["Subtitle"],
        "TOC_TITLE":    ["TOC Heading"],
        "TOC_HEADING_1":["TOC 1"],
        "TOC_HEADING_2":["TOC 2"],
        "TOC_HEADING_3":["TOC 3"],
        "TOC_HEADING_4":["TOC 4"],
        "HEADING_1":    ["Heading 1"],
        "HEADING_2":    ["Heading 2"],
        "HEADING_3":    ["Heading 3"],
        "HEADING_4":    ["Heading 4"],
        "HEADING_5":    ["Heading 5"],
        "HEADING_6":    ["Heading 6"],
        "PARAGRAPH":    ["Normal"],
        "BULLET_ITEM":  ["List Bullet", "List Paragraph"],
        "NUMBERED_ITEM":["List Number", "List Paragraph"],
    }
    patched_styles = set()
    for ptype, style_names in STYLE_TYPE_MAP.items():
        fn, fs = get_config_for_type(ptype, config)
        is_bold = (
            (ptype.startswith("HEADING_") and config.get("bold_headings", True)) or
            (ptype in ("TITLE", "COVER_PAGE") and config.get("bold_titles", True)) or
            (ptype == "TOC_TITLE" and config.get("bold_toc", True)) or
            (ptype in ("BULLET_ITEM", "NUMBERED_ITEM", "LIST_ITEM") and config.get("bold_lists", False))
        )
        for sname in style_names:
            if sname not in patched_styles:
                _patch_doc_style(doc, sname, fn, fs, is_bold)
                patched_styles.add(sname)

    # ── Main per-paragraph loop ────────────────────────────────────────────────
    for idx, para in enumerate(doc.paragraphs):
        elem_info = para_map.get(idx)
        if not elem_info:
            continue

        ptype = elem_info["type"]
        original_indent = elem_info.get("indent", 0.0)

        font_name, font_size = get_config_for_type(ptype, config)

        is_heading = ptype.startswith("HEADING_")
        is_title = ptype == "TITLE"
        is_toc_title = ptype == "TOC_TITLE"
        is_cover = ptype == "COVER_PAGE"
        is_list = ptype in ("BULLET_ITEM", "NUMBERED_ITEM", "LIST_ITEM")

        bold_run = (
            (is_heading and config.get("bold_headings", True)) or
            (is_title and config.get("bold_titles", True)) or
            (is_cover and config.get("bold_titles", True)) or
            (is_toc_title and config.get("bold_toc", True)) or
            (is_list and config.get("bold_lists", False))
        )

        highlight_color_name = None
        if config.get("highlight", False):
            hc = HIGHLIGHT_COLORS.get(ptype, WD_COLOR_INDEX.GRAY_25)
            highlight_color_name = color_map.get(hc)

        # ── 1. Paragraph-level pPr/rPr override ───────────────────────────────
        # Sits between style and run in cascade. Covers paragraphs with no runs.
        apply_para_direct_format(para, font_name, font_size, bold_run, highlight_color_name)

        # ── 2. Rewrite manual numbered headings ────────────────────────────────
        if is_heading:
            text = (para.text or "").strip()
            lvl, num, title = detect_numbered_heading(text)
            if num and title:
                clear_paragraph_runs(para)
                para.add_run(f"{num}\t")
                para.add_run(title)

        # ── 3. Run-level override — ALL runs including inside w:hyperlink ──────
        for r_el in _get_all_runs(para):
            # Skip runs that contain drawings (images)
            if r_el.xpath(".//w:drawing"):
                continue

            rPr = r_el.find(qn("w:rPr"))
            if rPr is None:
                rPr = OxmlElement("w:rPr")
                r_el.insert(0, rPr)

            # Font
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.insert(0, rFonts)
            rFonts.set(qn("w:ascii"), font_name)
            rFonts.set(qn("w:hAnsi"), font_name)
            rFonts.set(qn("w:cs"), font_name)
            rFonts.set(qn("w:eastAsia"), font_name)
            for attr in (qn("w:asciiTheme"), qn("w:hAnsiTheme"), qn("w:cstheme")):
                if rFonts.get(attr) is not None:
                    del rFonts.attrib[attr]

            # Size
            half_pts = str(int(font_size * 2))
            for tag in ("w:sz", "w:szCs"):
                el = rPr.find(qn(tag))
                if el is None:
                    el = OxmlElement(tag)
                    rPr.append(el)
                el.set(qn("w:val"), half_pts)

            # Bold
            for tag in ("w:b", "w:bCs"):
                el = rPr.find(qn(tag))
                if el is None:
                    el = OxmlElement(tag)
                    rPr.append(el)
                el.set(qn("w:val"), "1" if bold_run else "0")

            # Highlight
            if highlight_color_name:
                hl = rPr.find(qn("w:highlight"))
                if hl is None:
                    hl = OxmlElement("w:highlight")
                    rPr.append(hl)
                hl.set(qn("w:val"), highlight_color_name)

        # ── 4. List marker formatting ──────────────────────────────────────────
        if is_list:
            apply_numbering_runprops(
                para,
                font_name=font_name,
                font_size=font_size,
                bold=bold_run,
                highlight_val=highlight_color_name,
            )

        # ── 5. Preserve indentation ────────────────────────────────────────────
        if config.get("preserve_indentation", True) and original_indent and original_indent > 0:
            para.paragraph_format.left_indent = Inches(original_indent)

    # ── Table handling ─────────────────────────────────────────────────────────
    for e in elements:
        if e["type"] == "TABLE" and e.get("table_idx", -1) < len(doc.tables):
            table = doc.tables[e["table_idx"]]
            font_name, font_size = get_config_for_type("TABLE", config)
            half_pts = str(int(font_size * 2))
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        # Patch paragraph-level too
                        apply_para_direct_format(para, font_name, font_size, False,
                                                  "cyan" if config.get("highlight") else None)
                        for r_el in _get_all_runs(para):
                            if r_el.xpath(".//w:drawing"):
                                continue
                            rPr = r_el.find(qn("w:rPr"))
                            if rPr is None:
                                rPr = OxmlElement("w:rPr")
                                r_el.insert(0, rPr)
                            rFonts = rPr.find(qn("w:rFonts"))
                            if rFonts is None:
                                rFonts = OxmlElement("w:rFonts")
                                rPr.insert(0, rFonts)
                            rFonts.set(qn("w:ascii"), font_name)
                            rFonts.set(qn("w:hAnsi"), font_name)
                            rFonts.set(qn("w:cs"), font_name)
                            rFonts.set(qn("w:eastAsia"), font_name)
                            for tag in ("w:sz", "w:szCs"):
                                el = rPr.find(qn(tag))
                                if el is None:
                                    el = OxmlElement(tag)
                                    rPr.append(el)
                                el.set(qn("w:val"), half_pts)
                            if config.get("highlight", False):
                                hl = rPr.find(qn("w:highlight"))
                                if hl is None:
                                    hl = OxmlElement("w:highlight")
                                    rPr.append(hl)
                                hl.set(qn("w:val"), "cyan")

    doc.save(output_path)


# ================== ROUTES ==================
@app.route("/")
def index():
    return render_template_string(HTML_TEMPLATE)


@app.route("/analyze_structure", methods=["POST"])
def analyze_structure():
    file = request.files.get("file")
    if not file or not file.filename.lower().endswith(".docx"):
        return jsonify({"error": "Invalid file"}), 400

    original_name = secure_filename(file.filename)
    input_path = os.path.join(app.config["UPLOAD_FOLDER"], original_name)
    file.save(input_path)

    analysis = analyze_document_structure(input_path)
    return jsonify(analysis)


@app.route("/format", methods=["POST"])
def format_document():
    file = request.files.get("file")
    config_raw = request.form.get("config", "{}")

    try:
        config = json.loads(config_raw)
    except Exception:
        return jsonify({"error": "Invalid config JSON"}), 400

    if not file or not file.filename.lower().endswith(".docx"):
        return jsonify({"error": "Invalid file"}), 400

    original_name = secure_filename(file.filename)
    base, ext = os.path.splitext(original_name)

    input_path = os.path.join(app.config["UPLOAD_FOLDER"], original_name)
    output_filename = f"{base}_formatted{ext}"
    output_path = os.path.join(app.config["OUTPUT_FOLDER"], output_filename)

    file.save(input_path)

    analysis = analyze_document_structure(input_path)
    format_docx(input_path, analysis["elements"], output_path, config)

    return send_file(output_path, as_attachment=True, download_name=output_filename)


if __name__ == "__main__":
    app.run(debug=True, port=5000)
