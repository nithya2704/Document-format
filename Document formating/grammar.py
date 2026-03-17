import re as _re
import io
import os
import json
import uuid

from flask import (Blueprint, render_template, request,
                   jsonify, send_file, session, current_app)
from werkzeug.utils import secure_filename
from docx import Document
from textblob import TextBlob
import language_tool_python

grammar_bp = Blueprint(
    "grammar",
    __name__,
    template_folder="templates",
    url_prefix="/grammar"
)

OUTPUT = "outputs/grammar"
os.makedirs(OUTPUT, exist_ok=True)

tool = language_tool_python.LanguageTool('en-US')


# ── Session store helpers ─────────────────────────────────────────────────

def _get_path():
    """Return the current working-file path for this session."""
    sid = session.get("sid")
    if not sid:
        return None
    return current_app.config["GET_WORKING_PATH"](sid)


def _set_path(path: str):
    """Update the working-file path for this session."""
    sid = session.get("sid")
    if sid:
        current_app.config["SET_WORKING_PATH"](sid, path)


# ── Document helpers ──────────────────────────────────────────────────────

def extract_sentences_from_docx(path):
    doc = Document(path)
    sentences = []
    paragraph_index = 0

    for para in doc.paragraphs:
        paragraph_index += 1
        if not para.text.strip():
            continue
        blob = TextBlob(para.text)
        for sentence in blob.sentences:
            sentences.append({
                "paragraph": paragraph_index,
                "location_label": f"\u00b6{paragraph_index}",
                "source": "paragraph",
                "sentence": str(sentence),
            })

    for t_idx, table in enumerate(doc.tables, start=1):
        for r_idx, row in enumerate(table.rows, start=1):
            for c_idx, cell in enumerate(row.cells, start=1):
                cell_text = cell.text.strip()
                if not cell_text:
                    continue
                paragraph_index += 1
                label = f"T{t_idx} R{r_idx} C{c_idx}"
                blob = TextBlob(cell_text)
                for sentence in blob.sentences:
                    sentences.append({
                        "paragraph": paragraph_index,
                        "location_label": label,
                        "source": "table",
                        "sentence": str(sentence),
                    })

    return sentences


_ABBREV_PATTERN = _re.compile(
    r'^('
    r'[A-Z]{2,}'
    r'|[A-Za-z]([.][A-Za-z])+[.]?'
    r'|[A-Z][a-z]+[.]'
    r')$'
)


def _is_abbrev_or_acronym(text):
    token = text.strip().rstrip('.')
    return bool(
        _ABBREV_PATTERN.match(text.strip())
        or _ABBREV_PATTERN.match(token)
        or (len(token) >= 2 and token.isupper())
    )


def analyze_sentence(sentence):
    matches = tool.check(sentence)
    errors = []
    for match in matches:
        error_text = sentence[match.offset: match.offset + match.error_length]
        if _is_abbrev_or_acronym(error_text):
            continue
        errors.append({
            "error_text": error_text,
            "message": match.message,
            "suggestions": match.replacements[:6],
        })
    return errors


def analyze_docx(path):
    sentences = extract_sentences_from_docx(path)
    seen_sentences: dict = {}
    report = []
    total_errors = 0

    for s in sentences:
        sentence = s["sentence"]
        para_idx = s["paragraph"]
        label = s.get("location_label", f"\u00b6{para_idx}")
        source = s.get("source", "paragraph")

        if sentence in seen_sentences:
            first_idx = seen_sentences[sentence]
            report[first_idx]["occurrences"].append(para_idx)
            report[first_idx]["occurrence_labels"].append(label)
        else:
            errors = analyze_sentence(sentence)
            total_errors += len(errors)
            entry = {
                "paragraph":        para_idx,
                "location_label":   label,
                "source":           source,
                "occurrences":      [para_idx],
                "occurrence_labels": [label],
                "sentence":         sentence,
                "errors":           errors,
                "error_count":      len(errors),
            }
            seen_sentences[sentence] = len(report)
            report.append(entry)

    return report, len(sentences), total_errors


def set_paragraph_text(para, new_text):
    if not para.runs:
        para.add_run(new_text)
        return
    para.runs[0].text = new_text
    for run in para.runs[1:]:
        run.text = ""


def apply_corrections_to_docx(doc_path, report, decisions):
    doc = Document(doc_path)
    para_corrections: dict = {}

    for si, item in enumerate(report):
        for ei, err in enumerate(item["errors"]):
            key = f"{si}_{ei}"
            decision = decisions.get(key, {})
            if decision.get("action") == "accept":
                suggestion = decision.get("suggestion")
                error_text = err.get("error_text", "")
                if suggestion and error_text:
                    for para_idx in item.get("occurrences", [item["paragraph"]]):
                        para_corrections.setdefault(para_idx, []).append(
                            (error_text, suggestion)
                        )

    if not para_corrections:
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    para_counter = 0

    for para in doc.paragraphs:
        para_counter += 1
        if para_counter in para_corrections:
            new_text = para.text
            for old, replacement in para_corrections[para_counter]:
                new_text = new_text.replace(old, replacement, 1)
            if new_text != para.text:
                set_paragraph_text(para, new_text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if not cell.text.strip():
                    continue
                para_counter += 1
                if para_counter in para_corrections:
                    original_text = cell.text
                    new_text = original_text
                    for old, replacement in para_corrections[para_counter]:
                        new_text = new_text.replace(old, replacement, 1)
                    if new_text != original_text:
                        remaining = new_text
                        for p in cell.paragraphs:
                            if p.text and remaining:
                                set_paragraph_text(p, remaining[:len(p.text)])
                                remaining = remaining[len(p.text):]

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── Routes ────────────────────────────────────────────────────────────────

@grammar_bp.route("/", methods=["GET", "POST"])
def index():
    if request.method == "GET":
        return render_template("grammarUI.html")

    # POST — analyse the session's working file
    path = _get_path()
    if not path or not os.path.exists(path):
        return jsonify({"error": "No file uploaded yet. Please upload from the home page."}), 400

    report, sentence_count, total_errors = analyze_docx(path)
    return jsonify({
        "report":         report,
        "sentence_count": sentence_count,
        "total_errors":   total_errors,
    })


@grammar_bp.route("/download", methods=["POST"])
def download():
    decisions_raw = request.form.get("decisions", "{}")
    report_raw = request.form.get("report",    "[]")
    original_name = request.form.get("original_name", "document")

    path = _get_path()
    if not path or not os.path.exists(path):
        return jsonify({"error": "No file found for this session"}), 400

    decisions = json.loads(decisions_raw)
    report = json.loads(report_raw)

    corrected_buffer = apply_corrections_to_docx(path, report, decisions)

    # Persist the corrected file so the next tool works on it.
    sid = session.get("sid", "default")
    out_dir = os.path.join(OUTPUT, sid)
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, f"grammar_{uuid.uuid4().hex}.docx")
    with open(out_path, "wb") as fh:
        fh.write(corrected_buffer.getvalue())
    _set_path(out_path)

    corrected_buffer.seek(0)
    download_name = f"{os.path.splitext(original_name)[0]}_corrected.docx"
    return send_file(
        corrected_buffer,
        as_attachment=True,
        download_name=download_name,
        mimetype=(
            "application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.document"
        ),
    )


if __name__ == "__main__":
    from flask import Flask
    _app = Flask(__name__, template_folder="templates")
    _app.secret_key = "dev"
    _app.register_blueprint(grammar_bp)
    _app.run(debug=True)
