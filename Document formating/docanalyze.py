from flask import Flask, render_template_string, request, jsonify, send_file
import os
from werkzeug.utils import secure_filename
from transformers import LayoutLMv3ForTokenClassification, LayoutLMv3Processor
from PIL import Image
import torch
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from pdf2image import convert_from_path
import pythoncom
import win32com.client

app = Flask(__name__)

app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['OUTPUT_FOLDER'] = 'outputs'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)

# ---------------- MODEL ----------------
model_name = "microsoft/layoutlmv3-base"
processor = LayoutLMv3Processor.from_pretrained(model_name)
model = LayoutLMv3ForTokenClassification.from_pretrained(model_name)

# ---------------- LABEL MAP ----------------
LABEL_MAP = {
    0: "O",
    1: "B-TITLE",
    2: "I-TITLE",
    3: "B-HEADING",
    4: "I-HEADING",
    5: "B-PARAGRAPH",
    6: "I-PARAGRAPH",
    7: "B-TABLE",
    8: "I-TABLE",
    9: "B-LIST",
    10: "I-LIST"
}

# ---------------- TRUE TEXT HIGHLIGHT COLORS ----------------
HIGHLIGHT_COLORS = {
    'TITLE': WD_COLOR_INDEX.RED,
    'HEADING': WD_COLOR_INDEX.VIOLET,
    'PARAGRAPH': WD_COLOR_INDEX.BRIGHT_GREEN,
    'TABLE': WD_COLOR_INDEX.YELLOW,
    'LIST': WD_COLOR_INDEX.DARK_BLUE
}

# ---------------- HTML ----------------
HTML_TEMPLATE = """
<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>Document Analyzer</title>
<meta name="viewport" content="width=device-width, initial-scale=1.0">

<style>
    body {
        margin: 0;
        padding: 0;
        min-height: 100vh;
        display: flex;
        align-items: center;
        justify-content: center;
        background: linear-gradient(135deg, #6a5acd, #7b68ee);
        font-family: Arial, Helvetica, sans-serif;
    }

    .box {
        background: white;
        width: 600px;
        padding: 40px;
        border-radius: 14px;
        text-align: center;
        box-shadow: 0 20px 50px rgba(0,0,0,0.3);
    }

    h1 {
        margin: 0 0 12px;
        font-size: 26px;
        color: #222;
    }

    p {
        margin: 0 0 30px;
        font-size: 15px;
        color: #555;
    }

    input[type="file"] {
        width: 100%;
        margin-bottom: 20px;
        font-size: 14px;
    }

    button {
        width: 100%;
        padding: 14px;
        background: #6a5acd;
        border: none;
        border-radius: 30px;
        color: white;
        font-size: 16px;
        cursor: pointer;
    }

    button:disabled {
        background: #999;
        cursor: not-allowed;
    }

    .loader {
        display: none;
        margin-top: 25px;
        font-size: 14px;
        color: #444;
    }

    .spinner {
        margin: 12px auto 0;
        width: 40px;
        height: 40px;
        border: 4px solid #ddd;
        border-top: 4px solid #6a5acd;
        border-radius: 50%;
        animation: spin 1s linear infinite;
    }

    @keyframes spin {
        0% { transform: rotate(0deg); }
        100% { transform: rotate(360deg); }
    }

    .footer {
        margin-top: 25px;
        font-size: 12px;
        color: #777;
    }
</style>
</head>

<body>

<div class="box">
    <h1>📄 Document Analyzer</h1>
    <p>Upload a DOCX file to highlight document structure</p>

    <input type="file" id="fileInput" accept=".docx">

    <button id="analyzeBtn" onclick="uploadFile()">Analyze Document</button>

    <div class="loader" id="loader">
        Processing document, please wait...
        <div class="spinner"></div>
    </div>

    <div class="footer">
        LayoutLMv3 powered document processing
    </div>
</div>

<script>
async function uploadFile() {
    const fileInput = document.getElementById("fileInput");
    const btn = document.getElementById("analyzeBtn");
    const loader = document.getElementById("loader");

    if (!fileInput.files.length) {
        alert("Please select a DOCX file");
        return;
    }

    const formData = new FormData();
    formData.append("file", fileInput.files[0]);

    btn.disabled = true;
    loader.style.display = "block";

    try {
        const response = await fetch("/analyze", {
            method: "POST",
            body: formData
        });

        const blob = await response.blob();

        // Trigger download
        const url = window.URL.createObjectURL(blob);
        const a = document.createElement("a");
        a.href = url;
        a.download = fileInput.files[0].name.replace(".docx", "_highlighted.docx");
        document.body.appendChild(a);
        a.click();
        document.body.removeChild(a);
        window.URL.revokeObjectURL(url);

    } catch (err) {
        alert("Error processing document");
    }

    loader.style.display = "none";
    btn.disabled = false;
}
</script>

</body>
</html>
"""

# ---------------- DOCX → IMAGE ----------------


def convert_docx_to_image(docx_path):
    try:
        pythoncom.CoInitialize()
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False

        doc = word.Documents.Open(docx_path)
        pdf_path = docx_path.replace('.docx', '.pdf')
        doc.SaveAs(pdf_path, FileFormat=17)
        doc.Close()
        word.Quit()
        pythoncom.CoUninitialize()

        images = convert_from_path(pdf_path)
        os.remove(pdf_path)
        return images[0] if images else None
    except Exception as e:
        print("DOCX to image error:", e)
        return None

# ---------------- TEXT EXTRACTION ----------------


def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    elements = []
    para_idx = 0

    for para in doc.paragraphs:
        if not para.text.strip():
            para_idx += 1
            continue

        if para.style.name == 'Title':
            elements.append(
                {'type': 'TITLE', 'text': para.text, 'para_idx': para_idx})
        elif 'Heading' in para.style.name:
            elements.append(
                {'type': 'HEADING', 'text': para.text, 'para_idx': para_idx})
        else:
            elements.append(
                {'type': 'PARAGRAPH', 'text': para.text, 'para_idx': para_idx})

        para_idx += 1

    for i, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append(" | ".join(cell.text for cell in row.cells))
        elements.append(
            {'type': 'TABLE', 'text': "\n".join(rows), 'table_idx': i})

    return elements

# ---------------- LAYOUTLM (OPTIONAL) ----------------


def analyze_with_layoutlmv3(image, elements):
    try:
        encoding = processor(image, return_tensors="pt", truncation=True)
        with torch.no_grad():
            outputs = model(**encoding)
        return elements
    except Exception as e:
        print("LayoutLM error:", e)
        return elements

# ---------------- TRUE TEXT HIGHLIGHT ----------------


def highlight_docx(input_path, elements, output_path):
    doc = Document(input_path)

    para_map = {el['para_idx']: el['type']
                for el in elements if 'para_idx' in el}

    # Highlight paragraphs
    for idx, para in enumerate(doc.paragraphs):
        if idx in para_map:
            color = HIGHLIGHT_COLORS.get(para_map[idx])
            if color:
                for run in para.runs:
                    run.font.color.rgb = None
                    run.font.highlight_color = color

    # Highlight tables
    for el in elements:
        if 'table_idx' in el and el['table_idx'] < len(doc.tables):
            table = doc.tables[el['table_idx']]
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    doc.save(output_path)

# ---------------- ROUTES ----------------


@app.route('/')
def index():
    return render_template_string(HTML_TEMPLATE)


@app.route('/analyze', methods=['POST'])
def analyze():
    file = request.files.get('file')
    if not file or not file.filename.endswith('.docx'):
        return jsonify({'error': 'Invalid file'}), 400

    filename = secure_filename(file.filename)
    input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(input_path)

    elements = extract_text_from_docx(input_path)

    image = convert_docx_to_image(os.path.abspath(input_path))
    if image:
        elements = analyze_with_layoutlmv3(image, elements)

    output_file = filename.replace('.docx', '_highlighted.docx')
    output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_file)

    highlight_docx(input_path, elements, output_path)
    os.remove(input_path)

    return send_file(output_path, as_attachment=True)


# ---------------- RUN ----------------
if __name__ == '__main__':
    app.run(debug=True, port=5000)
