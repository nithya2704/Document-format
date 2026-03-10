from flask import Flask, request, render_template_string, jsonify, send_file
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx2pdf import convert
import os
import json
import uuid
import pythoncom
from werkzeug.utils import secure_filename

app = Flask(__name__)

UPLOAD = "uploads"
OUTPUT = "outputs"

os.makedirs(UPLOAD, exist_ok=True)
os.makedirs(OUTPUT, exist_ok=True)

# ---------------- ANALYSE ----------------


def analyse_docx(path):

    doc = Document(path)

    elements = []
    types_present = set()

    for i, p in enumerate(doc.paragraphs):

        text = p.text.strip()

        if not text:
            continue

        is_list = bool(p._element.xpath('.//w:numPr'))

        if is_list:
            elements.append({"type": "list", "index": i})
            types_present.add("list")
        else:
            elements.append({"type": "paragraph", "index": i})
            types_present.add("paragraph")

    return elements, sorted(types_present)


# ---------------- FORMAT ----------------

def format_docx(input_path, output_path, detected, cfg):

    doc = Document(input_path)

    alignment_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
    }

    idx_map = {d["index"]: d["type"] for d in detected}

    for i, p in enumerate(doc.paragraphs):

        etype = idx_map.get(i)

        if not etype:
            continue

        pf = p.paragraph_format

        align_value = cfg.get(f"{etype}_alignment", "left")
        p.alignment = alignment_map.get(align_value)

        spacing_value = float(cfg.get(f"{etype}_spacing", 1))
        pf.line_spacing = spacing_value

        if etype == "paragraph":

            options = cfg.get("paragraph_spacing_options", [])

            if "add_before" in options:
                pf.space_before = Pt(12)

            if "add_after" in options:
                pf.space_after = Pt(12)

            if "remove_before" in options:
                pf.space_before = Pt(0)

            if "remove_after" in options:
                pf.space_after = Pt(0)

    doc.save(output_path)


# ---------------- HTML UI ----------------

HTML = """

<!DOCTYPE html>
<html>

<head>

<title>DOCX Formatter</title>

<style>

body{
font-family:Arial;
background:#f3f4f6;
padding:40px;
}

.box{
background:white;
width:800px;
margin:auto;
padding:30px;
border-radius:12px;
}

button{
padding:12px;
width:100%;
background:#4f46e5;
color:white;
border:none;
border-radius:8px;
margin-top:20px;
cursor:pointer;
font-size:15px;
transition:all 0.25s ease;
}

button:hover{
background:#3730a3;
transform:translateY(-2px);
box-shadow:0 6px 14px rgba(0,0,0,0.2);
}

select{
width:100%;
padding:8px;
margin-top:5px;
}

.button-row{
display:flex;
gap:15px;
margin-top:20px;
}

.button-row button{
width:50%;
}

.spacing-group{
margin-top:15px;
}

.spacing-group label{
display:block;
margin-top:5px;
}

/* TOAST */

#toast{
visibility:hidden;
min-width:250px;
background:#10b981;
color:white;
text-align:center;
border-radius:8px;
padding:16px;
position:fixed;
z-index:2000;
right:30px;
top:30px;
box-shadow:0 8px 18px rgba(0,0,0,0.2);
}

#toast.show{
visibility:visible;
animation:fadein 0.5s, fadeout 0.5s 3s;
}

@keyframes fadein{
from{opacity:0; right:0;}
to{opacity:1; right:30px;}
}

@keyframes fadeout{
from{opacity:1;}
to{opacity:0;}
}

/* MODAL */

.modal{
display:none;
position:fixed;
z-index:1000;
left:0;
top:0;
width:100%;
height:100%;
background:rgba(0,0,0,0.6);
}

.modal-content{
background:white;
width:80%;
height:85%;
margin:3% auto;
border-radius:10px;
position:relative;
padding:20px;
}

.close{
position:absolute;
top:10px;
right:15px;
font-size:28px;
cursor:pointer;
}

iframe{
width:100%;
height:100%;
border:none;
}

</style>

</head>

<body>

<div id="toast"></div>

<div class="box">

<h2>📄 DOCX Advanced Formatter</h2>

<input type="file" id="file">

<br><br>

<button onclick="analyse()">Analyse Document</button>

<div id="ui"></div>

<div id="actions"></div>

</div>

<div id="previewModal" class="modal">

<div class="modal-content">

<span class="close" onclick="closePreview()">×</span>

<iframe id="previewFrame"></iframe>

</div>

</div>

<script>

let types=[];

function showToast(msg){

let t=document.getElementById("toast");

t.innerText=msg;

t.className="show";

setTimeout(()=>{t.className=t.className.replace("show","");},3500);

}

async function analyse(){

let fd=new FormData();
fd.append("file",file.files[0]);

let res=await fetch("/analyse",{method:"POST",body:fd});
let data=await res.json();

types=data.types;

let html="<h3>Formatting Options</h3>";

types.forEach(t=>{

html+=`

<h4>${t}</h4>

<label>Alignment</label>

<select id='${t}_alignment'>
<option value="left">Left</option>
<option value="right">Right</option>
<option value="justify">Justify</option>
</select>

<label>Line Spacing</label>

<select id='${t}_spacing'>
<option value="1">Single</option>
<option value="1.5">1.5</option>
<option value="2">Double</option>
</select>

`;

});

if(types.includes("paragraph")){

html+=`

<div class="spacing-group">

<h4>Paragraph Spacing</h4>

<label><input type="checkbox" class="spacingOpt" value="add_before"> Add Space Before Paragraph</label>

<label><input type="checkbox" class="spacingOpt" value="add_after"> Add Space After Paragraph</label>

<label><input type="checkbox" class="spacingOpt" value="remove_before"> Remove Space Before Paragraph</label>

<label><input type="checkbox" class="spacingOpt" value="remove_after"> Remove Space After Paragraph</label>

</div>

`;

}

html+="<button onclick='generate()'>Generate</button>";

ui.innerHTML=html;

}


async function generate(){

let cfg={};

types.forEach(t=>{

cfg[`${t}_alignment`]=document.getElementById(`${t}_alignment`).value;
cfg[`${t}_spacing`]=document.getElementById(`${t}_spacing`).value;

});

let spacingOpts=[];

document.querySelectorAll(".spacingOpt").forEach(cb=>{
if(cb.checked) spacingOpts.push(cb.value);
});

cfg["paragraph_spacing_options"]=spacingOpts;

window.currentConfig=cfg;

showToast("Document generated successfully. You can preview or download.");

actions.innerHTML=`

<div class="button-row">

<button onclick="previewDoc()">Preview</button>

<button onclick="downloadDoc()">Download DOCX</button>

</div>

`;

}


async function previewDoc(){

showToast("Preparing preview...");

let fd=new FormData();

fd.append("file",file.files[0]);
fd.append("config",JSON.stringify(window.currentConfig));

let res=await fetch("/preview",{method:"POST",body:fd});

let blob=await res.blob();

let url=URL.createObjectURL(blob);

document.getElementById("previewFrame").src=url;

document.getElementById("previewModal").style.display="block";

}


function closePreview(){

document.getElementById("previewModal").style.display="none";

}


async function downloadDoc(){

showToast("Downloading document...");

let fd=new FormData();

fd.append("file",file.files[0]);
fd.append("config",JSON.stringify(window.currentConfig));

let res=await fetch("/format",{method:"POST",body:fd});

let blob=await res.blob();

let a=document.createElement("a");

a.href=URL.createObjectURL(blob);

a.download="formatted.docx";

a.click();

}

</script>

</body>
</html>

"""


# ---------------- ROUTES ----------------

@app.route("/")
def home():
    return render_template_string(HTML)


@app.route("/analyse", methods=["POST"])
def analyse():

    f = request.files["file"]

    path = os.path.join(UPLOAD, secure_filename(f.filename))

    f.save(path)

    detected, types = analyse_docx(path)

    return jsonify({"types": types})


@app.route("/preview", methods=["POST"])
def preview():

    pythoncom.CoInitialize()

    f = request.files["file"]
    cfg = json.loads(request.form["config"])

    in_path = os.path.join(UPLOAD, secure_filename(f.filename))
    f.save(in_path)

    out_docx = os.path.join(OUTPUT, f"formatted_{uuid.uuid4()}.docx")

    detected, _ = analyse_docx(in_path)

    format_docx(in_path, out_docx, detected, cfg)

    pdf_path = os.path.join(OUTPUT, f"preview_{uuid.uuid4()}.pdf")

    convert(out_docx, pdf_path)

    pythoncom.CoUninitialize()

    return send_file(pdf_path)


@app.route("/format", methods=["POST"])
def format_doc():

    f = request.files["file"]
    cfg = json.loads(request.form["config"])

    in_path = os.path.join(UPLOAD, secure_filename(f.filename))
    f.save(in_path)

    out_path = os.path.join(OUTPUT, "formatted.docx")

    detected, _ = analyse_docx(in_path)

    format_docx(in_path, out_path, detected, cfg)

    return send_file(out_path, as_attachment=True)


if __name__ == "__main__":
    app.run(debug=True, threaded=False)
