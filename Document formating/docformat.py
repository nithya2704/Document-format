from flask import Flask, render_template_string, request, send_file, jsonify
import os, json
from werkzeug.utils import secure_filename

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX

import pythoncom
import win32com.client
from pdf2image import convert_from_path

from docx.oxml import OxmlElement, ns


# ================== FLASK SETUP ==================
app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['OUTPUT_FOLDER'] = 'outputs'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)

# ================== COLORS ==================
HIGHLIGHT_COLORS = {
    "TITLE": WD_COLOR_INDEX.YELLOW,
    "HEADING": WD_COLOR_INDEX.BRIGHT_GREEN,
    "PARAGRAPH": WD_COLOR_INDEX.GRAY_25,
    "TABLE": WD_COLOR_INDEX.TURQUOISE
}

# ================== UI ==================
HTML_TEMPLATE = """
<!DOCTYPE html>
<html>
<head>
<title>Smart Document Formatter</title>
<style>
body {
    background: linear-gradient(135deg,#6a5acd,#7b68ee);
    font-family: Arial;
    padding: 40px;
}
.box {
    background: white;
    width: 750px;
    padding: 30px;
    border-radius: 15px;
    margin: auto;
}
h1 { text-align:center; }
label { font-weight:bold; }
select,input { width:100%; padding:8px; margin-bottom:10px; }
button {
    width:100%;
    padding:14px;
    background:#6a5acd;
    border:none;
    border-radius:30px;
    color:white;
    font-size:16px;
}
</style>
</head>

<body>
<div class="box">
<h1>📄 Smart Document Formatter</h1>

<input type="file" id="file">

<h3>Title Style</h3>
<label>Font</label>
<select id="title_font"><option>Times New Roman</option><option>Arial</option><option>Calibri</option></select>
<label>Size</label>
<input type="number" id="title_size" value="26">

<h3>Heading Style</h3>
<label>Font</label>
<select id="heading_font"><option>Calibri</option><option>Arial</option><option>Times New Roman</option></select>
<label>Size</label>
<input type="number" id="heading_size" value="18">

<h3>Paragraph Style</h3>
<label>Font</label>
<select id="para_font"><option>Calibri</option><option>Times New Roman</option><option>Arial</option></select>
<label>Size</label>
<input type="number" id="para_size" value="12">

<label><input type="checkbox" id="bold_titles" checked> Bold Titles</label><br>
<label><input type="checkbox" id="highlight" checked> Highlight Sections</label>

<br><br>
<button onclick="submitDoc()">Format Document</button>
</div>

<script>
async function submitDoc(){
    let file = document.getElementById("file").files[0];
    if(!file){ alert("Upload DOCX"); return; }

    let config = {
        title_font: title_font.value,
        title_size: title_size.value,
        heading_font: heading_font.value,
        heading_size: heading_size.value,
        para_font: para_font.value,
        para_size: para_size.value,
        bold_titles: bold_titles.checked,
        highlight: highlight.checked
    };

    let fd = new FormData();
    fd.append("file", file);
    fd.append("config", JSON.stringify(config));

    let res = await fetch("/analyze", {method:"POST", body:fd});
    let blob = await res.blob();

    let a = document.createElement("a");
    a.href = URL.createObjectURL(blob);
    a.download = file.name.replace(".docx","_formatted.docx");
    a.click();
}
</script>
</body>
</html>
"""

# ================== UTILITIES ==================
def extract_text_structure(docx_path):
    doc = Document(docx_path)
    elements = []
    idx = 0

    for p in doc.paragraphs:
        if not p.text.strip():
            idx += 1
            continue

        if p.style.name == "Title":
            t = "TITLE"
        elif "Heading" in p.style.name:
            t = "HEADING"
        else:
            t = "PARAGRAPH"

        elements.append({"type": t, "para_idx": idx})
        idx += 1

    for i, _ in enumerate(doc.tables):
        elements.append({"type": "TABLE", "table_idx": i})

    return elements

def add_border_to_run_images(run, border_pt=0.25):
    """
    Adds a border to an inline image inside a run.
    border_pt is in points (0.25 pt requested).
    """
    drawings = run._element.xpath('.//w:drawing')
    if not drawings:
        return

    # drawing = drawings[0]
    for drawing in drawings:
        inlines = drawing.xpath('.//wp:inline')
        anchors = drawing.xpath('.//wp:anchor')

        for container in inlines + anchors:
            spPr = container.xpath('.//pic:spPr')
            if not spPr:
                continue

            spPr = spPr[0]
    ## if not inline:
    ##     return

        ## inline = inline[0]
    ## Navigate to pic:spPr (shape properties)
    ## spPr = inline.xpath('.//pic:spPr')
    ## if not spPr:
    ##     return

    ## spPr = spPr[0]

    # Remove existing border if present (avoid duplicates)
            for ln in spPr.xpath('.//a:ln'):
                spPr.remove(ln)

    # Create graphic frame properties
    ## graphicFramePr = OxmlElement('wp:graphicFramePr')
    ## graphicFrameLocks = OxmlElement('a:graphicFrameLocks')
    ## graphicFrameLocks.set(ns.qn('noChangeAspect'), '1')
    ## graphicFramePr.append(graphicFrameLocks)
    ## inline.append(graphicFramePr)

    ## graphic = inline.xpath('.//a:graphic')
    ## if not graphic:
    ##     return
    ## graphic = graphic[0]

    # Create border
            ln = OxmlElement('a:ln')
            ln.set('w', str(int(border_pt * 12700)))  # pt → EMUs

            solidFill = OxmlElement('a:solidFill')
            srgbClr = OxmlElement('a:srgbClr')
            srgbClr.set('val', '000000')  # black
            solidFill.append(srgbClr)

            ln.append(solidFill)

    # Optional: square corners
            prstDash = OxmlElement('a:prstDash')
            prstDash.set('val', 'solid')
            ln.append(prstDash)

            spPr.append(ln)

    # Apply to graphic
    ## graphic = inline.xpath('.//a:graphic')[0]
    ## graphic.append(ln)


def format_docx(input_path, elements, output_path, config):
    doc = Document(input_path)

    para_map = {e["para_idx"]: e["type"] for e in elements if "para_idx" in e}

    # -------- FORMAT PARAGRAPHS --------
    for idx, para in enumerate(doc.paragraphs):
        ptype = para_map.get(idx)

        for run in para.runs:
            # IMAGE HANDLING
            if run._element.xpath('.//w:drawing'):
                add_border_to_run_images(run, border_pt=0.25)
                continue

            if not ptype:
                continue

            if ptype == "TITLE":
                run.font.name = config["title_font"]
                run.font.size = Pt(int(config["title_size"]))
                run.font.bold = config["bold_titles"]

            elif ptype == "HEADING":
                run.font.name = config["heading_font"]
                run.font.size = Pt(int(config["heading_size"]))
                run.font.bold = True

            elif ptype == "PARAGRAPH":
                run.font.name = config["para_font"]
                run.font.size = Pt(int(config["para_size"]))

            if config["highlight"]:
                run.font.highlight_color = HIGHLIGHT_COLORS.get(ptype)

    # -------- FORMAT TABLES --------
    for e in elements:
        if e["type"] == "TABLE" and e["table_idx"] < len(doc.tables):
            table = doc.tables[e["table_idx"]]
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            # IMAGE IN TABLE
                            if run._element.xpath('.//w:drawing'):
                                add_border_to_run_images(run, border_pt=0.25)
                            elif config["highlight"]:
                                run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    doc.save(output_path)


# ================== ROUTES ==================
@app.route("/")
def index():
    return render_template_string(HTML_TEMPLATE)

@app.route("/analyze", methods=["POST"])
def analyze():
    file = request.files.get("file")
    config = json.loads(request.form.get("config"))

    if not file or not file.filename.endswith(".docx"):
        return jsonify({"error": "Invalid file"}), 400
    
    original_name = secure_filename(file.filename)
    base, ext = os.path.splitext(original_name)

    input_path = os.path.join(app.config["UPLOAD_FOLDER"], original_name)
    output_filename = f"{base}_formatted{ext}"
    output_path = os.path.join(app.config["OUTPUT_FOLDER"], output_filename)

    # fname = secure_filename(file.filename)
    # in_path = os.path.join(app.config["UPLOAD_FOLDER"], fname)
    # out_path = os.path.join(app.config["OUTPUT_FOLDER"], fname.replace(".docx","_formatted.docx"))

    file.save(input_path)

    elements = extract_text_structure(input_path)
    format_docx(input_path, elements, output_path, config)

    return send_file(output_path, as_attachment=True, download_name=output_filename)

# ================== RUN ==================
if __name__ == "__main__":
    app.run(debug=True, port=5000)
