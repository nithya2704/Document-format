#!/usr/bin/env python3
"""
Font Formatter - Apply consistent fonts to Word documents
Fixed version: Applies table header bg color to ALL tables EXCEPT cover page
"""
import os
import re
import json
import uuid
import platform
from collections import Counter

from flask import (Blueprint, Flask, request, jsonify,
                   send_file, render_template, session, current_app)

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ── Platform
IS_WINDOWS = platform.system() == "Windows"
WORD_AVAILABLE = False
pythoncom = None
convert = None
if IS_WINDOWS:
    try:
        import pythoncom
        from docx2pdf import convert
        WORD_AVAILABLE = True
    except ImportError:
        pass

# ── Blueprint
font_bp = Blueprint("font", __name__, url_prefix="/font",
                    template_folder="templates")

OUTPUT = "outputs/font"
os.makedirs(OUTPUT, exist_ok=True)

# =============================================================================
#  SESSION HELPERS
# =============================================================================


def _get_path():
    sid = session.get("sid")
    if not sid:
        return None
    return current_app.config["GET_WORKING_PATH"](sid)


def _set_path(path: str):
    sid = session.get("sid")
    if sid:
        current_app.config["SET_WORKING_PATH"](sid, path)


def _session_out_dir():
    sid = session.get("sid", "default")
    d = os.path.join(OUTPUT, sid)
    os.makedirs(d, exist_ok=True)
    return d

# =============================================================================
#  UTILITIES
# =============================================================================


def is_in_footer_or_header(para) -> bool:
    _FH = {qn("w:hdr"), qn("w:ftr")}
    node = para._element.getparent()
    while node is not None:
        if node.tag in _FH:
            return True
        node = node.getparent()
    return False


def get_paragraph_indentation(para) -> float:
    try:
        if para.paragraph_format.left_indent:
            return para.paragraph_format.left_indent.inches
    except Exception:
        pass
    return 0.0


def has_page_break_before(para) -> bool:
    try:
        pPr = para._element.find(qn("w:pPr"))
        if pPr is not None:
            pb = pPr.find(qn("w:pageBreakBefore"))
            if pb is not None:
                val = pb.get(qn("w:val"), "true")
                if val.lower() not in ("false", "0", "off"):
                    return True
        for br in para._element.xpath(".//w:br"):
            if br.get(qn("w:type"), "") == "page":
                return True
    except Exception:
        pass
    return False


# ── Numbered heading detection
NUMBERED_HEADING_RE = re.compile(r"^\s*(\d+(?:\.\d+)+)\.?(.*)$")


def detect_numbered_heading(text: str):
    if not text:
        return None, None, None
    m = NUMBERED_HEADING_RE.match(text)
    if not m:
        return None, None, None
    num = m.group(1)
    title = m.group(2).strip()
    parts = num.split(".")
    if len(parts) < 2:
        return None, None, None
    return min(6, len(parts)), num, title


# ── Numbered list resolution
_NUMBERED_FMTS = frozenset({
    "decimal", "lowerLetter", "upperLetter", "lowerRoman",
    "upperRoman", "ordinal", "cardinalText", "decimalZero",
})


def _resolve_numFmt(para) -> str:
    try:
        numbering_part = para.part.numbering_part
    except Exception:
        numbering_part = None

    def _fmt(numId_val, ilvl_val):
        if numId_val == "0":
            return ""
        if numbering_part is None:
            return "bullet"
        nbr = numbering_part._element
        abs_nodes = nbr.xpath(
            f".//w:num[@w:numId='{numId_val}']/w:abstractNumId")
        if not abs_nodes:
            return "bullet"
        abstract_id = abs_nodes[0].get(qn("w:val"), "0")
        fmt_nodes = nbr.xpath(
            f".//w:abstractNum[@w:abstractNumId='{abstract_id}']"
            f"/w:lvl[@w:ilvl='{ilvl_val}']/w:numFmt")
        return fmt_nodes[0].get(qn("w:val"), "") if fmt_nodes else "bullet"

    pPr = para._element.find(qn("w:pPr"))
    if pPr is not None:
        numPr = pPr.find(qn("w:numPr"))
        if numPr is not None:
            nid = numPr.find(qn("w:numId"))
            ilvl = numPr.find(qn("w:ilvl"))
            nid_v = nid.get(qn("w:val"), "0") if nid is not None else "0"
            ilvl_v = ilvl.get(qn("w:val"), "0") if ilvl is not None else "0"
            if nid_v == "0":
                return ""
            return _fmt(nid_v, ilvl_v)

    style = para.style if para.style else None
    while style is not None:
        try:
            s_el = style.element
            s_pPr = s_el.find(qn("w:pPr"))
            if s_pPr is not None:
                s_numPr = s_pPr.find(qn("w:numPr"))
                if s_numPr is not None:
                    nid = s_numPr.find(qn("w:numId"))
                    ilvl = s_numPr.find(qn("w:ilvl"))
                    nid_v = nid.get(
                        qn("w:val"), "0") if nid is not None else "0"
                    ilvl_v = ilvl.get(
                        qn("w:val"), "0") if ilvl is not None else "0"
                    if nid_v == "0":
                        return ""
                    return _fmt(nid_v, ilvl_v)
        except Exception:
            pass
        try:
            style = style.base_style
        except Exception:
            break
    return ""


def get_list_type(para) -> str:
    try:
        sn = para.style.name if para.style else ""
    except Exception:
        sn = ""
    if "List Bullet" in sn or "List Number" in sn or sn.startswith("List "):
        return "LIST_ITEM"
    fmt = _resolve_numFmt(para)
    return "LIST_ITEM" if fmt else None

# ── Cover page detection


def _detect_cover_end(doc) -> int:
    paragraphs = doc.paragraphs

    # Strategy 1: explicit page-break
    for i, para in enumerate(paragraphs[:60]):
        for run in para.runs:
            for br in run._element.xpath(".//w:br"):
                if br.get(qn("w:type"), "") == "page":
                    return i
        if i > 0 and has_page_break_before(para):
            return i - 1

    # Strategy 2: first Heading style
    for i, para in enumerate(paragraphs[:40]):
        sn = para.style.name if para.style else ""
        if sn.startswith("Heading") and i > 0:
            return i - 1

    # Strategy 3: keyword heuristic
    COVER_KW = re.compile(
        r"\b(prepared|author|version|copyright|confidential|restricted|"
        r"january|february|march|april|may|june|july|august|"
        r"september|october|november|december|\d{4})\b", re.IGNORECASE)
    non_empty = [(i, p) for i, p in enumerate(
        paragraphs) if (p.text or "").strip()]
    cover_candidate = -1
    for rank, (i, para) in enumerate(non_empty[:25]):
        text = (para.text or "").strip()
        sn = para.style.name if para.style else ""
        if (len(text) < 200 or para.alignment == WD_ALIGN_PARAGRAPH.CENTER or
                bool(COVER_KW.search(text)) or sn in ("Title", "Subtitle")):
            cover_candidate = i
        elif rank > 2:
            break
    return cover_candidate

# ── Table header row detection
# ✅ FIX: Row 0 is ALWAYS the header. Also honour explicit tblHeader XML flag.


def _is_table_header_row(table, row_idx: int) -> bool:
    """Row 0 is always treated as a header row. Also honours tblHeader XML flag."""
    if row_idx == 0:
        return True
    try:
        row = table.rows[row_idx]
        trPr = row._tr.find(qn("w:trPr"))
        if trPr is not None:
            tblHeader = trPr.find(qn("w:tblHeader"))
            if tblHeader is not None:
                val = tblHeader.get(qn("w:val"), "true")
                if val.lower() not in ("false", "0", "off"):
                    return True
    except Exception:
        pass
    return False


def _is_table_footer_row(table, row_idx: int, total_rows: int) -> bool:
    if row_idx != total_rows - 1:
        return False
    try:
        row_text = " ".join(
            cell.text for cell in table.rows[row_idx].cells).lower()
        FOOTER_KW = re.compile(
            r"\b(total|sum|grand\s+total|subtotal|average|avg|count)\b", re.IGNORECASE)
        return bool(FOOTER_KW.search(row_text))
    except Exception:
        return False

# ── Font resolution


def _make_theme_resolver(doc):
    _cache = {}
    try:
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        theme_part = doc.part.part_related_by(RT.THEME)
        theme_root = theme_part._element
        A = "http://schemas.openxmlformats.org/drawingml/2006/main"

        def _find(tags):
            el = theme_root
            for tag in tags:
                el = el.find(f"{{{A}}}{tag}")
                if el is None:
                    return ""
            return el.get("typeface", "") if el is not None else ""
        major = _find(["themeElements", "fontScheme", "majorFont", "latin"])
        minor = _find(["themeElements", "fontScheme", "minorFont", "latin"])
        _cache = {
            "majorHAnsi": major, "majorAscii": major, "majorBidi": major, "majorEastAsia": major,
            "minorHAnsi": minor, "minorAscii": minor, "minorBidi": minor, "minorEastAsia": minor,
        }
    except Exception:
        pass
    return lambda v: _cache.get(v, "")


def _get_para_font(para, theme_resolver) -> str:
    def _from_rFonts(el):
        if el is None:
            return ""
        for attr in (qn("w:ascii"), qn("w:hAnsi"), qn("w:cs")):
            v = el.get(attr, "")
            if v:
                return v
        for attr in (qn("w:asciiTheme"), qn("w:hAnsiTheme"), qn("w:cstheme")):
            v = el.get(attr, "")
            if v:
                resolved = theme_resolver(v)
                if resolved:
                    return resolved
        return ""

    try:
        run_fonts = []
        for run in para.runs:
            if not (run.text or "").strip():
                continue
            rPr = run._r.find(qn("w:rPr"))
            if rPr is not None:
                f = _from_rFonts(rPr.find(qn("w:rFonts")))
                if f:
                    run_fonts.append(f)
            if not run_fonts and run.font and run.font.name:
                run_fonts.append(run.font.name)
        if run_fonts:
            return Counter(run_fonts).most_common(1)[0][0]

        pPr = para._element.find(qn("w:pPr"))
        if pPr is not None:
            rPr = pPr.find(qn("w:rPr"))
            if rPr is not None:
                f = _from_rFonts(rPr.find(qn("w:rFonts")))
                if f:
                    return f

        style = para.style
        while style is not None:
            try:
                if style.font and style.font.name:
                    return style.font.name
                s_el = style.element
                s_pPr = s_el.find(qn("w:pPr"))
                if s_pPr is not None:
                    s_rPr = s_pPr.find(qn("w:rPr"))
                    if s_rPr is not None:
                        f = _from_rFonts(s_rPr.find(qn("w:rFonts")))
                        if f:
                            return f
                s_rPr2 = s_el.find(qn("w:rPr"))
                if s_rPr2 is not None:
                    f = _from_rFonts(s_rPr2.find(qn("w:rFonts")))
                    if f:
                        return f
            except Exception:
                pass
            try:
                style = style.base_style
            except Exception:
                break
    except Exception:
        pass
    return ""

# =============================================================================
#  SKIP-REGION HELPERS  (cover, TOC, Table of Figures)
# =============================================================================


# The ordered list of element types shown in the UI
USER_TYPES = [
    "HEADING_1", "HEADING_2", "HEADING_3",
    "HEADING_4", "HEADING_5", "HEADING_6",
    "PARAGRAPH", "LIST_ITEM", "TABLE_TEXT",
]


def _style_base(para) -> str:
    """Return the base style name of a paragraph, stripping ' Char' suffixes."""
    try:
        return (para.style.name if para.style else "").split(" Char")[0].strip()
    except Exception:
        return ""


# Style names that Word uses for TOC / List of Figures / List of Tables entries
_TOC_STYLE_PREFIXES = (
    "toc ",        # "TOC 1" … "TOC 9"
    "TOC ",
    "table of contents",
    "contents heading",
    "figure caption",
    "table caption",
    "caption",
    "list of figures",
    "list of tables",
)

# Keywords that identify a heading as a TOC/LOF/LOT section title
_TOC_HEADING_KW = re.compile(
    r"^\s*(table\s+of\s+(contents?|figures?|tables?)"
    r"|contents?"
    r"|list\s+of\s+(figures?|tables?|abbreviations?|acronyms?))\s*$",
    re.IGNORECASE,
)


def _detect_toc_ranges(doc) -> set:
    """
    Return a set of paragraph indices that belong to:
      - Table of Contents
      - Table of Figures / List of Figures
      - Table of Tables / List of Tables

    Strategy:
    1. Any paragraph whose style name starts with a known TOC/caption prefix.
    2. A heading whose text matches _TOC_HEADING_KW triggers skipping of all
       following paragraphs until the next heading of the same or higher level,
       or until an explicit page-break is encountered.
    """
    skip = set()
    paragraphs = doc.paragraphs

    i = 0
    while i < len(paragraphs):
        para = paragraphs[i]
        sn = _style_base(para).lower()
        text = (para.text or "").strip()

        # Rule 1: style-name match — always skip these
        if any(sn.startswith(pfx.lower()) for pfx in _TOC_STYLE_PREFIXES):
            skip.add(i)
            i += 1
            continue

        # Rule 2: heading text matches TOC/LOF/LOT pattern
        if _TOC_HEADING_KW.match(text):
            skip.add(i)
            # Determine the heading level so we know when the section ends
            heading_level = None
            if sn.startswith("heading"):
                try:
                    heading_level = int(sn.split()[-1])
                except Exception:
                    heading_level = 1

            i += 1
            while i < len(paragraphs):
                p = paragraphs[i]
                p_sn = _style_base(p).lower()
                p_text = (p.text or "").strip()

                # Stop at a page-break
                if has_page_break_before(p):
                    break

                # Stop when we hit a heading of the same or higher level
                if p_sn.startswith("heading"):
                    try:
                        lvl = int(p_sn.split()[-1])
                    except Exception:
                        lvl = 1
                    if heading_level is None or lvl <= heading_level:
                        break

                # Also stop if the paragraph text looks like a new section heading
                lvl2, _, _ = detect_numbered_heading(p_text)
                if lvl2 and (heading_level is None or lvl2 <= heading_level):
                    break

                skip.add(i)
                i += 1
            continue

        i += 1

    return skip


# =============================================================================
#  ANALYSIS
# =============================================================================


HIGHLIGHT_COLORS = {
    "HEADING_1": WD_COLOR_INDEX.BRIGHT_GREEN,
    "HEADING_2": WD_COLOR_INDEX.TURQUOISE,
    "HEADING_3": WD_COLOR_INDEX.PINK,
    "HEADING_4": WD_COLOR_INDEX.VIOLET,
    "HEADING_5": WD_COLOR_INDEX.TEAL,
    "HEADING_6": WD_COLOR_INDEX.DARK_BLUE,
    "PARAGRAPH": WD_COLOR_INDEX.WHITE,
    "LIST_ITEM": WD_COLOR_INDEX.YELLOW,
    "TABLE_TEXT": WD_COLOR_INDEX.RED,
}

_COLOR_MAP = {
    WD_COLOR_INDEX.BRIGHT_GREEN: "green",
    WD_COLOR_INDEX.TURQUOISE: "cyan",
    WD_COLOR_INDEX.PINK: "magenta",
    WD_COLOR_INDEX.VIOLET: "darkMagenta",
    WD_COLOR_INDEX.TEAL: "darkCyan",
    WD_COLOR_INDEX.DARK_BLUE: "darkBlue",
    WD_COLOR_INDEX.WHITE: "white",
    WD_COLOR_INDEX.YELLOW: "yellow",
    WD_COLOR_INDEX.RED: "red"
}


def ft_analyze_document_structure(docx_path: str) -> dict:
    """Analyze document and detect element types + table header colors."""
    doc = Document(docx_path)
    theme_resolver = _make_theme_resolver(doc)

    elements = []
    element_counts = {}
    detected_types = set()
    sample_texts = {}
    _font_votes = {}
    detected_table_bg_colors = set()

    # Detect regions to ignore
    cover_end = _detect_cover_end(doc)
    toc_skip_indices = _detect_toc_ranges(doc)

    # Build body position map
    body_children = list(doc.element.body)
    body_elem_to_pos = {id(child): pos for pos,
                        child in enumerate(body_children)}

    para_body_pos = {}
    para_seq = 0
    for child in body_children:
        if child.tag == qn("w:p"):
            para_body_pos[para_seq] = body_elem_to_pos[id(child)]
            para_seq += 1

    cover_body_threshold = -1
    if cover_end >= 0 and cover_end in para_body_pos:
        cover_body_threshold = para_body_pos[cover_end]

    # Build set of paragraph element IDs inside table cells so they are not
    # double-classified in the main paragraph loop.
    table_para_ids = set()
    for tbl in doc.tables:
        for row in tbl.rows:
            seen_tc_ids = set()
            for cell in row.cells:
                tc_id = id(cell._tc)
                if tc_id in seen_tc_ids:
                    continue
                seen_tc_ids.add(tc_id)
                for cp in cell.paragraphs:
                    table_para_ids.add(id(cp._element))

    # ── STEP 1: Body paragraphs ────────────────────────────────────────────
    for idx, para in enumerate(doc.paragraphs):
        if is_in_footer_or_header(para):
            continue
        if not (para.text or "").strip():
            continue
        if id(para._element) in table_para_ids:
            continue
        if cover_end >= 0 and idx <= cover_end:
            continue
        if idx in toc_skip_indices:
            continue

        sb = _style_base(para)
        raw_text = (para.text or "").strip()
        ptype = None

        # Prefer numbered heading first
        lvl, _, _ = detect_numbered_heading(raw_text)
        if lvl:
            ptype = f"HEADING_{lvl}"
        elif sb.startswith("Heading"):
            try:
                level = max(1, min(6, int(sb.split()[-1])))
                ptype = f"HEADING_{level}"
            except Exception:
                ptype = "HEADING_1"
        elif get_list_type(para):
            ptype = "LIST_ITEM"
        else:
            ptype = "PARAGRAPH"

        elements.append({
            "type": ptype,
            "para_idx": idx,
            "indent": get_paragraph_indentation(para),
        })
        element_counts[ptype] = element_counts.get(ptype, 0) + 1
        detected_types.add(ptype)

        if ptype not in sample_texts:
            sample_texts[ptype] = raw_text[:250]

        font = _get_para_font(para, theme_resolver)
        if font:
            _font_votes.setdefault(ptype, []).append(font)

    # ── STEP 2: Tables ──────────────────────────────────────────────────────
    for tbl_idx, table in enumerate(doc.tables):
        tbl_body_pos = body_elem_to_pos.get(id(table._tbl), -1)
        on_cover = (cover_body_threshold >= 0 and 0 <=
                    tbl_body_pos <= cover_body_threshold)

        # Detect existing header bg colors from first row only
        if not on_cover:
            try:
                if table.rows:
                    first_row = table.rows[0]
                    seen_tc = set()
                    for cell in first_row.cells:
                        tc_id = id(cell._tc)
                        if tc_id in seen_tc:
                            continue
                        seen_tc.add(tc_id)

                        tcPr = cell._tc.find(qn("w:tcPr"))
                        if tcPr is None:
                            continue

                        shd = tcPr.find(qn("w:shd"))
                        if shd is None:
                            continue

                        fill = (shd.get(qn("w:fill"), "") or "").strip()
                        if fill and fill.upper() not in ("AUTO", "FFFFFF", ""):
                            detected_table_bg_colors.add(f"#{fill.upper()}")
                        else:
                            shade = shd.get(qn("w:themeShade"))
                            tint = shd.get(qn("w:themeTint"))
                            if shade or tint:
                                detected_table_bg_colors.add("#D3D3D3")
            except Exception:
                pass

        elements.append({
            "type": "TABLE",
            "table_idx": tbl_idx,
            "on_cover": on_cover,
        })
        element_counts["TABLE_TEXT"] = element_counts.get("TABLE_TEXT", 0) + 1
        detected_types.add("TABLE_TEXT")

        # Detect each table-cell paragraph as TABLE_TEXT
        for row_idx, row in enumerate(table.rows):
            seen_tc_ids = set()
            for cell_idx, cell in enumerate(row.cells):
                tc_id = id(cell._tc)
                if tc_id in seen_tc_ids:
                    continue
                seen_tc_ids.add(tc_id)

                for para_idx_in_cell, cp in enumerate(cell.paragraphs):
                    if not (cp.text or "").strip():
                        continue
                    if on_cover:
                        continue

                    elements.append({
                        "type": "TABLE_TEXT",
                        "table_idx": tbl_idx,
                        "row_idx": row_idx,
                        "cell_idx": cell_idx,
                        "para_idx_in_cell": para_idx_in_cell,
                        "indent": get_paragraph_indentation(cp),
                    })
                    element_counts["TABLE_TEXT"] = element_counts.get(
                        "TABLE_TEXT", 0) + 1
                    detected_types.add("TABLE_TEXT")

                    if "TABLE_TEXT" not in sample_texts:
                        sample_texts["TABLE_TEXT"] = (
                            cp.text or "").strip()[:250]

                    font = _get_para_font(cp, theme_resolver)
                    if font:
                        _font_votes.setdefault("TABLE_TEXT", []).append(font)

    # Keep UI-facing output aligned with existing types
    _USER_ORDER = {t: i for i, t in enumerate(USER_TYPES)}

    def _sort_key(x):
        return _USER_ORDER.get(x, 99)

    visible_types = [t for t in detected_types if t in _USER_ORDER]

    current_fonts = {
        pt: Counter(votes).most_common(1)[0][0]
        for pt, votes in _font_votes.items() if votes and pt in _USER_ORDER
    }

    return {
        "elements": elements,
        "detected_elements": sorted(visible_types, key=_sort_key),
        "element_counts": element_counts,
        "sample_texts": sample_texts,
        "current_fonts": current_fonts,
        "detected_table_bgs": sorted(list(detected_table_bg_colors)),
    }

# =============================================================================
#  STANDARD COLORS
# =============================================================================


STANDARD_TABLE_BG_COLORS = [
    "#4472C4", "#70AD47", "#FFC000", "#ED7D31", "#A5A5A5",
    "#44546A", "#BDD7EE", "#C6EFCE", "#FFE699", "#F4B084",
]

# =============================================================================
#  FORMATTING
# =============================================================================


def _apply_run_props(rPr, font_name, font_size, bold):
    """Apply font/size/bold to rPr element."""
    if font_name:
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        for attr, val in (("w:ascii", font_name), ("w:hAnsi", font_name),
                          ("w:cs", font_name), ("w:eastAsia", font_name)):
            rFonts.set(qn(attr), val)
        for attr in (qn("w:asciiTheme"), qn("w:hAnsiTheme"), qn("w:cstheme")):
            if rFonts.get(attr) is not None:
                del rFonts.attrib[attr]
        half_pts = str(int(font_size) * 2)
        for tag in ("w:sz", "w:szCs"):
            el = rPr.find(qn(tag))
            if el is None:
                el = OxmlElement(tag)
                rPr.append(el)
            el.set(qn("w:val"), half_pts)

    for tag in ("w:b", "w:bCs"):
        el = rPr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            rPr.append(el)
        el.set(qn("w:val"), "1" if bold else "0")


def _apply_para_format(para, font_name, font_size, bold, highlight_name=None):
    """Apply formatting to paragraph + all runs."""
    pPr = para._element.get_or_add_pPr()
    rPr = pPr.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        pPr.append(rPr)

    _apply_run_props(rPr, font_name, font_size, bold)

    if highlight_name:
        hl = rPr.find(qn("w:highlight"))
        if hl is None:
            hl = OxmlElement("w:highlight")
            rPr.append(hl)
        hl.set(qn("w:val"), highlight_name)

    for r_elem in para._element.xpath(".//w:r"):
        nested = False
        for tbl in para._element.xpath(".//w:tbl"):
            if tbl.xpath(".//w:r[generate-id() = generate-id(current())]"):
                nested = True
                break
        if nested:
            continue

        r_rPr = r_elem.find(qn("w:rPr"))
        if r_rPr is None:
            r_rPr = OxmlElement("w:rPr")
            r_elem.insert(0, r_rPr)
        _apply_run_props(r_rPr, font_name, font_size, bold)
        if highlight_name:
            hl = r_rPr.find(qn("w:highlight"))
            if hl is None:
                hl = OxmlElement("w:highlight")
                r_rPr.append(hl)
            hl.set(qn("w:val"), highlight_name)


def _apply_format_to_para(para, elem, config):
    """Resolve font/size/bold/highlight for one element and apply it."""
    ptype = elem["type"]
    font_name = config.get(
        f"{ptype.lower()}_font") or config.get("paragraph_font")
    font_size = config.get(f"{ptype.lower()}_size") or config.get(
        "paragraph_size") or 12
    try:
        font_size = int(font_size)
    except Exception:
        font_size = 12

    is_heading = ptype.startswith("HEADING_")
    bold = (is_heading and config.get("bold_headings", True)) or \
           (ptype == "LIST_ITEM" and config.get("bold_lists", False))

    do_highlight = config.get("highlight", False)
    hl_name = _COLOR_MAP.get(HIGHLIGHT_COLORS.get(ptype)
                             ) if do_highlight else None

    _apply_para_format(para, font_name, font_size,
                       bold, highlight_name=hl_name)


def _set_color_on_rPr(rPr, rgb_hex: str):
    """Set w:color on an rPr element to the given 6-char hex string (no #)."""
    color_el = rPr.find(qn("w:color"))
    if color_el is None:
        color_el = OxmlElement("w:color")
        rPr.append(color_el)
    color_el.set(qn("w:val"), rgb_hex.upper())
    # Remove theme colour overrides from the color element itself
    for attr in (qn("w:themeColor"), qn("w:themeTint"), qn("w:themeShade")):
        if color_el.get(attr) is not None:
            del color_el.attrib[attr]
    # Also remove theme colour overrides that may sit directly on rPr —
    # Word honours w:themeColor on rPr even when w:color has an explicit val,
    # which causes the explicit colour to be ignored in the rendered output.
    for attr in (qn("w:themeColor"), qn("w:themeTint"), qn("w:themeShade")):
        if rPr.get(attr) is not None:
            del rPr.attrib[attr]


def ft_format_docx(input_path: str, elements: list, output_path: str, config: dict):
    """Format document."""
    doc = Document(input_path)

    # Body paragraph lookup
    para_map = {
        e["para_idx"]: e
        for e in elements
        if e.get("type") != "TABLE_TEXT" and "para_idx" in e
    }

    # Table-text lookup (IMPORTANT)
    table_text_map = {
        (
            e.get("table_idx"),
            e.get("row_idx"),
            e.get("cell_idx"),
            e.get("para_idx_in_cell"),
        ): e
        for e in elements if e.get("type") == "TABLE_TEXT"
    }

    # Table info (for cover detection)
    tbl_map = {
        e["table_idx"]: e
        for e in elements if e.get("type") == "TABLE"
    }

    # ── PASS 1: Body paragraphs ─────────────────────────────────────
    for idx, para in enumerate(doc.paragraphs):
        if is_in_footer_or_header(para):
            continue

        elem = para_map.get(idx)
        if not elem:
            continue

        _apply_format_to_para(para, elem, config)

    # ── PASS 2: TABLE TEXT (NEW FIX) ───────────────────────────────
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            seen_tc_ids = set()
            for cell_idx, cell in enumerate(row.cells):
                tc_id = id(cell._tc)
                if tc_id in seen_tc_ids:
                    continue
                seen_tc_ids.add(tc_id)

                for para_idx_in_cell, cp in enumerate(cell.paragraphs):
                    if not (cp.text or "").strip():
                        continue

                    elem = table_text_map.get((
                        table_idx,
                        row_idx,
                        cell_idx,
                        para_idx_in_cell,
                    ))

                    if not elem:
                        continue

                    _apply_format_to_para(cp, elem, config)

    # ── PASS 3: Table header background ────────────────────────────
    table_heading_bg = config.get("table_heading_bg")
    if table_heading_bg:
        rgb = table_heading_bg.lstrip("#").upper()

        for tbl_idx, table in enumerate(doc.tables):
            elem = tbl_map.get(tbl_idx)
            on_cover = elem.get("on_cover", False) if elem else False

            if on_cover:
                continue

            try:
                for row_idx, row in enumerate(table.rows):
                    if not _is_table_header_row(table, row_idx):
                        break

                    seen_tc = set()
                    for cell in row.cells:
                        tc_id = id(cell._tc)
                        if tc_id in seen_tc:
                            continue
                        seen_tc.add(tc_id)

                        tc = cell._tc
                        tcPr = tc.find(qn("w:tcPr"))
                        if tcPr is None:
                            tcPr = OxmlElement("w:tcPr")
                            tc.insert(0, tcPr)

                        shd = tcPr.find(qn("w:shd"))
                        if shd is None:
                            shd = OxmlElement("w:shd")
                            tcPr.append(shd)

                        shd.set(qn("w:val"), "clear")
                        shd.set(qn("w:color"), "auto")
                        shd.set(qn("w:fill"), rgb)

                        for attr in (
                            qn("w:themeShade"),
                            qn("w:themeTint"),
                            qn("w:themeFill"),
                            qn("w:themeFillShade"),
                            qn("w:themeFillTint"),
                        ):
                            if shd.get(attr) is not None:
                                del shd.attrib[attr]
            except Exception:
                pass

    # ── PASS 3b: Table header text colour ──────────────────────
    table_heading_text_color = config.get(
        "table_heading_text_color", "").strip()
    if table_heading_text_color:
        rgb = table_heading_text_color.lstrip("#").upper()

        for tbl_idx, table in enumerate(doc.tables):
            elem = tbl_map.get(tbl_idx)
            on_cover = elem.get("on_cover", False) if elem else False
            if on_cover:
                continue

            try:
                for row_idx, row in enumerate(table.rows):
                    if not _is_table_header_row(table, row_idx):
                        break

                    seen_tc = set()
                    for cell in row.cells:
                        tc_id = id(cell._tc)
                        if tc_id in seen_tc:
                            continue
                        seen_tc.add(tc_id)

                        for para in cell.paragraphs:
                            # Apply to paragraph-level rPr
                            pPr = para._element.get_or_add_pPr()
                            p_rPr = pPr.find(qn("w:rPr"))
                            if p_rPr is None:
                                p_rPr = OxmlElement("w:rPr")
                                pPr.append(p_rPr)
                            _set_color_on_rPr(p_rPr, rgb)

                            # Apply to every run
                            for r_elem in para._element.xpath(".//w:r"):
                                r_rPr = r_elem.find(qn("w:rPr"))
                                if r_rPr is None:
                                    r_rPr = OxmlElement("w:rPr")
                                    r_elem.insert(0, r_rPr)
                                _set_color_on_rPr(r_rPr, rgb)
            except Exception:
                pass

    # ── PASS 3c: Paragraph text colour ──────────────────────────
    paragraph_text_color = config.get("paragraph_text_color", "").strip()
    if paragraph_text_color:
        rgb = paragraph_text_color.lstrip("#").upper()
        for idx, para in enumerate(doc.paragraphs):
            if is_in_footer_or_header(para):
                continue
            elem = para_map.get(idx)
            if not elem or elem.get("type") != "PARAGRAPH":
                continue
            try:
                pPr = para._element.get_or_add_pPr()
                p_rPr = pPr.find(qn("w:rPr"))
                if p_rPr is None:
                    p_rPr = OxmlElement("w:rPr")
                    pPr.append(p_rPr)
                _set_color_on_rPr(p_rPr, rgb)
                for r_elem in para._element.xpath(".//w:r"):
                    r_rPr = r_elem.find(qn("w:rPr"))
                    if r_rPr is None:
                        r_rPr = OxmlElement("w:rPr")
                        r_elem.insert(0, r_rPr)
                    _set_color_on_rPr(r_rPr, rgb)
            except Exception:
                pass

    # ── PASS 3d: List items text colour ─────────────────────────
    list_text_color = config.get("list_text_color", "").strip()
    if list_text_color:
        rgb = list_text_color.lstrip("#").upper()
        for idx, para in enumerate(doc.paragraphs):
            if is_in_footer_or_header(para):
                continue
            elem = para_map.get(idx)
            if not elem or elem.get("type") != "LIST_ITEM":
                continue
            try:
                pPr = para._element.get_or_add_pPr()
                p_rPr = pPr.find(qn("w:rPr"))
                if p_rPr is None:
                    p_rPr = OxmlElement("w:rPr")
                    pPr.append(p_rPr)
                _set_color_on_rPr(p_rPr, rgb)
                for r_elem in para._element.xpath(".//w:r"):
                    r_rPr = r_elem.find(qn("w:rPr"))
                    if r_rPr is None:
                        r_rPr = OxmlElement("w:rPr")
                        r_elem.insert(0, r_rPr)
                    _set_color_on_rPr(r_rPr, rgb)
            except Exception:
                pass

    # ── PASS 3e: Table body (non-header) text colour ────────────
    table_body_text_color = config.get("table_body_text_color", "").strip()
    if table_body_text_color:
        rgb = table_body_text_color.lstrip("#").upper()

        for tbl_idx, table in enumerate(doc.tables):
            elem = tbl_map.get(tbl_idx)
            on_cover = elem.get("on_cover", False) if elem else False
            if on_cover:
                continue

            try:
                total_rows = len(table.rows)
                for row_idx, row in enumerate(table.rows):
                    # Skip header rows — only colour body rows
                    if _is_table_header_row(table, row_idx):
                        continue

                    seen_tc = set()
                    for cell in row.cells:
                        tc_id = id(cell._tc)
                        if tc_id in seen_tc:
                            continue
                        seen_tc.add(tc_id)

                        for para in cell.paragraphs:
                            pPr = para._element.get_or_add_pPr()
                            p_rPr = pPr.find(qn("w:rPr"))
                            if p_rPr is None:
                                p_rPr = OxmlElement("w:rPr")
                                pPr.append(p_rPr)
                            _set_color_on_rPr(p_rPr, rgb)

                            for r_elem in para._element.xpath(".//w:r"):
                                r_rPr = r_elem.find(qn("w:rPr"))
                                if r_rPr is None:
                                    r_rPr = OxmlElement("w:rPr")
                                    r_elem.insert(0, r_rPr)
                                _set_color_on_rPr(r_rPr, rgb)
            except Exception:
                pass

    # ── PASS 3f: Table header text font & size ──────────────────
    table_heading_font = config.get("table_heading_font", "").strip()
    table_heading_font_size_raw = config.get("table_heading_font_size")
    try:
        table_heading_font_size = int(
            table_heading_font_size_raw) if table_heading_font_size_raw else None
    except (ValueError, TypeError):
        table_heading_font_size = None

    if table_heading_font or table_heading_font_size:
        for tbl_idx, table in enumerate(doc.tables):
            elem = tbl_map.get(tbl_idx)
            on_cover = elem.get("on_cover", False) if elem else False
            if on_cover:
                continue

            try:
                for row_idx, row in enumerate(table.rows):
                    if not _is_table_header_row(table, row_idx):
                        break

                    seen_tc = set()
                    for cell in row.cells:
                        tc_id = id(cell._tc)
                        if tc_id in seen_tc:
                            continue
                        seen_tc.add(tc_id)

                        for para in cell.paragraphs:
                            # Apply to paragraph-level rPr
                            pPr = para._element.get_or_add_pPr()
                            p_rPr = pPr.find(qn("w:rPr"))
                            if p_rPr is None:
                                p_rPr = OxmlElement("w:rPr")
                                pPr.append(p_rPr)

                            if table_heading_font:
                                rFonts = p_rPr.find(qn("w:rFonts"))
                                if rFonts is None:
                                    rFonts = OxmlElement("w:rFonts")
                                    p_rPr.insert(0, rFonts)
                                for attr, val in (
                                    ("w:ascii", table_heading_font),
                                    ("w:hAnsi", table_heading_font),
                                    ("w:cs", table_heading_font),
                                    ("w:eastAsia", table_heading_font),
                                ):
                                    rFonts.set(qn(attr), val)
                                for attr in (qn("w:asciiTheme"), qn("w:hAnsiTheme"), qn("w:cstheme")):
                                    if rFonts.get(attr) is not None:
                                        del rFonts.attrib[attr]

                            if table_heading_font_size:
                                half_pts = str(table_heading_font_size * 2)
                                for tag in ("w:sz", "w:szCs"):
                                    el = p_rPr.find(qn(tag))
                                    if el is None:
                                        el = OxmlElement(tag)
                                        p_rPr.append(el)
                                    el.set(qn("w:val"), half_pts)

                            # Apply to every run in the cell paragraph
                            for r_elem in para._element.xpath(".//w:r"):
                                r_rPr = r_elem.find(qn("w:rPr"))
                                if r_rPr is None:
                                    r_rPr = OxmlElement("w:rPr")
                                    r_elem.insert(0, r_rPr)

                                if table_heading_font:
                                    rFonts = r_rPr.find(qn("w:rFonts"))
                                    if rFonts is None:
                                        rFonts = OxmlElement("w:rFonts")
                                        r_rPr.insert(0, rFonts)
                                    for attr, val in (
                                        ("w:ascii", table_heading_font),
                                        ("w:hAnsi", table_heading_font),
                                        ("w:cs", table_heading_font),
                                        ("w:eastAsia", table_heading_font),
                                    ):
                                        rFonts.set(qn(attr), val)
                                    for attr in (qn("w:asciiTheme"), qn("w:hAnsiTheme"), qn("w:cstheme")):
                                        if rFonts.get(attr) is not None:
                                            del rFonts.attrib[attr]

                                if table_heading_font_size:
                                    half_pts = str(table_heading_font_size * 2)
                                    for tag in ("w:sz", "w:szCs"):
                                        el = r_rPr.find(qn(tag))
                                        if el is None:
                                            el = OxmlElement(tag)
                                            r_rPr.append(el)
                                        el.set(qn("w:val"), half_pts)
            except Exception:
                pass

    doc.save(output_path)

# =============================================================================
#  HELPER
# =============================================================================


def _is_dark_color(hex_color: str) -> bool:
    """Return True if hex color (no #) is dark enough to warrant white text."""
    try:
        r = int(hex_color[0:2], 16)
        g = int(hex_color[2:4], 16)
        b = int(hex_color[4:6], 16)
        return (0.299 * r + 0.587 * g + 0.114 * b) / 255 < 0.5
    except Exception:
        return True

# =============================================================================
#  HTML PREVIEW  — PDF-style layout (white page on grey background)
# =============================================================================


def docx_to_html_preview(docx_path: str, elements: list = None, highlight: bool = False) -> str:
    """Convert docx to HTML with PDF-style preview and optional element highlighting."""
    import html as _html_mod
    doc = Document(docx_path)

    # Build element type map by paragraph index
    para_type_map = {}
    if elements:
        for e in elements:
            if "para_idx" in e:
                para_type_map[e["para_idx"]] = e["type"]

    # ✅ FIX: Solid pastel background colors — no opacity, no white wash
    HIGHLIGHT_MAP = {
        "HEADING_1": "#00ff00",
        "HEADING_2": "#00ffff",
        "HEADING_3": "#ff00ff",
        "HEADING_4": "#800080",
        "HEADING_5": "#e67e22",
        "HEADING_6": "#e74c3c",
        "PARAGRAPH": None,        # paragraphs are never highlighted
        "LIST_ITEM": "#ffff00",
        "TABLE_TEXT": "#ff0000",
    }

    # ✅ PDF-style layout matching alignment.py
    css = (
        "<!DOCTYPE html><html><head><meta charset='utf-8'><style>"
        "*, *::before, *::after{box-sizing:border-box;margin:0;padding:0;}"
        "body{font-family:Arial,sans-serif;margin:0;padding:0;background:#e8e8e8;}"
        ".page{background:white;width:794px;min-height:1123px;margin:30px auto;"
        "padding:72px 80px;box-shadow:0 2px 12px rgba(0,0,0,0.18);}"
        "h1,h2,h3,h4,h5,h6{margin:0.4em 0 0.2em;line-height:1.3;}"
        "h1{font-size:24px;} h2{font-size:20px;} h3{font-size:18px;}"
        "h4{font-size:16px;} h5{font-size:14px;} h6{font-size:13px;}"
        "p{margin:0.3em 0;line-height:1.5;font-size:13px;}"
        "ul,ol{margin:0.5em 0 0.5em 1.8em;padding:0;}"
        "li{margin:0.2em 0;line-height:1.5;}"
        "table{border-collapse:collapse;width:100%;margin:0.8em 0;}"
        "td,th{border:1px solid #c8c8c8;padding:6px 10px;vertical-align:top;font-size:0.95em;}"
        "tr:nth-child(even) td{background:#f9f9f9;}"
        ".hr{border-top:2px dashed #999;margin:20px 0;}"
        "</style></head><body><div class='page'>"
    )

    parts = [css]
    para_map = {id(p._element): p for p in doc.paragraphs}
    table_map = {id(t._tbl): t for t in doc.tables}
    para_seq_idx = {id(p._element): i for i, p in enumerate(doc.paragraphs)}

    def render_runs(para, bg_color=None):
        out = []
        for run in para.runs:
            t = _html_mod.escape(run.text or "")
            if not t:
                continue
            rPr = run._r.find(qn("w:rPr"))
            bold = ital = False
            if rPr is not None:
                if rPr.find(qn("w:b")) is not None:
                    bold = True
                if rPr.find(qn("w:i")) is not None:
                    ital = True
            if bold:
                t = f"<strong>{t}</strong>"
            if ital:
                t = f"<em>{t}</em>"
            # ✅ FIX: solid inline highlight, no opacity layer
            if bg_color and highlight:
                t = f"<span style='background:{bg_color};border-radius:2px;padding:0 1px;'>{t}</span>"
            out.append(t)
        return "".join(out)

    para_counter = 0
    for child in doc.element.body:
        if child.tag == qn("w:p"):
            para = para_map.get(id(child))
            if para is None or is_in_footer_or_header(para):
                para_counter += 1
                continue

            text = (para.text or "").strip()
            seq_idx = para_seq_idx.get(id(para._element), -1)
            ptype = para_type_map.get(seq_idx, "")
            bg_color = HIGHLIGHT_MAP.get(
                ptype) if highlight and ptype else None

            # Detect heading level from style
            sn = para.style.name if para.style else ""
            htag = "p"
            for i in range(1, 7):
                if f"Heading {i}" in sn:
                    htag = f"h{i}"
                    break

            if not text:
                parts.append("<p>&nbsp;</p>")
                para_counter += 1
                continue

            inner = render_runs(para, bg_color)

            if has_page_break_before(para):
                parts.append('<div class="hr"></div>')

            # Block-level background for non-heading types
            if bg_color and highlight:
                parts.append(
                    f"<{htag} style='background:{bg_color};'>{inner}</{htag}>")
            else:
                parts.append(f"<{htag}>{inner}</{htag}>")
            para_counter += 1

        elif child.tag == qn("w:tbl"):
            table = table_map.get(id(child))
            if table is None:
                continue

            rows = ['<table>']
            for row_idx, row in enumerate(table.rows):
                is_header = _is_table_header_row(table, row_idx)
                rows.append("<tr>")

                seen_tc = set()
                for cell in row.cells:
                    tc_id = id(cell._tc)
                    if tc_id in seen_tc:
                        continue
                    seen_tc.add(tc_id)

                    cell_parts = []
                    for cp in cell.paragraphs:
                        if is_in_footer_or_header(cp):
                            continue
                        # Pass TABLE_TEXT highlight colour into cell runs
                        cell_bg = HIGHLIGHT_MAP.get(
                            "TABLE_TEXT") if highlight else None
                        cell_parts.append(render_runs(cp, cell_bg))

                    content = "<br>".join(
                        p for p in cell_parts if p) or "&nbsp;"
                    tag = "th" if is_header else "td"

                    # Preserve cell background color
                    style = ""
                    tcPr = cell._tc.find(qn("w:tcPr"))
                    if tcPr is not None:
                        shd = tcPr.find(qn("w:shd"))
                        if shd is not None:
                            fill = shd.get(qn("w:fill"), "").upper()
                            if fill and fill not in ("AUTO", "FFFFFF", ""):
                                text_col = "white" if _is_dark_color(
                                    fill) else "inherit"
                                style = f' style="background:#{fill};color:{text_col};font-weight:600;"'

                    rows.append(f"<{tag}{style}>{content}</{tag}>")
                rows.append("</tr>")

            rows.append("</table>")
            parts.append("".join(rows))

    parts.append("</div></body></html>")
    return "".join(parts)

# =============================================================================
#  FLASK ROUTES
# =============================================================================


@font_bp.route("/")
def home():
    return render_template("fontUI.html")


@font_bp.route("/analyse", methods=["POST"])
def ft_analyse():
    path = _get_path()
    if not path or not os.path.exists(path):
        return jsonify({"error": "No file in session"}), 400
    try:
        data = ft_analyze_document_structure(path)
        data["standard_table_bgs"] = STANDARD_TABLE_BG_COLORS
        return jsonify(data)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@font_bp.route("/original", methods=["POST"])
def ft_original():
    path = _get_path()
    if not path or not os.path.exists(path):
        return jsonify({"error": "No file in session"}), 400
    try:
        if IS_WINDOWS and WORD_AVAILABLE:
            try:
                pythoncom.CoInitialize()
                pdf_path = os.path.join(
                    _session_out_dir(), f"orig_{uuid.uuid4().hex}.pdf")
                convert(path, pdf_path)
                pythoncom.CoUninitialize()
                return send_file(pdf_path, mimetype="application/pdf")
            except Exception:
                try:
                    pythoncom.CoUninitialize()
                except Exception:
                    pass
        html = docx_to_html_preview(path, elements=None, highlight=False)
        return html, 200, {"Content-Type": "text/html; charset=utf-8"}
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@font_bp.route("/preview", methods=["POST"])
def ft_preview():
    path = _get_path()
    if not path or not os.path.exists(path):
        return jsonify({"error": "No file in session"}), 400
    try:
        cfg = json.loads(request.form.get("config", "{}"))
        out_docx = os.path.join(
            _session_out_dir(), f"ft_preview_{uuid.uuid4().hex}.docx")
        data = ft_analyze_document_structure(path)
        # ft_format_docx now bakes highlight colours into the docx runs
        # when cfg["highlight"] is True, so the PDF render picks them up too.
        ft_format_docx(path, data["elements"], out_docx, cfg)

        if IS_WINDOWS and WORD_AVAILABLE:
            pdf_path = os.path.join(
                _session_out_dir(), f"ft_preview_{uuid.uuid4().hex}.pdf")
            pythoncom.CoInitialize()
            try:
                convert(out_docx, pdf_path)
            finally:
                pythoncom.CoUninitialize()
            return send_file(pdf_path, mimetype="application/pdf")

        # Non-Windows or Word not installed — fall back to HTML renderer
        highlight = cfg.get("highlight", False)
        html = docx_to_html_preview(
            out_docx, elements=data["elements"], highlight=highlight)
        return html, 200, {"Content-Type": "text/html; charset=utf-8"}
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@font_bp.route("/format", methods=["POST"])
def ft_format():
    path = _get_path()
    if not path or not os.path.exists(path):
        return jsonify({"error": "No file in session"}), 400
    try:
        cfg = json.loads(request.form.get("config", "{}"))
        cfg["highlight"] = False
        out_path = os.path.join(
            _session_out_dir(), f"font_formatted_{uuid.uuid4().hex}.docx")
        data = ft_analyze_document_structure(path)
        ft_format_docx(path, data["elements"], out_path, cfg)
        _set_path(out_path)
        return send_file(out_path, as_attachment=True, download_name="font_formatted.docx")
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# =============================================================================
#  STANDALONE
# =============================================================================


if __name__ == "__main__":
    app = Flask(__name__, template_folder="templates")
    app.secret_key = "dev"
    app.register_blueprint(font_bp)

    @app.route("/")
    def root():
        from flask import redirect
        return redirect("/font/")

    print("Font Formatter → http://127.0.0.1:5001/font/")
    app.run(debug=True, port=5001, threaded=False)
